#!/usr/bin/env python3
"""Integrate dry-lab partial causal closure results into locked manuscript text."""

from __future__ import annotations

import csv
import re
from pathlib import Path

import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"

MAIN_MD = OUT / "main_text_clean_version.md"
RESULTS_MD = OUT / "results_final_rewritten.md"
DISCUSSION_MD = OUT / "discussion_final_rewritten.md"
SOURCE_DOCX = OUT / "submission_ready_manuscript.docx"
UPDATED_DOCX = OUT / "submission_ready_manuscript_causal_consistency_updated.docx"

REQUIRED_FILES = {
    "FIGURE_CAUSAL_CLOSURE_FINAL.png": OUT / "FIGURE_CAUSAL_CLOSURE_FINAL.png",
    "FIGURE_CAUSAL_CLOSURE_FINAL.pdf": OUT / "FIGURE_CAUSAL_CLOSURE_FINAL.pdf",
    "final_closure_state_report.md": OUT / "final_closure_state_report.md",
    "closure_model_comparison.tsv": OUT / "closure_model_comparison.tsv",
    "regime_conditioned_deltaZ_consistency.tsv": OUT / "regime_conditioned_deltaZ_consistency.tsv",
    "causal_closure_score_summary.tsv": OUT / "causal_closure_score_summary.tsv",
    "counterfactual_bias_by_regime.tsv": OUT / "counterfactual_bias_by_regime.tsv",
    "edge_stability_across_regimes.tsv": OUT / "edge_stability_across_regimes.tsv",
    "perturbation_deltaZ.tsv": OUT / "perturbation_deltaZ.tsv",
    "counterfactual_consistency_scores.tsv": OUT / "counterfactual_consistency_scores.tsv",
    "null_perturbation_control_results.tsv": OUT / "null_perturbation_control_results.tsv",
    "cross_dataset_perturbation_similarity.tsv": OUT / "cross_dataset_perturbation_similarity.tsv",
    "W_matrix_by_regime.tsv": OUT / "W_matrix_by_regime.tsv",
    "regime_conditioned_grn_summary.tsv": OUT / "regime_conditioned_grn_summary.tsv",
}

FORBIDDEN_PHRASES = [
    "causal proof",
    "causal closure established",
    "perturbations drive",
    "experimentally validated",
    "W(Z) demonstrates causality",
    "causal mechanism confirmed",
    "full causal closure",
]


def read_tsv(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path, sep="\t")


def get_summary_values() -> dict[str, float | str]:
    closure = read_tsv(OUT / "closure_model_comparison.tsv")
    ccs = read_tsv(OUT / "causal_closure_score_summary.tsv")
    counter = read_tsv(OUT / "counterfactual_consistency_scores.tsv")
    vals: dict[str, float | str] = {}
    if not closure.empty:
        row = closure.iloc[0]
        vals["global_mse"] = float(row["global_MSE"])
        vals["conditioned_mse"] = float(row["regime_conditioned_MSE_LOO"])
        vals["error_reduction"] = float(row["prediction_error_reduction_fraction"])
        vals["closure_class"] = str(row["final_closure_classification"])
        vals["reversal_frequency"] = float(row["counterfactual_direction_reversal_frequency"])
        vals["instability_ratio"] = float(row["cross_regime_instability_ratio"])
    if not ccs.empty:
        for pathway in ["RA", "BMP", "NOTCH", "FGF", "SHH"]:
            sub = ccs[ccs["pathway"].eq(pathway)]
            if len(sub):
                vals[f"{pathway}_ccs"] = float(sub.iloc[0]["CCS_available_evidence_rescaled"])
    if not counter.empty:
        for pathway in ["RA", "BMP", "NOTCH", "SHH"]:
            sub = counter[counter["pathway"].eq(pathway)]
            if len(sub):
                vals[f"{pathway}_cos"] = float(sub.iloc[0]["observed_vs_counterfactual_cosine"])
    return vals


def causal_results_paragraph(vals: dict[str, float | str]) -> str:
    return (
        "A supplementary dry-lab perturbation-consistency audit further tested whether posterior regime-probability shifts were "
        "compatible with the learned regulatory structure without changing the nine main figures. A global W_GRN-based "
        f"counterfactual model was insufficient, whereas regime-conditioned W(Z) reduced the mean squared error from "
        f"{vals.get('global_mse', float('nan')):.3g} to {vals.get('conditioned_mse', float('nan')):.3g} "
        f"(error-reduction proxy {vals.get('error_reduction', float('nan')):.3g}; Supplementary Fig. 1 and Supplementary Tables 1-5). "
        f"The improvement was partial rather than complete: RA and BMP showed the strongest perturbation-consistency support "
        f"(CCS {vals.get('RA_ccs', float('nan')):.3g} and {vals.get('BMP_ccs', float('nan')):.3g}), while NOTCH was limited by "
        "single-member pathway coverage and SHH/FGF showed weaker directional support. Counterfactual direction reversals "
        f"remained frequent ({vals.get('reversal_frequency', float('nan')):.2g}), so these outputs support partial dry-lab "
        "perturbation consistency in regime-conditioned subspaces, not FULL_CLOSURE."
    )


def discussion_limitations_paragraph(vals: dict[str, float | str]) -> str:
    return (
        "The perturbation-consistency layer also separates representation closure from causal closure. The latent-state-regime "
        "mixture is strongly supported as a representation of the processed state space, but the dry-lab perturbation audit "
        "supports only local, regime-conditioned causal consistency. Regime-conditioned W(Z) improved the reconstruction of "
        f"posterior shifts, yet cross-regime instability remained high (instability ratio {vals.get('instability_ratio', float('nan')):.3g}) "
        f"and counterfactual sign reversals persisted ({vals.get('reversal_frequency', float('nan')):.2g}). A complete causal-closure "
        "claim therefore requires direct wet-lab perturbation time series or matched perturbation-multiome validation in the same "
        "regime contexts. The present study supports a computational causal-consistency layer, not wet-lab causal validation."
    )


def supplementary_caption(vals: dict[str, float | str]) -> str:
    return (
        "# Supplementary Figure 1 | Dry-lab perturbation consistency and partial causal closure of regime-conditioned dynamics.\n\n"
        "Existing perturbation-derived outputs were used to test whether pathway-associated perturbations were directionally "
        "consistent with changes in latent-state-regime posterior probabilities. Panel A summarizes pathway-level ΔZ posterior shifts "
        "for RA, BMP, NOTCH, FGF and SHH contrasts. Panel B compares observed ΔZ vectors with counterfactual ΔZ vectors derived "
        "from the locked W_GRN axis effects. Panel C reports cross-dataset perturbation consistency; direct pathway-member "
        "posterior shifts were available from one direct perturbation dataset and are therefore marked as not computable for "
        "cross-dataset invariance rather than inferred from pseudo-replicates. Panel D compares real pathway-associated shifts "
        "with shuffled-label null perturbations. Panel E reports the causal closure score for available evidence, and Panel F "
        "shows the evidence-weight fraction. The resulting classification was PARTIAL_CLOSURE, not FULL_CLOSURE: regime-conditioned "
        f"W(Z) reduced reconstruction error relative to the global W model, but counterfactual reversals remained frequent "
        f"({vals.get('reversal_frequency', float('nan')):.2g}) and cross-regime instability remained high. These analyses support "
        "dry-lab perturbation-consistent structure in regime-conditioned subspaces, not wet-lab causal validation."
    )


def insert_after(text: str, anchor: str, insert: str) -> str:
    if insert in text:
        return text
    idx = text.find(anchor)
    if idx == -1:
        raise ValueError(f"Anchor not found: {anchor[:80]}")
    end = idx + len(anchor)
    return text[:end] + "\n\n" + insert + text[end:]


def update_markdowns(vals: dict[str, float | str]) -> tuple[str, str]:
    results_insert = causal_results_paragraph(vals)
    discussion_insert = discussion_limitations_paragraph(vals)

    results = RESULTS_MD.read_text()
    result_anchor = (
        "At the same time, overlap does not imply identity. Symmetrized KL divergence shows regime-specific distributional "
        "differences, supporting adult_repair and salamander_blastema as overlapping but distinct latent-state-regime mixtures. "
        "Figure 9 closes the Results by showing distributional overlap without scalar separability."
    )
    results_updated = insert_after(results, result_anchor, results_insert)
    (OUT / "results_with_causal_consistency_insert.md").write_text(results_updated + ("\n" if not results_updated.endswith("\n") else ""))

    discussion = DISCUSSION_MD.read_text()
    discussion_anchor = (
        "Several limitations remain. The analysis is locked at processed state-score and proxy layers, with cross-species alignment, "
        "batch/species entanglement and orthology constraints. It does not establish complete causal regulatory circuitry, direct "
        "chromatin memory, lineage history or spatial morphogen control. Those claims require same-system perturbation time series, "
        "lineage tracing, spatial measurements and multiome validation."
    )
    discussion_updated = insert_after(discussion, discussion_anchor, discussion_insert)
    (OUT / "discussion_with_causal_limitations.md").write_text(discussion_updated + ("\n" if not discussion_updated.endswith("\n") else ""))

    main = MAIN_MD.read_text()
    main_updated = insert_after(main, result_anchor, results_insert)
    main_updated = insert_after(main_updated, discussion_anchor, discussion_insert)
    (OUT / "main_text_clean_version_causal_updated.md").write_text(main_updated + ("\n" if not main_updated.endswith("\n") else ""))
    return results_insert, discussion_insert


def sentence_split(text: str) -> list[str]:
    # Conservative split for audit; keeps abbreviations acceptable for phrase scans.
    parts = re.split(r"(?<=[.!?])\s+", text.replace("\n", " "))
    return [p.strip() for p in parts if p.strip()]


def safe_replacement(sentence: str) -> str:
    s = sentence
    repl = {
        "causal proof": "dry-lab causal-consistency support",
        "causal closure established": "partial computational causal closure was observed",
        "perturbations drive": "perturbations are associated with",
        "experimentally validated": "wet-lab causal validation",
        "W(Z) demonstrates causality": "W(Z) is compatible with regime-conditioned perturbation consistency",
        "causal mechanism confirmed": "causal mechanism remains computationally supported but not experimentally confirmed",
    }
    for bad, good in repl.items():
        s = re.sub(re.escape(bad), good, s, flags=re.IGNORECASE)
    return s


def claim_safety_audit() -> None:
    files = [
        OUT / "main_text_clean_version_causal_updated.md",
        OUT / "results_with_causal_consistency_insert.md",
        OUT / "discussion_with_causal_limitations.md",
        OUT / "supplementary_causal_closure_caption.md",
        OUT / "supplementary_causal_closure_package_plan.md",
    ]
    rows = []
    for path in files:
        text = path.read_text() if path.exists() else ""
        for sent in sentence_split(text):
            lower = sent.lower()
            hits = [p for p in FORBIDDEN_PHRASES if p.lower() in lower]
            if hits:
                rows.append({
                    "file": path.name,
                    "forbidden_phrase": "; ".join(hits),
                    "flagged_sentence": sent,
                    "safe_replacement": safe_replacement(sent),
                })

    lines = ["# Causal Claim Safety Audit", ""]
    if not rows:
        lines += [
            "No exact forbidden causal-overclaim phrases were detected in the updated manuscript text.",
            "",
            "Forbidden phrases scanned:",
            "",
        ]
        lines += [f"- {p}" for p in FORBIDDEN_PHRASES]
        lines += [
            "",
            "Interpretation lock: the updated manuscript uses partial dry-lab perturbation consistency and partial computational causal closure language, not COMPLETE_CLOSURE or wet-lab causal validation.",
        ]
    else:
        lines += ["| file | forbidden phrase | flagged sentence | safe replacement |", "| --- | --- | --- | --- |"]
        for r in rows:
            lines.append(
                "| {file} | {forbidden_phrase} | {flagged_sentence} | {safe_replacement} |".format(
                    **{k: str(v).replace("|", "/") for k, v in r.items()}
                )
            )
    (OUT / "causal_claim_safety_audit.md").write_text("\n".join(lines) + "\n")


def integration_plan() -> None:
    rows = [
        ("FIGURE_CAUSAL_CLOSURE_FINAL.png", "Supplementary Figure", "Supplementary Fig. 1 raster; not part of locked 9 main figures"),
        ("FIGURE_CAUSAL_CLOSURE_FINAL.pdf", "Supplementary Figure", "Vector/PDF companion for Supplementary Fig. 1"),
        ("final_closure_state_report.md", "Supplementary Methods", "Narrative audit for PARTIAL_CLOSURE classification and claim boundary"),
        ("closure_model_comparison.tsv", "Supplementary Table", "H1 global W versus H2 regime-conditioned W(Z) comparison"),
        ("regime_conditioned_deltaZ_consistency.tsv", "Supplementary Table", "Regime-conditioned observed/counterfactual ΔZ consistency"),
        ("causal_closure_score_summary.tsv", "Supplementary Table", "Pathway-level causal closure score summary"),
        ("counterfactual_bias_by_regime.tsv", "Supplementary Table", "Regime-specific counterfactual bias and reversal frequency"),
        ("edge_stability_across_regimes.tsv", "Source Data", "Edge-level stability across W(Z) matrices"),
        ("perturbation_deltaZ.tsv", "Source Data", "Pathway and target-level ΔZ posterior shifts"),
        ("counterfactual_consistency_scores.tsv", "Source Data", "Observed versus counterfactual ΔZ alignment"),
        ("null_perturbation_control_results.tsv", "Source Data", "Shuffled-label null perturbation control"),
        ("cross_dataset_perturbation_similarity.tsv", "Source Data", "Cross-dataset invariance status; not computable for single direct dataset"),
        ("W_matrix_by_regime.tsv", "Source Data", "Regime-conditioned W(Z) expected edge matrix"),
        ("regime_conditioned_grn_summary.tsv", "Supplementary Table", "Regime-conditioned GRN module summary"),
        ("regime_axis_conditioning_gates.tsv", "Methods", "Regime posterior state-context gates used for W(Z) expectation"),
        ("pathway_response_conditioning_gates.tsv", "Methods", "Pathway ΔZ response gates used for W(Z) expectation"),
    ]
    with (OUT / "supplementary_causal_closure_integration_plan.tsv").open("w", newline="") as fh:
        writer = csv.writer(fh, delimiter="\t")
        writer.writerow(["file", "destination", "integration_note"])
        for row in rows:
            writer.writerow(row)


def package_plan(vals: dict[str, float | str]) -> None:
    text = f"""# Supplementary Causal Closure Package Plan

## Placement

The dry-lab perturbation consistency analysis is integrated as Supplementary Fig. 1 and Supplementary Tables 1-5. It does not change the locked 9-main-figure architecture and does not create a Figure 10.

## Claim Boundary

Use: perturbation-consistent structure, partial computational causal closure, regime-conditioned perturbation consistency, dry-lab causal consistency.

Do not claim: COMPLETE_CLOSURE, wet-lab causal validation, direct perturbation-driven regime transitions, or W(Z)-based causal demonstration.

## Numerical Lock

- Closure classification: `{vals.get('closure_class', 'PARTIAL_CLOSURE')}`
- Global W MSE: `{vals.get('global_mse', float('nan')):.4g}`
- Regime-conditioned W(Z) MSE: `{vals.get('conditioned_mse', float('nan')):.4g}`
- Error-reduction proxy: `{vals.get('error_reduction', float('nan')):.4g}`
- RA CCS: `{vals.get('RA_ccs', float('nan')):.4g}`
- BMP CCS: `{vals.get('BMP_ccs', float('nan')):.4g}`
- Counterfactual direction reversal frequency: `{vals.get('reversal_frequency', float('nan')):.4g}`

## Required Manuscript Sentence

The analysis supports partial dry-lab perturbation consistency in regime-conditioned subspaces, not FULL_CLOSURE.
"""
    (OUT / "supplementary_causal_closure_package_plan.md").write_text(text)


def insert_docx_paragraph_after(doc: Document, anchor: str, paragraph: str) -> bool:
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == paragraph.strip():
            return True
        if para.text.strip() == anchor.strip():
            new_para = para.insert_paragraph_before("")
            # Move newly inserted paragraph after anchor by XML manipulation.
            p = para._p
            new_p = new_para._p
            p.addnext(new_p)
            new_para.style = para.style
            new_para.add_run(paragraph)
            return True
    return False


def update_docx(results_insert: str, discussion_insert: str) -> None:
    doc = Document(str(SOURCE_DOCX))
    result_anchor = (
        "At the same time, overlap does not imply identity. Symmetrized KL divergence shows regime-specific distributional differences, "
        "supporting adult_repair and salamander_blastema as overlapping but distinct latent-state-regime mixtures. Figure 9 closes the Results "
        "by showing distributional overlap without scalar separability."
    )
    discussion_anchor = (
        "Several limitations remain. The analysis is locked at processed state-score and proxy layers, with cross-species alignment, "
        "batch/species entanglement and orthology constraints. It does not establish complete causal regulatory circuitry, direct chromatin "
        "memory, lineage history or spatial morphogen control. Those claims require same-system perturbation time series, lineage tracing, "
        "spatial measurements and multiome validation."
    )
    if not insert_docx_paragraph_after(doc, result_anchor, results_insert):
        raise RuntimeError("Could not locate Results anchor in DOCX")
    if not insert_docx_paragraph_after(doc, discussion_anchor, discussion_insert):
        raise RuntimeError("Could not locate Discussion anchor in DOCX")
    doc.save(str(UPDATED_DOCX))


def main() -> None:
    missing = [name for name, path in REQUIRED_FILES.items() if not path.exists()]
    if missing:
        print("warning: missing optional inputs:", ", ".join(missing))
    vals = get_summary_values()
    results_insert, discussion_insert = update_markdowns(vals)
    package_plan(vals)
    (OUT / "supplementary_causal_closure_caption.md").write_text(supplementary_caption(vals) + "\n")
    integration_plan()
    claim_safety_audit()
    update_docx(results_insert, discussion_insert)
    print("manuscript causal consistency integration complete")
    print(UPDATED_DOCX)


if __name__ == "__main__":
    main()
