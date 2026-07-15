from __future__ import annotations

import html
import shutil
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
PKG = ROOT / "FINAL_SUBMISSION_PACKAGE"

FIG_PNG = OUT / "MODEL_EVOLUTION_MAP_Figure1A.png"
FIG_PDF = OUT / "MODEL_EVOLUTION_MAP_Figure1A.pdf"
FIG_SVG = OUT / "MODEL_EVOLUTION_MAP_Figure1A.svg"
SECTION_MD = OUT / "model_lineage_section.md"
CAPTION_MD = OUT / "model_evolution_map_caption.md"
UPDATED_MD = OUT / "final_claim_safe_main_text_with_model_lineage.md"
UPDATED_DOCX = OUT / "FINAL_MANUSCRIPT_WITH_MODEL_EVOLUTION_MAP.docx"
UPDATED_LEGENDS = OUT / "FINAL_FIGURE_LEGENDS_WITH_MODEL_EVOLUTION.docx"

BASE_DOCX = OUT / "FINAL_MANUSCRIPT_WITH_MAIN_FIGURES.docx"
BASE_MD = OUT / "final_claim_safe_main_text.md"
BASE_LEGENDS = OUT / "FINAL_FIGURE_LEGENDS.docx"

W, H = 5400, 2500
MARGIN = 130

COL = {
    "ink": "#111827",
    "muted": "#5F6368",
    "blue": "#1F4D78",
    "old_bg": "#FEE2E2",
    "old_edge": "#B91C1C",
    "fail_bg": "#FFF7ED",
    "fail_edge": "#C2410C",
    "trans_bg": "#FEF3C7",
    "trans_edge": "#92400E",
    "new_bg": "#E0F2FE",
    "new_edge": "#075985",
    "claim_bg": "#DCFCE7",
    "claim_edge": "#166534",
    "card": "#F8FAFC",
    "border": "#CBD5E1",
}


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Helvetica.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for p in candidates:
        if Path(p).exists():
            return ImageFont.truetype(p, size)
    return ImageFont.load_default()


F = {
    "title": font(58, True),
    "subtitle": font(30),
    "stage": font(30, True),
    "body": font(28),
    "small": font(24),
    "tiny": font(21),
    "badge": font(28, True),
}


class SVG:
    def __init__(self, w: int, h: int):
        self.w = w
        self.h = h
        self.items = [
            f'<svg xmlns="http://www.w3.org/2000/svg" width="{w}" height="{h}" viewBox="0 0 {w} {h}">',
            '<rect width="100%" height="100%" fill="white"/>',
        ]

    def rect(self, x, y, w, h, fill, stroke, sw=3, rx=28):
        self.items.append(
            f'<rect x="{x:.1f}" y="{y:.1f}" width="{w:.1f}" height="{h:.1f}" rx="{rx:.1f}" fill="{fill}" stroke="{stroke}" stroke-width="{sw}"/>'
        )

    def text(self, x, y, text, size=24, fill="#111827", weight="normal", anchor="start"):
        self.items.append(
            f'<text x="{x:.1f}" y="{y:.1f}" font-family="Arial, Helvetica, sans-serif" font-size="{size}" font-weight="{weight}" fill="{fill}" text-anchor="{anchor}">{html.escape(str(text))}</text>'
        )

    def line(self, x1, y1, x2, y2, stroke="#334155", sw=5, arrow=False):
        marker = ' marker-end="url(#arrow)"' if arrow else ""
        self.items.append(f'<line x1="{x1:.1f}" y1="{y1:.1f}" x2="{x2:.1f}" y2="{y2:.1f}" stroke="{stroke}" stroke-width="{sw}"{marker}/>')

    def defs(self):
        self.items.insert(
            2,
            '<defs><marker id="arrow" markerWidth="12" markerHeight="12" refX="10" refY="6" orient="auto"><path d="M2,2 L10,6 L2,10 Z" fill="#334155"/></marker></defs>',
        )

    def save(self, path: Path):
        self.defs()
        self.items.append("</svg>")
        path.write_text("\n".join(self.items), encoding="utf-8")


def rgb(h: str) -> tuple[int, int, int]:
    h = h.lstrip("#")
    return tuple(int(h[i : i + 2], 16) for i in (0, 2, 4))


def wrap(text: str, chars: int) -> list[str]:
    words = text.split()
    lines, cur = [], []
    for w in words:
        if cur and len(" ".join(cur + [w])) > chars:
            lines.append(" ".join(cur))
            cur = [w]
        else:
            cur.append(w)
    if cur:
        lines.append(" ".join(cur))
    return lines


def draw_text(draw: ImageDraw.ImageDraw, xy, text, fnt, fill=COL["ink"], anchor=None):
    draw.text(xy, str(text), font=fnt, fill=fill, anchor=anchor)


def rounded(draw: ImageDraw.ImageDraw, box, fill, outline, width=4, radius=36):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw: ImageDraw.ImageDraw, x1, y1, x2, y2):
    draw.line([x1, y1, x2, y2], fill=rgb("#334155"), width=7)
    # Simple triangular head.
    draw.polygon([(x2, y2), (x2 - 28, y2 - 16), (x2 - 28, y2 + 16)], fill=rgb("#334155"))


def card(draw, svg: SVG, x, y, w, h, title, bullets, face, edge, footer=None):
    rounded(draw, [x, y, x + w, y + h], face, edge, 5, 38)
    svg.rect(x, y, w, h, face, edge, 4, 38)
    draw_text(draw, (x + 36, y + 55), title, F["stage"], edge)
    svg.text(x + 36, y + 65, title, 30, edge, "bold")
    yy = y + 130
    for b in bullets:
        for k, line in enumerate(wrap(b, 27)):
            prefix = "• " if k == 0 else "  "
            draw_text(draw, (x + 40, yy), prefix + line, F["small"], COL["ink"])
            svg.text(x + 40, yy + 23, prefix + line, 24, COL["ink"])
            yy += 38
        yy += 10
    if footer:
        rounded(draw, [x + 36, y + h - 102, x + w - 36, y + h - 38], "#FFFFFF", edge, 3, 20)
        svg.rect(x + 36, y + h - 102, w - 72, 64, "#FFFFFF", edge, 3, 20)
        draw_text(draw, (x + w / 2, y + h - 70), footer, F["badge"], edge, anchor="mm")
        svg.text(x + w / 2, y + h - 62, footer, 28, edge, "bold", "middle")


def build_figure():
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    svg = SVG(W, H)

    draw_text(draw, (MARGIN, 92), "Figure 1A | Model evolution map: replacement, not refinement", F["title"], COL["ink"])
    draw_text(
        draw,
        (MARGIN, 150),
        "Explicit lineage from the deprecated PGCS/Phi scalar model to the locked latent-state-regime mixture framework.",
        F["subtitle"],
        COL["muted"],
    )
    svg.text(MARGIN, 115, "Figure 1A | Model evolution map: replacement, not refinement", 58, COL["ink"], "bold")
    svg.text(MARGIN, 175, "Explicit lineage from the deprecated PGCS/Phi scalar model to the locked latent-state-regime mixture framework.", 30, COL["muted"])

    labels = [
        (
            "Old Model",
            ["PGCS / Phi scalar", "single order-parameter assumption", "global score expected to rank states"],
            COL["old_bg"],
            COL["old_edge"],
            "SCALAR ASSUMED",
        ),
        (
            "Failure Layer",
            ["Phi AUC ≈ 0.480", "permutation instability", "KS detects shape mismatch, not separability"],
            COL["fail_bg"],
            COL["fail_edge"],
            "FAILURE LOCKED",
        ),
        (
            "Transition",
            ["scalar embedding breakdown", "cross-species inconsistency", "no valid global threshold"],
            COL["trans_bg"],
            COL["trans_edge"],
            "BREAKDOWN",
        ),
        (
            "New Model",
            ["latent state regime mixture P(Z | S, W_GRN)", "regime-conditioned dynamics", "no global scalar"],
            COL["new_bg"],
            COL["new_edge"],
            "REPLACEMENT",
        ),
        (
            "Final Claim",
            ["observed states are mixtures", "overlap does not imply identity", "model is replacement, not refinement"],
            COL["claim_bg"],
            COL["claim_edge"],
            "NOT REFINEMENT",
        ),
    ]

    card_w = 900
    gap = 110
    y = 430
    h = 920
    xs = [MARGIN + i * (card_w + gap) for i in range(5)]
    for i, (title, bullets, face, edge, footer) in enumerate(labels):
        card(draw, svg, xs[i], y, card_w, h, title, bullets, face, edge, footer)
        if i < 4:
            ay = y + h / 2
            arrow(draw, xs[i] + card_w + 18, ay, xs[i + 1] - 25, ay)
            svg.line(xs[i] + card_w + 18, ay, xs[i + 1] - 25, ay, "#334155", 6, True)

    # Bottom explanatory strip.
    strip_x, strip_y, strip_w, strip_h = MARGIN, 1605, W - 2 * MARGIN, 520
    rounded(draw, [strip_x, strip_y, strip_x + strip_w, strip_y + strip_h], COL["card"], COL["border"], 4, 34)
    svg.rect(strip_x, strip_y, strip_w, strip_h, COL["card"], COL["border"], 4, 34)
    draw_text(draw, (strip_x + 42, strip_y + 70), "Interpretive lock", F["stage"], COL["blue"])
    svg.text(strip_x + 42, strip_y + 82, "Interpretive lock", 30, COL["blue"], "bold")
    lock_lines = [
        "The scalar hypothesis is not partially retained: Phi/PGCS is used only as a deprecated diagnostic proxy.",
        "The accepted representation is P(Z | S, W_GRN), where adult_repair, embryonic_reactivation, salamander_blastema and salamander_intact are latent state regimes.",
        "Regime-conditioned dynamics replace the global scalar model; dry-lab perturbation results remain PARTIAL_CLOSURE only.",
    ]
    yy = strip_y + 135
    for line in lock_lines:
        for wrapped in wrap(line, 126):
            draw_text(draw, (strip_x + 52, yy), wrapped, F["body"], COL["ink"])
            svg.text(strip_x + 52, yy + 26, wrapped, 28, COL["ink"])
            yy += 46
        yy += 8

    # Explicit no-claim badge.
    badge = "NO GLOBAL SCALAR • NO WET-LAB CAUSAL CLAIM • REPLACEMENT MODEL"
    rounded(draw, [strip_x + strip_w - 1720, strip_y + strip_h - 116, strip_x + strip_w - 50, strip_y + strip_h - 42], "#FFFFFF", COL["blue"], 4, 24)
    draw_text(draw, (strip_x + strip_w - 885, strip_y + strip_h - 78), badge, F["badge"], COL["blue"], anchor="mm")
    svg.rect(strip_x + strip_w - 1720, strip_y + strip_h - 116, 1670, 74, "#FFFFFF", COL["blue"], 4, 24)
    svg.text(strip_x + strip_w - 885, strip_y + strip_h - 68, badge, 28, COL["blue"], "bold", "middle")

    foot = "Abbreviations: PGCS, prior scalar cell-state score; Phi, deprecated latent scalar proxy; GRN, gene regulatory network."
    draw_text(draw, (MARGIN, H - 115), foot, F["tiny"], COL["muted"])
    svg.text(MARGIN, H - 95, foot, 21, COL["muted"])

    img.save(FIG_PNG, dpi=(300, 300))
    img.save(FIG_PDF, "PDF", resolution=300.0)
    svg.save(FIG_SVG)


def section_text() -> str:
    return (
        "## Model lineage and replacement logic\n\n"
        "The final framework is a replacement of the earlier PGCS/Phi scalar model, not a refinement of it. "
        "The old model assumed that cell fate structure could be compressed into a single scalar order parameter. "
        "That assumption failed under the locked validation layer: Phi showed random-level discrimination (AUC approximately 0.480), unstable permutation and bootstrap behavior, and KS-detected distributional shape differences that did not provide a separable biological threshold. "
        "These failures define a scalar embedding breakdown, especially across mammalian repair and salamander blastema contexts. "
        "The manuscript therefore replaces the scalar model with the latent-state-regime mixture P(Z|S,W_GRN), in which observed states carry posterior mass over adult_repair, embryonic_reactivation, salamander_blastema and salamander_intact regimes. "
        "This lineage is summarized in Figure 1A and fixes the interpretive rule used throughout the manuscript: the accepted model is a regime-mixture replacement, not a partial rescue of Phi.\n"
    )


def caption_text() -> str:
    return (
        "Figure 1A | Model evolution map. The earlier PGCS/Phi scalar model assumed that cell fate could be represented by a single order parameter. "
        "The validation layer rejected this assumption through random-level Phi discrimination (AUC approximately 0.480), permutation instability and KS shape mismatch without separability. "
        "This scalar embedding breakdown motivates replacement by the latent-state-regime mixture P(Z|S,W_GRN), with regime-conditioned dynamics and no global scalar. "
        "The final claim is replacement, not refinement; Phi is retained only as a deprecated diagnostic proxy."
    )


def insert_paragraph_after(paragraph: Paragraph, text: str | None = None, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_paragraph_before(paragraph: Paragraph, text: str | None = None, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_caption_style(p: Paragraph):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.italic = True


def add_docx_updates():
    doc = Document(BASE_DOCX)
    for p in doc.paragraphs:
        if "scalar deprecated scalar proxy" in p.text:
            p.text = p.text.replace("scalar deprecated scalar proxy", "deprecated scalar proxy")
    # Insert a short prose section after the paragraph that introduces the replacement framework.
    replacement_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("The replacement is a latent-state-regime mixture framework."):
            replacement_idx = i
            break
    if replacement_idx is None:
        raise RuntimeError("Could not find replacement-framework paragraph.")
    anchor = doc.paragraphs[replacement_idx]
    h = insert_paragraph_after(anchor, "Model lineage and replacement logic", "Heading 2")
    prose = insert_paragraph_after(
        h,
        (
            "The final framework is a replacement of the earlier PGCS/Phi scalar model, not a refinement of it. "
            "The old model assumed that cell fate structure could be compressed into a single scalar order parameter. "
            "That assumption failed under the locked validation layer: Phi showed random-level discrimination (AUC approximately 0.480), unstable permutation and bootstrap behavior, and KS-detected distributional shape differences that did not provide a separable biological threshold. "
            "These failures define a scalar embedding breakdown, especially across mammalian repair and salamander blastema contexts. "
            "The manuscript therefore replaces the scalar model with the latent-state-regime mixture P(Z|S,W_GRN), in which observed states carry posterior mass over adult_repair, embryonic_reactivation, salamander_blastema and salamander_intact regimes. "
            "This lineage is summarized in Figure 1A and fixes the interpretive rule used throughout the manuscript: the accepted model is a regime-mixture replacement, not a partial rescue of Phi."
        ),
    )
    prose.paragraph_format.space_after = Pt(8)

    # Insert Figure 1A immediately before Result 1 to make lineage visible at the start of Results.
    result_anchor = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("Result 1."):
            result_anchor = p
            break
    if result_anchor is None:
        raise RuntimeError("Could not find Result 1 heading.")
    fig_p = insert_paragraph_after(result_anchor)
    fig_p.paragraph_format.page_break_before = False
    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_p.add_run().add_picture(str(FIG_PNG), width=Inches(6.35))
    cap = insert_paragraph_after(fig_p, caption_text())
    set_caption_style(cap)
    doc.save(UPDATED_DOCX)

    legends = Document(BASE_LEGENDS)
    for p in legends.paragraphs:
        if "scalar deprecated scalar proxy" in p.text:
            p.text = p.text.replace("scalar deprecated scalar proxy", "deprecated scalar proxy")
    # Add legend after subtitle paragraph, before Figure 1 legend.
    first_fig_idx = next((i for i, p in enumerate(legends.paragraphs) if p.text.strip() == "Figure 1"), None)
    if first_fig_idx is None:
        raise RuntimeError("Could not find Figure 1 legend heading.")
    before = legends.paragraphs[first_fig_idx]
    h1a = insert_paragraph_before(before, "Figure 1A")
    h1a.style = legends.styles["Heading 2"]
    leg = insert_paragraph_after(h1a, caption_text())
    leg.paragraph_format.space_after = Pt(6)
    legends.save(UPDATED_LEGENDS)


def update_markdown():
    SECTION_MD.write_text(section_text(), encoding="utf-8")
    CAPTION_MD.write_text("# " + caption_text() + "\n", encoding="utf-8")
    text = BASE_MD.read_text(encoding="utf-8")
    insert_after = "Under this interpretation, latent state regime mixture explains observed structure while preserving overlap, non-identifiability and regime-specific divergence.\n"
    if insert_after in text:
        text = text.replace(insert_after, insert_after + "\n" + section_text() + "\n")
    else:
        text += "\n\n" + section_text()
    UPDATED_MD.write_text(text, encoding="utf-8")


def update_package():
    if not PKG.exists():
        return
    shutil.copy2(UPDATED_DOCX, PKG / UPDATED_DOCX.name)
    shutil.copy2(UPDATED_LEGENDS, PKG / UPDATED_LEGENDS.name)
    shutil.copy2(UPDATED_MD, PKG / UPDATED_MD.name)
    main_fig_dir = PKG / "main_figures"
    main_fig_dir.mkdir(exist_ok=True)
    for src in [FIG_PNG, FIG_PDF, FIG_SVG, SECTION_MD, CAPTION_MD]:
        shutil.copy2(src, main_fig_dir / src.name if src.suffix.lower() in {".png", ".pdf", ".svg"} else PKG / src.name)


def main():
    build_figure()
    update_markdown()
    add_docx_updates()
    update_package()


if __name__ == "__main__":
    main()
