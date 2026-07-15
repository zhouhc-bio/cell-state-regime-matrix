from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw


INPUT = Path(
    "/Users/hanchengdezhuanqiangongju/Library/Group Containers/"
    "group.com.apple.notes/Accounts/2E5C596C-7223-45E9-B7AC-9CE08155F4D4/"
    "Media/F32EEE6A-28A5-4608-9D5F-8C4203FEF13C/"
    "1_D2BDED3F-C668-40A5-A821-537EF98A1B03/细胞命运论文010.docx"
)
OUTPUT = Path("细胞命运论文010_FIGURE_IMAGE_TITLES_REMOVED.docx")
WORK = Path("work/figure_title_cleanup_010")
ORIGINAL_MEDIA = WORK / "original_media"
CLEAN_MEDIA = WORK / "cleaned_media"
REPORT = Path("figure_image_title_cleanup_check.md")


# Rectangles are in pixel coordinates in each embedded source image.
# The masks remove only figure-level titles, legacy Figure-number headers,
# internal review/status notes, and slide-like conclusion strips.
# Panel labels, axes, legends, colorbars, data labels, regime labels,
# statistical annotations, model equations, and scientific panel titles remain.
MASKS: dict[str, list[tuple[int, int, int, int, str, tuple[int, int, int] | None]]] = {
    "image1.png": [
        (0, 0, 5400, 210, "embedded Figure 1A title/subtitle", (255, 255, 255)),
        (120, 1560, 5280, 2165, "presentation-style interpretive lock/conclusion box", (255, 255, 255)),
    ],
    "image2.png": [
        (0, 0, 5000, 300, "embedded Figure 1 title/subtitle", (246, 248, 252)),
        (540, 2295, 3180, 2385, "presentation-style final visual rule sentence", (255, 255, 255)),
    ],
    "image3.png": [
        (0, 0, 4200, 300, "embedded Figure 2 title/subtitle", (246, 248, 252)),
        (60, 425, 4140, 665, "large figure-level title inside image", (255, 255, 255)),
        (50, 4240, 4145, 4565, "presentation-style conclusion boxes", (255, 255, 255)),
    ],
    "image4.png": [
        (0, 0, 4425, 300, "embedded Figure 3 title/subtitle", (246, 248, 252)),
        (80, 410, 4350, 645, "large figure-level title inside image", (255, 255, 255)),
        (55, 4070, 4370, 4365, "presentation-style conclusion boxes", (255, 255, 255)),
    ],
    "image5.png": [
        (0, 0, 6400, 300, "embedded Figure 4 title/subtitle", (246, 248, 252)),
        (70, 3500, 6325, 4535, "presentation-style conclusion boxes", (255, 255, 255)),
    ],
    "image6.png": [
        (0, 0, 6400, 300, "embedded Figure 5 title/subtitle", (246, 248, 252)),
        (3820, 2750, 6380, 3860, "internal interpretation-lock box", (255, 255, 255)),
        (70, 3650, 6325, 4535, "presentation-style allowed/not-allowed boxes", (255, 255, 255)),
    ],
    "image7.png": [
        (0, 0, 4200, 300, "embedded Figure 6 title/subtitle", (246, 248, 252)),
        (60, 385, 4140, 615, "large figure-level title inside image", (255, 255, 255)),
        (55, 4985, 4145, 5285, "presentation-style conclusion boxes", (255, 255, 255)),
    ],
    "image8.png": [
        (0, 0, 4200, 280, "embedded Figure 7 title/subtitle", (246, 248, 252)),
        (250, 340, 3970, 470, "legacy internal Figure 1 title", (246, 248, 252)),
    ],
    "image9.png": [
        (0, 0, 7600, 280, "embedded Figure 8 title/subtitle", (246, 248, 252)),
        (250, 395, 3580, 560, "legacy internal Figure 2 title", (246, 248, 252)),
        (3950, 395, 7285, 560, "legacy internal Figure 3 title", (246, 248, 252)),
    ],
    "image10.png": [
        (0, 0, 4200, 280, "embedded Figure 9 title/subtitle", (246, 248, 252)),
        (250, 345, 2860, 500, "legacy internal Figure 4 title", (246, 248, 252)),
        (55, 3315, 4145, 3615, "presentation-style conclusion boxes", (255, 255, 255)),
    ],
    "image11.png": [
        (0, 0, 4200, 250, "embedded Supplementary Figure 1 title", (255, 255, 255)),
    ],
}


def count_refs(docx_path: Path) -> int:
    with zipfile.ZipFile(docx_path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    return len(re.findall(r"\[REF::[^\]]+\]", xml))


def media_count(docx_path: Path) -> int:
    with zipfile.ZipFile(docx_path) as zf:
        return sum(1 for name in zf.namelist() if name.startswith("word/media/"))


def extract_media() -> None:
    ORIGINAL_MEDIA.mkdir(parents=True, exist_ok=True)
    CLEAN_MEDIA.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(INPUT) as zf:
        for name in zf.namelist():
            if name.startswith("word/media/"):
                target = ORIGINAL_MEDIA / Path(name).name
                target.write_bytes(zf.read(name))


def clean_media() -> list[str]:
    log: list[str] = []
    for src in sorted(ORIGINAL_MEDIA.glob("image*.png"), key=lambda p: int(p.stem.replace("image", ""))):
        im = Image.open(src).convert("RGB")
        draw = ImageDraw.Draw(im)
        masks = MASKS.get(src.name, [])
        for x0, y0, x1, y1, reason, fill in masks:
            draw.rectangle((x0, y0, x1, y1), fill=fill or (255, 255, 255))
            log.append(f"- `{src.name}`: masked `{reason}` at ({x0}, {y0}, {x1}, {y1})")
        im.save(CLEAN_MEDIA / src.name, optimize=True)
    return log


def rebuild_docx() -> None:
    temp = WORK / "tmp_rebuild.docx"
    if temp.exists():
        temp.unlink()
    with zipfile.ZipFile(INPUT, "r") as zin, zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/media/"):
                clean = CLEAN_MEDIA / Path(item.filename).name
                if clean.exists():
                    data = clean.read_bytes()
            zout.writestr(item, data)
    shutil.move(str(temp), str(OUTPUT))


def make_report(log: list[str]) -> None:
    before_refs = count_refs(INPUT)
    after_refs = count_refs(OUTPUT)
    before_media = media_count(INPUT)
    after_media = media_count(OUTPUT)
    before_paragraphs = len(Document(INPUT).paragraphs)
    after_paragraphs = len(Document(OUTPUT).paragraphs)
    status = "PASS" if (before_refs == after_refs and before_media == after_media and before_paragraphs == after_paragraphs) else "CHECK"
    REPORT.write_text(
        "\n".join(
            [
                "# Figure Image Title Cleanup Check",
                "",
                f"Input: `{INPUT}`",
                f"Output: `{OUTPUT.resolve()}`",
                "",
                f"Structural status: `{status}`",
                f"Paragraph count: {before_paragraphs} -> {after_paragraphs}",
                f"Embedded media count: {before_media} -> {after_media}",
                f"[REF::...] citation slot count: {before_refs} -> {after_refs}",
                "",
                "Removed image-internal artifacts:",
                *log,
                "",
                "Preserved by design: panel labels, panel titles, axes, legends, colorbars, data labels, regime labels, statistical annotations, model equations, and manuscript figure legends.",
            ]
        ),
        encoding="utf-8",
    )


def main() -> None:
    if not INPUT.exists():
        raise FileNotFoundError(INPUT)
    extract_media()
    log = clean_media()
    rebuild_docx()
    make_report(log)
    print(OUTPUT.resolve())
    print(REPORT.resolve())


if __name__ == "__main__":
    main()
