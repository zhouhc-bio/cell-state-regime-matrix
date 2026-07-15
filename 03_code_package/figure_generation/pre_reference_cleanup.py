#!/usr/bin/env python3
"""Remove review-only artifacts from the clean review draft."""

from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path.cwd()
INPUT = ROOT / "Regenerative_cell_fate_latent_regime_CLEAN_REVIEW_DRAFT.docx"
OUTPUT = ROOT / "Regenerative_cell_fate_latent_regime_PRE_REFERENCE_CLEAN.docx"
CHECK = ROOT / "cleanup_check.md"


ADMIN_TEXT = {
    "Data availability": (
        "All public datasets and derived source-data tables used in this study "
        "will be listed before submission."
    ),
    "Code availability": (
        "Analysis code and figure-generation scripts will be made available before submission."
    ),
    "Acknowledgements": "Acknowledgements will be finalized before submission.",
    "Author contributions": "Author contributions will be finalized before submission.",
    "Competing interests": "The authors declare no competing interests.",
}

REVIEW_NOTE_RE = re.compile(
    r"(Fig\.|Figure|Supplementary Fig\.|Supplementary Figure)\s+[^.]*"
    r"inserted for review; final journal submission may require separate figure upload\.",
    re.IGNORECASE,
)

BAD_TEXT_PATTERNS = [
    "inserted for review",
    "final journal submission may require separate figure upload",
    "TODO",
    "FIXME",
    "REVIEW_ONLY",
    "INTERNAL_NOTE",
    "MISSING_SOURCE",
    "REBUILD",
    "REANNOTATE",
]


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing | .//w:pict"))


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def set_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def ref_count(docx_path: Path) -> tuple[int, int]:
    with zipfile.ZipFile(docx_path) as z:
        xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
    refs = re.findall(r"\[REF::[^\]]+\]", xml)
    return len(refs), len(set(refs))


def media_count(docx_path: Path) -> int:
    with zipfile.ZipFile(docx_path) as z:
        return sum(1 for n in z.namelist() if n.startswith("word/media/"))


def clean_doc() -> tuple[int, int, int, int]:
    before_refs, before_unique = ref_count(INPUT)
    before_media = media_count(INPUT)
    doc = Document(INPUT)

    removed_review_notes = 0
    for paragraph in list(doc.paragraphs):
        text = " ".join(paragraph.text.split())
        if paragraph_has_drawing(paragraph):
            continue
        if REVIEW_NOTE_RE.fullmatch(text) or (
            "inserted for review" in text
            and "final journal submission may require separate figure upload" in text
        ):
            delete_paragraph(paragraph)
            removed_review_notes += 1

    # Reload paragraph list after deletions and standardize administrative sections.
    admin_replacements = 0
    paragraphs = doc.paragraphs
    for idx, paragraph in enumerate(paragraphs):
        heading = paragraph.text.strip()
        if heading in ADMIN_TEXT:
            for next_paragraph in paragraphs[idx + 1 :]:
                if next_paragraph.style.name.startswith("Heading"):
                    break
                if paragraph_has_drawing(next_paragraph):
                    continue
                if next_paragraph.text.strip():
                    if next_paragraph.text.strip() != ADMIN_TEXT[heading]:
                        set_paragraph_text(next_paragraph, ADMIN_TEXT[heading])
                        admin_replacements += 1
                    break

    doc.save(OUTPUT)
    after_refs, after_unique = ref_count(OUTPUT)
    after_media = media_count(OUTPUT)
    if (before_refs, before_unique) != (after_refs, after_unique):
        raise RuntimeError(
            f"REF slot count changed: before {(before_refs, before_unique)}, "
            f"after {(after_refs, after_unique)}"
        )
    if before_media != after_media:
        raise RuntimeError(f"Media count changed: before {before_media}, after {after_media}")
    return removed_review_notes, admin_replacements, after_refs, after_media


def main() -> None:
    removed_review_notes, admin_replacements, refs, media = clean_doc()
    # Initial check note; render/visual checks append to this file later.
    CHECK.write_text(
        "\n".join(
            [
                "# Cleanup Check",
                "",
                f"- Review-only figure insertion lines removed: {removed_review_notes}.",
                f"- Administrative placeholder/section text replacements: {admin_replacements}.",
                f"- Embedded media count after cleanup: {media}.",
                f"- Citation slots preserved: {refs} `[REF::...]` slots.",
                "- Figure render/visual-label check: pending render QA.",
                "- Ready for manual reference insertion: pending final verification.",
                "",
            ]
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
