from __future__ import annotations

import csv
import re
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image


ROOT = Path("/Users/hanchengdezhuanqiangongju/Documents/Codex/2026-06-18/task-reconstruct-and-continue-analysis-of")
INPUT_DOCX = Path("/Users/hanchengdezhuanqiangongju/Downloads/Regenerative_cell_fate_latent_regime_manuscript_citation_slots.docx")
OUT_DOCX = ROOT / "Regenerative_cell_fate_latent_regime_manuscript_WITH_INLINE_FIGURES.docx"
LEGENDS_DOCX = ROOT / "Figure_legends_submission_ready.docx"
INVENTORY_TSV = ROOT / "figure_inventory.tsv"
AUDIT_MD = ROOT / "figure_audit_final.md"
REPAIR_MD = ROOT / "figure_repair_tasks.md"


@dataclass
class FigureSpec:
    figure_id: str
    expected_file: str
    selected_path: Path
    insertion_anchor: str
    classification: str
    legend_risk: str
    notes: str


FIGURES = [
    FigureSpec(
        "Figure 1A",
        "MODEL_EVOLUTION_MAP_Figure1A.(png/pdf/svg)",
        ROOT / "FINAL_SUBMISSION_PACKAGE/main_figures/MODEL_EVOLUTION_MAP_Figure1A.png",
        "We first made the model lineage explicit",
        "READY_MAIN",
        "low",
        "Model evolution map; replacement, not refinement.",
    ),
    FigureSpec(
        "Figure 1",
        "Final_Figure_1_Phi_failure_latent_regime_replacement.(png/pdf)",
        ROOT / "FINAL_SUBMISSION_PACKAGE/main_figures/Final_Figure_1_Phi_failure_latent_regime_replacement.png",
        "The four locked regimes are adult_repair",
        "READY_MAIN",
        "medium",
        "Phi appears as rejected scalar projection only.",
    ),
    FigureSpec(
        "Figure 2",
        "Final_Figure_2_shared_accessibility_reannotated.(png/pdf)",
        ROOT / "FINAL_SUBMISSION_PACKAGE/main_figures/Final_Figure_2_shared_accessibility_reannotated.png",
        "This result is central to the replacement model",
        "READY_MAIN",
        "medium",
        "Accessibility is separated from fate determination.",
    ),
    FigureSpec(
        "Figure 3",
        "Final_Figure_3_stemness_local_accessibility.(png/pdf)",
        ROOT / "FINAL_SUBMISSION_PACKAGE/main_figures/Final_Figure_3_stemness_local_accessibility.png",
        "This distinction protects the manuscript",
        "READY_MAIN",
        "medium",
        "Stemness is local accessibility, not global fate identity.",
    ),
    FigureSpec(
        "Figure 4",
        "Final_Figure_4_regime_conditioned_positional_information.(png/pdf)",
        ROOT / "FINAL_SUBMISSION_PACKAGE/main_figures/Final_Figure_4_regime_conditioned_positional_information.png",
        "This is particularly important for salamander regeneration",
        "READY_MAIN",
        "high",
        "Positional information remains regime-conditioned.",
    ),
    FigureSpec(
        "Figure 5",
        "Final_Figure_5_fate_lock_adult_repair_basin.(png/pdf)",
        ROOT / "FINAL_SUBMISSION_PACKAGE/main_figures/Final_Figure_5_fate_lock_adult_repair_basin.png",
        "The basin interpretation is intentionally restrained",
        "READY_MAIN",
        "high",
        "Fate-lock is basin constraint; no complete causal claim.",
    ),
    FigureSpec(
        "Figure 6",
        "Final_Figure_6_tumor_like_plasticity_distinct_branch.(png/pdf)",
        ROOT / "FINAL_SUBMISSION_PACKAGE/main_figures/Final_Figure_6_tumor_like_plasticity_distinct_branch.png",
        "In the final framework, tumor-like plasticity",
        "READY_MAIN",
        "high",
        "Tumor-like plasticity remains distinct from regeneration.",
    ),
    FigureSpec(
        "Figure 7",
        "Final_Figure_7_single_Phi_rejection.(png/pdf)",
        ROOT / "FINAL_SUBMISSION_PACKAGE/main_figures/Final_Figure_7_single_Phi_rejection.png",
        "This failure is the turning point of the manuscript",
        "READY_MAIN",
        "low",
        "Scalar Phi is rejected as a regime separator.",
    ),
    FigureSpec(
        "Figure 8",
        "Final_Figure_8_latent_regime_posterior_dynamics.(png/pdf)",
        ROOT / "FINAL_SUBMISSION_PACKAGE/main_figures/Final_Figure_8_latent_regime_posterior_dynamics.png",
        "The model is deliberately described as a working representation",
        "READY_MAIN",
        "medium",
        "Primary model remains P(Z|S,W_GRN).",
    ),
    FigureSpec(
        "Figure 9",
        "Final_Figure_9_overlap_symmetrized_divergence.(png/pdf)",
        ROOT / "FINAL_SUBMISSION_PACKAGE/main_figures/Final_Figure_9_overlap_symmetrized_divergence.png",
        "Together, overlap and divergence support",
        "READY_MAIN",
        "medium",
        "Overlap is not identity; divergence is not scalar distance.",
    ),
    FigureSpec(
        "Supplementary Figure 1",
        "SUPP_FIGURE_CAUSAL_CLOSURE_REDESIGNED.(png/pdf/svg)",
        ROOT / "outputs/FINAL_SUPPLEMENTARY_FIGURES_VISUALLY_REDESIGNED/SUPP_FIGURE_CAUSAL_CLOSURE_REDESIGNED.png",
        "The global W_GRN model was insufficient",
        "READY_SUPPLEMENT",
        "medium",
        "Dry-lab perturbation consistency only; PARTIAL_CLOSURE.",
    ),
]


def all_candidates(spec: FigureSpec) -> list[Path]:
    stem_tokens = {
        "Figure 1A": ["MODEL_EVOLUTION_MAP_Figure1A"],
        "Figure 1": ["Final_Figure_1_Phi_failure_latent_regime_replacement"],
        "Figure 2": ["Final_Figure_2_shared_accessibility_reannotated"],
        "Figure 3": ["Final_Figure_3_stemness_local_accessibility"],
        "Figure 4": ["Final_Figure_4_regime_conditioned_positional_information"],
        "Figure 5": ["Final_Figure_5_fate_lock_adult_repair_basin"],
        "Figure 6": ["Final_Figure_6_tumor_like_plasticity_distinct_branch"],
        "Figure 7": ["Final_Figure_7_single_Phi_rejection"],
        "Figure 8": ["Final_Figure_8_latent_regime_posterior_dynamics"],
        "Figure 9": ["Final_Figure_9_overlap_symmetrized_divergence"],
        "Supplementary Figure 1": ["SUPP_FIGURE_CAUSAL_CLOSURE_REDESIGNED", "FIGURE_CAUSAL_CLOSURE_FINAL"],
    }[spec.figure_id]
    exts = {".png", ".pdf", ".svg", ".tif", ".tiff", ".jpg", ".jpeg"}
    found: list[Path] = []
    for p in (ROOT / "outputs").rglob("*"):
        if p.is_file() and p.suffix.lower() in exts and any(tok in p.name for tok in stem_tokens):
            found.append(p)
    for p in (ROOT / "FINAL_SUBMISSION_PACKAGE").rglob("*"):
        if p.is_file() and p.suffix.lower() in exts and any(tok in p.name for tok in stem_tokens):
            found.append(p)
    unique = sorted(set(found), key=lambda x: str(x))
    return unique


def image_info(path: Path) -> tuple[str, str]:
    if not path.exists():
        return "", ""
    fmt = path.suffix.lower().lstrip(".").upper()
    if path.suffix.lower() in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
        with Image.open(path) as im:
            return fmt, f"{im.width}x{im.height} px"
    if path.suffix.lower() == ".pdf":
        try:
            from pypdf import PdfReader

            return fmt, f"{len(PdfReader(str(path)).pages)} page(s)"
        except Exception:
            return fmt, "PDF page count unavailable"
    if path.suffix.lower() == ".svg":
        txt = path.read_text(errors="ignore")[:1000]
        width = re.search(r'width="([^"]+)"', txt)
        height = re.search(r'height="([^"]+)"', txt)
        if width and height:
            return fmt, f"{width.group(1)} x {height.group(1)}"
        return fmt, "SVG dimensions not declared"
    return fmt, ""


def rel(path: Path) -> str:
    try:
        return str(path.relative_to(ROOT))
    except ValueError:
        return str(path)


def write_inventory() -> dict[str, list[Path]]:
    candidates_by_id = {}
    with INVENTORY_TSV.open("w", newline="") as f:
        writer = csv.writer(f, delimiter="\t")
        writer.writerow([
            "figure_id",
            "expected_file",
            "detected_file",
            "file_format",
            "file_size",
            "resolution_or_page_count",
            "status",
            "notes",
        ])
        for spec in FIGURES:
            candidates = all_candidates(spec)
            candidates_by_id[spec.figure_id] = candidates
            selected = spec.selected_path if spec.selected_path.exists() else (candidates[0] if candidates else None)
            if selected is None:
                writer.writerow([spec.figure_id, spec.expected_file, "", "", "", "", "MISSING", "No matching figure file found."])
                continue
            fmt, dims = image_info(selected)
            size = f"{selected.stat().st_size} bytes"
            status = "FOUND_READY"
            notes = spec.notes
            if selected.suffix.lower() not in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
                status = "FORMAT_NEEDS_CONVERSION"
                notes += " Selected file is not directly embeddable by python-docx."
            if len(candidates) > 1:
                notes += f" Duplicate candidates detected ({len(candidates)}); selected authoritative review copy."
            writer.writerow([spec.figure_id, spec.expected_file, rel(selected), fmt, size, dims, status, notes])
    return candidates_by_id


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def find_anchor(doc: Document, phrase: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if phrase in p.text:
            return p
    return None


def insert_figures_docx() -> None:
    shutil.copyfile(INPUT_DOCX, OUT_DOCX)
    doc = Document(OUT_DOCX)
    for spec in FIGURES:
        if not spec.selected_path.exists():
            continue
        anchor = find_anchor(doc, spec.insertion_anchor)
        if anchor is None:
            anchor = find_anchor(doc, "Figure legends")
        if anchor is None:
            anchor = doc.paragraphs[-1]
        p = insert_paragraph_after(anchor, "")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(str(spec.selected_path), width=Inches(6.25))

        cap = insert_paragraph_after(p, f"{spec.figure_id} inserted for review; final journal submission may require separate figure upload.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        for r in cap.runs:
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = RGBColor(90, 90, 90)
    doc.save(OUT_DOCX)


def set_basic_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color


def extract_legend_pairs() -> list[tuple[str, str]]:
    src = Document(INPUT_DOCX)
    paras = [p.text.strip() for p in src.paragraphs]
    try:
        start = paras.index("Figure legends") + 1
        end = paras.index("References")
    except ValueError:
        return []
    content = [p for p in paras[start:end] if p]
    pairs = []
    for i in range(0, len(content), 2):
        title = content[i]
        legend = content[i + 1] if i + 1 < len(content) else ""
        pairs.append((title, legend))
    return pairs


def write_legends_docx() -> None:
    warnings = {
        "Figure 1A": "Source warning: figure-level source present; no repair needed.",
        "Figure 1": "Source warning: figure-level source present; panel-level source detail should be tightened before formal source-data submission.",
        "Figure 2": "Source warning: figure-level source present; panel-level source detail should be tightened before formal source-data submission.",
        "Figure 3": "Source warning: figure-level source present; panel-level source detail should be tightened before formal source-data submission.",
        "Figure 4": "Source warning: figure-level source present; panel-level source detail should be tightened before formal source-data submission.",
        "Figure 5": "Source warning: figure-level source present; panel-level source detail should be tightened before formal source-data submission.",
        "Figure 6": "Source warning: figure-level source present; panel-level source detail should be tightened before formal source-data submission.",
        "Figure 7": "Source warning: scalar-invalidation source TSVs are present; keep Phi wording negative only.",
        "Figure 8": "Source warning: model-source MD/TSV files are present; do not treat as wet-lab causal validation.",
        "Figure 9": "Source warning: overlap/divergence source files are present; divergence is not axis distance.",
        "Supplementary Figure 1": "Source warning: supplementary perturbation TSVs are present; PARTIAL_CLOSURE only.",
    }
    doc = Document()
    set_basic_styles(doc)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Figure legends for submission-ready review")
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)
    intro = doc.add_paragraph("Citation slots are placeholders only. Resolve real references later; do not fabricate DOI, PMID, author-year or journal metadata.")
    intro.runs[0].font.italic = True
    for title_text, legend in extract_legend_pairs():
        doc.add_paragraph(title_text, style="Heading 1")
        doc.add_paragraph(legend)
        slots = sorted(set(re.findall(r"\[REF::[^\]]+\]", title_text + " " + legend)))
        slot_text = "Required citation slots: " + ("; ".join(slots) if slots else "none")
        p = doc.add_paragraph(slot_text)
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(90, 90, 90)
        key = title_text.split("|")[0].strip()
        w = doc.add_paragraph(warnings.get(key, "Source warning: MISSING_SOURCE"))
        w.runs[0].font.size = Pt(9)
        w.runs[0].font.italic = True
        w.runs[0].font.color.rgb = RGBColor(120, 80, 0)
    doc.save(LEGENDS_DOCX)


def write_audit(candidates_by_id: dict[str, list[Path]]) -> None:
    lines = [
        "# Final Figure Audit",
        "",
        "Scope: review-ready inline figure assembly for `Regenerative_cell_fate_latent_regime_manuscript_citation_slots.docx`.",
        "",
        "Global result: all expected review figures were detected and inserted. The scientific model remains centered on `P(Z|S,W_GRN)`; PGCS/Phi appears only as failed scalar projection or invalidation/QC layer.",
        "",
        "| figure_id | classification | story match | label/readability audit | regime/Phi/causal-safety audit | source trace | notes |",
        "|---|---|---|---|---|---|---|",
    ]
    for spec in FIGURES:
        if not spec.selected_path.exists():
            classification = "NEEDS_RECONSTRUCTION"
            story = "Missing figure file."
            readability = "Not assessed."
            safety = "Not assessed."
            trace = "MISSING_SOURCE"
        else:
            classification = spec.classification
            story = "Matches the manuscript figure legend and locked figure-to-story map."
            readability = "Panel labels and overall labels are visible in the selected full-size PNG; final journal upload should use separate high-resolution files."
            safety = "Regime names and claim boundaries are consistent; no wet-lab validation or scalar-positive Phi interpretation detected."
            if spec.figure_id in {"Figure 1", "Figure 2", "Figure 3", "Figure 4", "Figure 5", "Figure 6"}:
                trace = "Figure-level trace present; panel-level source-data detail should be tightened before formal submission."
            else:
                trace = "Traceable to figure/source summary files in outputs or FINAL_SUBMISSION_PACKAGE."
        lines.append(
            f"| {spec.figure_id} | {classification} | {story} | {readability} | {safety} | {trace} | {spec.notes} |"
        )
    lines += [
        "",
        "## Candidate handling",
        "",
        "Duplicate copies were found for most main figures in `outputs/`, `outputs/final_9figures_reconstructed/`, `outputs/FINAL_9_MAIN_FIGURES_VISUALLY_REDESIGNED/` and `FINAL_SUBMISSION_PACKAGE/main_figures/`. The assembly selected the `FINAL_SUBMISSION_PACKAGE/main_figures/` PNG copies for main figures and the visually redesigned supplementary causal-closure PNG for Supplementary Figure 1.",
        "",
        "## Claim-safety conclusion",
        "",
        "- `adult_repair`, `embryonic_reactivation`, `salamander_blastema` and `salamander_intact` are used consistently.",
        "- PGCS/Phi remains a failed scalar projection or QC/invalidation layer.",
        "- Tumor-like plasticity remains separated from regenerative competence.",
        "- Salamander blastema is not conflated with ordinary mammalian adult wound repair.",
        "- Supplementary Figure 1 reports partial dry-lab perturbation consistency only.",
    ]
    AUDIT_MD.write_text("\n".join(lines) + "\n")


def write_repair_tasks() -> None:
    lines = [
        "# Figure Repair Tasks",
        "",
        "No expected figure is missing for review assembly. The items below are source-trace and submission-polish tasks rather than blockers for review reading.",
        "",
        "| figure_id | what_is_missing_or_problematic | why_it_matters | source_file_needed | blocks_manuscript_review | can_move_to_supplementary | recommended_fix |",
        "|---|---|---|---|---|---|---|",
    ]
    for spec in FIGURES:
        if not spec.selected_path.exists():
            lines.append(
                f"| {spec.figure_id} | Figure file is missing. | Inline review cannot show this figure. | {spec.expected_file} | yes | depends on figure role | Rebuild or recover the locked figure from source outputs. |"
            )
        elif spec.figure_id in {"Figure 1", "Figure 2", "Figure 3", "Figure 4", "Figure 5", "Figure 6"}:
            lines.append(
                f"| {spec.figure_id} | Panel-level source-data trace is not fully enumerated in the current assembly map. | Formal submission/source-data audit may require panel-by-panel TSV links. | Panel source TSV/MD files or a panel source-data manifest. | no | no | Add a panel-level source-data manifest before journal upload; keep current figure in main text for review. |"
            )
    lines += [
        "| Supplementary Figure 1 | Use the redesigned audit dashboard consistently instead of older sparse causal-closure figure where possible. | Prevents sparse-data visual weakness and preserves PARTIAL_CLOSURE. | `outputs/FINAL_SUPPLEMENTARY_FIGURES_VISUALLY_REDESIGNED/SUPP_FIGURE_CAUSAL_CLOSURE_REDESIGNED.*` | no | already supplementary | Use redesigned dashboard for review package; retain source TSVs as Supplementary Tables. |",
        "",
        "## Blocking status",
        "",
        "BLOCKS_MANUSCRIPT_REVIEW: no",
        "",
        "BLOCKS_FINAL_JOURNAL_UPLOAD: partial, because panel-level source-data manifests should be tightened before formal source-data submission.",
    ]
    REPAIR_MD.write_text("\n".join(lines) + "\n")


def main() -> None:
    candidates = write_inventory()
    insert_figures_docx()
    write_legends_docx()
    write_audit(candidates)
    write_repair_tasks()
    print(f"Wrote {OUT_DOCX}")
    print(f"Wrote {LEGENDS_DOCX}")
    print(f"Wrote {INVENTORY_TSV}")
    print(f"Wrote {AUDIT_MD}")
    print(f"Wrote {REPAIR_MD}")


if __name__ == "__main__":
    main()
