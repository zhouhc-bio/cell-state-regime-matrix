from __future__ import annotations

import argparse
import json
import shutil
from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


TITLE = "状态体制条件化潜在结构而非标量可塑性解释单细胞数据集中的再生性细胞命运组织"
SHORT_TITLE = "状态体制条件化潜在结构与再生性细胞命运组织"
KEYWORDS = "关键词：细胞命运；再生；单细胞分析；潜在状态域体制；可塑性；位置信息；基因调控网络"

TOC_HEADINGS = [
    "摘要",
    "引言",
    "结果",
    "讨论",
    "局限性",
    "方法",
    "数据可用性",
    "代码可用性",
    "致谢",
    "作者贡献",
    "利益冲突",
    "图注",
    "参考文献",
]


def toc_line(title: str, page: str | int) -> str:
    # Sichuan University guide asks for continuous dot leaders between title
    # and page number. ASCII dots render more reliably than repeated full-width
    # ellipsis in LibreOffice/Word font substitution.
    display = "程序可用性" if title == "代码可用性" else title
    return f"{display}{'.' * max(12, 34 - len(display) * 2)}{page}"


def add_toc_runs(paragraph, title: str, page: str | int):
    display = "程序可用性" if title == "代码可用性" else title
    dots = "." * max(12, 34 - len(display) * 2)
    for text, font in ((display, "宋体"), (dots, "Times New Roman"), (str(page), "宋体")):
        run = paragraph.add_run(text)
        set_east_asia_font(run, font if font != "Times New Roman" else "宋体", "Times New Roman", 12)
    return paragraph


def audit_docx(path: Path) -> dict:
    markers = ["fldChar", "instrText", "ADDIN EN.CITE", "EN.CITE.DATA", "EndNote"]
    with ZipFile(path) as z:
        text = "".join(
            z.read(n).decode("utf-8", "ignore")
            for n in z.namelist()
            if n.startswith("word/") and n.endswith(".xml")
        )
        media = [n for n in z.namelist() if n.startswith("word/media/")]
    return {m: text.count(m) for m in markers} | {"media_count": len(media)}


def set_east_asia_font(run, east_asia="宋体", latin="Times New Roman", size_pt=None, bold=None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def paragraph_has_image(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//*[local-name()='blip']"))


def apply_paragraph_base(paragraph, size=12, east_asia="宋体", latin="Times New Roman", first_line=True):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_line:
        pf.first_line_indent = Pt(24)
    else:
        pf.first_line_indent = Pt(0)
    for run in paragraph.runs:
        set_east_asia_font(run, east_asia=east_asia, latin=latin, size_pt=size)


def clear_para(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def set_para_text(paragraph, text, style=None, size=12, east_asia="宋体", latin="Times New Roman", bold=None):
    clear_para(paragraph)
    if style:
        paragraph.style = style
    run = paragraph.add_run(text)
    set_east_asia_font(run, east_asia=east_asia, latin=latin, size_pt=size, bold=bold)
    return paragraph


def insert_after(paragraph, text="", style=None):
    new_p = deepcopy(paragraph._p)
    for child in list(new_p):
        new_p.remove(child)
    paragraph._p.addnext(new_p)
    out = Paragraph(new_p, paragraph._parent)
    if style:
        out.style = style
    if text:
        out.add_run(text)
    return out


def add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    set_east_asia_font(run, "宋体", "Times New Roman", 10.5)


def apply_heading(paragraph, level):
    if level == 1:
        paragraph.style = "Heading 1"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = paragraph.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(20)
        pf.space_before = Pt(20)
        pf.space_after = Pt(20)
        pf.first_line_indent = Pt(0)
        for run in paragraph.runs:
            set_east_asia_font(run, "黑体", "Times New Roman", 15, True)
    elif level == 2:
        paragraph.style = "Heading 2"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = paragraph.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(20)
        pf.space_before = Pt(20)
        pf.space_after = Pt(0)
        pf.first_line_indent = Pt(0)
        for run in paragraph.runs:
            set_east_asia_font(run, "宋体", "Times New Roman", 14, True)
    elif level == 3:
        paragraph.style = "Heading 3"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = paragraph.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(20)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Pt(0)
        for run in paragraph.runs:
            set_east_asia_font(run, "楷体", "Times New Roman", 12, True)


def apply_page_setup(doc):
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)

        # Header / book eyebrow.
        header = section.header
        for p in header.paragraphs:
            clear_para(p)
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.first_line_indent = Pt(0)
        hp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        hp.paragraph_format.line_spacing = Pt(14)
        r = hp.add_run(f"四川大学本科毕业论文\t{SHORT_TITLE}")
        set_east_asia_font(r, "宋体", "Times New Roman", 10.5)
        add_bottom_border(hp)

        footer = section.footer
        for p in footer.paragraphs:
            clear_para(p)
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.first_line_indent = Pt(0)
        fp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fp.paragraph_format.line_spacing = Pt(14)
        add_page_field(fp)


def style_document(doc, toc_pages=None):
    apply_page_setup(doc)

    # Chinese manuscript style: use Chinese figure callouts while preserving
    # EndNote fields by editing only normal text runs.
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = run.text.replace("Supplementary Fig.", "补充图").replace("Fig.", "图")
        for idx, run in enumerate(para.runs[:-1]):
            if run.text == "Fig" and para.runs[idx + 1].text.startswith("."):
                run.text = "图"
                para.runs[idx + 1].text = para.runs[idx + 1].text[1:]

    # Base style sweep first.
    for para in doc.paragraphs:
        txt = para.text.strip()
        if paragraph_has_image(para):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para.paragraph_format.line_spacing = None
            continue
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        apply_paragraph_base(para, size=12, east_asia="宋体", first_line=True)
        if txt.startswith("Figure ") or txt.startswith("Supplementary Figure"):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Pt(0)
            for run in para.runs:
                set_east_asia_font(run, "宋体", "Times New Roman", 10.5, True)

    paras = doc.paragraphs
    if len(paras) >= 16:
        title_entries = [
            (0, TITLE, 18, "黑体", True),
            (1, "", 12, "宋体", None),
            (2, "专业：[专业名称]", 14, "楷体", None),
            (3, "学生：[作者 1]，[作者 2]，[作者 3]", 14, "楷体", None),
            (4, "指导教师：[指导教师姓名]", 14, "楷体", None),
            (5, "", 12, "宋体", None),
            (6, "四川大学本科毕业论文（设计）", 14, "黑体", True),
            (7, "", 12, "宋体", None),
            (8, "2026年6月", 12, "宋体", None),
        ]
        for idx in range(16):
            text, size, font, bold = "", 12, "宋体", None
            for spec in title_entries:
                if spec[0] == idx:
                    _, text, size, font, bold = spec
                    break
            set_para_text(paras[idx], text, size=size, east_asia=font, bold=bold)
            paras[idx].alignment = WD_ALIGN_PARAGRAPH.CENTER
            paras[idx].paragraph_format.first_line_indent = Pt(0)
            paras[idx].paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paras[idx].paragraph_format.line_spacing = Pt(20)
        # The source manuscript already contains a page break after the title
        # block. Do not add a second one, otherwise a blank page appears.

    # Convert the duplicated title heading before abstract into the formal abstract heading.
    first_big_heading_done = False
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt.startswith("再生性细胞命运由潜在调控状态混合表征"):
            set_para_text(para, "摘要", style="Heading 1", size=15, east_asia="黑体", bold=True)
            apply_heading(para, 1)
            first_big_heading_done = True
        elif first_big_heading_done and txt == "摘要":
            set_para_text(para, "", size=12)
            break

    # Insert keywords and a static TOC placeholder after the abstract body if not present.
    have_toc = any(p.text.strip() == "目录" for p in doc.paragraphs)
    if not have_toc:
        abstract_body = None
        for p in doc.paragraphs:
            if p.text.strip().startswith("修复与再生通常起始于损伤诱导的可塑性"):
                abstract_body = p
                break
        if abstract_body is not None:
            kw = insert_after(abstract_body, KEYWORDS)
            apply_paragraph_base(kw, size=10.5, east_asia="楷体", first_line=False)
            kw.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Page break after abstract keywords.
            br = insert_after(kw, "")
            br.add_run().add_break(WD_BREAK.PAGE)
            toc_title = insert_after(br, "目录", "Heading 1")
            apply_heading(toc_title, 1)
            prev = toc_title
            for h in TOC_HEADINGS:
                page = "待更新" if not toc_pages else str(toc_pages.get(h, "待更新"))
                prev = insert_after(prev, "")
                add_toc_runs(prev, h, page)
                prev.paragraph_format.first_line_indent = Pt(0)
                prev.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                prev.paragraph_format.line_spacing = Pt(20)
                prev.alignment = WD_ALIGN_PARAGRAPH.LEFT
            end = insert_after(prev, "")
            end.add_run().add_break(WD_BREAK.PAGE)

    # Update existing TOC placeholders when page numbers are known.
    if toc_pages:
        for para in doc.paragraphs:
            txt = para.text.strip()
            for h in TOC_HEADINGS:
                if txt.startswith(h + "."):
                    set_para_text(para, toc_line(h, toc_pages.get(h, "待更新")), size=12, east_asia="宋体")
                    para.paragraph_format.first_line_indent = Pt(0)
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Headings and special paragraphs.
    heading1 = {"摘要", "目录", "引言", "结果", "讨论", "局限性", "方法", "数据可用性", "代码可用性", "致谢", "作者贡献", "利益冲突", "图注", "图例", "参考文献"}
    heading2_prefixes = {
        "模型谱系", "命运控制框架", "共享可达性", "Stemness", "位置程序", "Fate-lock",
        "肿瘤样可塑性", "已弃用的标量", "潜在调控状态后验", "重叠与发散", "补充性干实验",
        "项目状态重建", "状态表征", "标量模型失效判定", "潜在调控状态混合模型",
        "轨迹与后验分析", "调控网络和扰动一致性分析", "文稿重建和图件绑定",
    }
    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt:
            continue
        if paragraph_has_image(para):
            continue
        if txt in heading1:
            apply_heading(para, 1)
        elif any(txt.startswith(prefix) for prefix in heading2_prefixes):
            apply_heading(para, 2)
        elif txt.startswith("Figure ") or txt.startswith("Supplementary Figure"):
            # Caption title lines.
            if txt.startswith("Figure "):
                new_txt = txt.replace("Figure ", "图 ", 1)
            else:
                new_txt = txt.replace("Supplementary Figure ", "附图 ", 1)
            set_para_text(para, new_txt, size=10.5, east_asia="宋体", bold=True)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Pt(0)
        elif txt.startswith("a，") or txt.startswith("a-c，") or txt.startswith("面板"):
            apply_paragraph_base(para, size=10.5, east_asia="宋体", first_line=True)
        elif txt.startswith("[") and "Author" in txt:
            apply_paragraph_base(para, size=10.5, east_asia="宋体", first_line=False)

    # Abstract body and keywords should use KaiTi 5 pt.
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt.startswith("修复与再生通常") or txt.startswith("关键词："):
            apply_paragraph_base(para, size=10.5, east_asia="楷体", first_line=(not txt.startswith("关键词：")))
            if txt.startswith("修复与再生通常"):
                para.paragraph_format.first_line_indent = Pt(42)
            else:
                para.paragraph_format.first_line_indent = Pt(0)
            for run in para.runs:
                if txt.startswith("关键词："):
                    set_east_asia_font(run, "楷体", "Times New Roman", 10.5)

    # References: compact 5 pt, hanging indent; do not change EndNote fields.
    in_refs = False
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt == "参考文献":
            in_refs = True
            continue
        if in_refs and txt:
            apply_paragraph_base(para, size=10.5, east_asia="宋体", first_line=False)
            para.paragraph_format.left_indent = Pt(18)
            para.paragraph_format.first_line_indent = Pt(-18)

    # Keep figures within the SCU text block width.
    max_width = Cm(16.2)
    for shape in doc.inline_shapes:
        if shape.width and shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = max_width
            shape.height = int(shape.height * ratio)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True)
    ap.add_argument("--output", required=True)
    ap.add_argument("--backup", required=True)
    ap.add_argument("--report", required=True)
    ap.add_argument("--toc-pages-json")
    args = ap.parse_args()

    inp = Path(args.input)
    out = Path(args.output)
    backup = Path(args.backup)
    report = Path(args.report)
    backup.parent.mkdir(parents=True, exist_ok=True)
    report.parent.mkdir(parents=True, exist_ok=True)
    if not backup.exists():
        shutil.copy2(inp, backup)

    before = audit_docx(inp)
    toc_pages = None
    if args.toc_pages_json:
        toc_pages = json.loads(Path(args.toc_pages_json).read_text(encoding="utf-8"))

    doc = Document(inp)
    style_document(doc, toc_pages=toc_pages)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    after = audit_docx(out)

    report.write_text(
        "# 四川大学本科毕业论文格式化报告\n\n"
        f"- 输入文件：`{inp}`\n"
        f"- 输出文件：`{out}`\n"
        f"- 备份文件：`{backup}`\n"
        "- 执行内容：A4 页面、页边距、页眉页脚、页码、固定 20 磅行距、中文字体层级、摘要/关键词、目录占位、标题、图注和参考文献版式。\n"
        "- 未执行内容：未伪造作者、专业、导师；未手工改写 EndNote 引文和参考文献条目。\n\n"
        "## EndNote/图片结构计数\n\n"
        "| 项目 | 处理前 | 处理后 |\n|---|---:|---:|\n"
        + "\n".join(f"| {k} | {before.get(k)} | {after.get(k)} |" for k in ["fldChar", "instrText", "ADDIN EN.CITE", "EN.CITE.DATA", "EndNote", "media_count"])
        + "\n",
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
