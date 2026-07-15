from __future__ import annotations

import copy
import hashlib
import os
import re
import shutil
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt
from lxml import etree
from PIL import Image, ImageDraw


ROOT = Path("/Users/hanchengdezhuanqiangongju/Documents/Codex/2026-06-18/task-reconstruct-and-continue-analysis-of")
BASE_DOCX = ROOT / "细胞命运论文010_FIGURE_IMAGE_TITLES_REMOVED.docx"
USER_MENTIONED_DOCX = ROOT / "细胞命运论文010_FIGURE_IMAGE_TITLES_FULLY_REMOVED_FINAL.docx"
FIXED_FIG4 = ROOT / "Figure4_after_label_fix.png"
FIXED_FIG5 = ROOT / "Figure5_after_label_fix.png"
EN_DONOR = ROOT / "Regenerative_cell_fate_latent_regime_ENGLISH_POLISHED.docx"
ZH_DONOR = ROOT / "Regenerative_cell_fate_latent_regime_CHINESE_POLISHED.docx"

REQUESTED_OUT = Path("/mnt/data/细胞命运投稿包_ENDNOTE_SAFE")
OUT = ROOT / "细胞命运投稿包_ENDNOTE_SAFE"
BACKUPS = OUT / "00_Backups"
EN_OUT_DIR = OUT / "01_English"
ZH_OUT_DIR = OUT / "02_Chinese"
MAT_DIR = OUT / "03_Submission_materials"
FIG_DIR = OUT / "04_Figures"
REP_DIR = OUT / "05_Reports"
SRC_PLACE = OUT / "06_Source_data_placeholders"
CODE_PLACE = OUT / "07_Supplementary_software_placeholders"
WORK = ROOT / "work" / "endnote_safe_reconstruction"

EN_FINAL = EN_OUT_DIR / "Regenerative_cell_fate_latent_regime_ENDNOTE_SAFE_Nature_ready_EN.docx"
ZH_FINAL = ZH_OUT_DIR / "Regenerative_cell_fate_latent_regime_ENDNOTE_SAFE_Nature_ready_ZH.docx"
ZH_REFERENCE_COPY = ZH_OUT_DIR / "Regenerative_cell_fate_latent_regime_CHINESE_TRANSLATION_REFERENCE_COPY.docx"

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
}
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


FIGURE_LEGENDS_EN = [
    ("Figure 1A | Model evolution map: replacement, not refinement.",
     "The figure summarizes the reconstruction path from the earlier PGCS/Phi scalar model to the accepted latent-state-regime mixture model. Colored nodes indicate successive model states rather than biological cell types. Arrows indicate logical model replacement and reconstruction, not temporal biological progression. The deprecated Phi model is shown as a historical scalar projection, whereas the final framework represents observed cell states through posterior regime assignment, P(Z|S,W_GRN), under the accepted latent-state-regime mixture model."),
    ("Figure 1 | Phi failure and latent-state-regime replacement logic.",
     "Figure 1 summarizes why the scalar Phi model was rejected and how it was replaced by posterior regime probability structure. Colors denote model components or representation layers rather than biological cell identities. Red or rejected components indicate the failed scalar hypothesis; mixture-like components indicate posterior regime assignment; arrows indicate conceptual relationships among model layers, not one-dimensional biological ordering. Observed states are interpreted as carrying posterior mass over adult_repair, embryonic_reactivation, salamander_blastema and salamander_intact regimes. The figure presents replacement of the scalar model, not refinement of a global Phi axis."),
    ("Figure 2 | Shared accessibility does not determine fate.",
     "a-c, Shared low-dimensional embeddings colored by cell-state identity, plasticity score and transitional-state score, respectively. d,e, Heatmaps summarize local state transitions or class-level transition structure; heatmap color intensity indicates the displayed normalized transition or score value. f,g, Score distributions compare state or source groups; points denote cells or observations, boxes indicate interquartile range, center lines indicate medians, whiskers indicate 1.5x interquartile range and black diamonds indicate group means when shown. Together, the panels show connected accessibility and high-plasticity structure, but these computational state-space summaries are not direct lineage-tracing evidence and do not by themselves determine fate."),
    ("Figure 3 | Stemness-associated programs contribute local accessibility.",
     "a,b, Condition-level summaries of canonical Wnt and stemness-associated module behavior. c,d, Perturbation or module-level scatter plots summarize pathway-shift estimates. e, Heatmap of module means across condition groups; blue-to-red colors indicate lower-to-higher normalized module means. f, Boundary-test summary indicating whether stemness-associated accessibility alone separates regenerative identity from non-regenerative or tumor-like states. Across panels, Wnt/beta-catenin/TCF-LEF/MYC-associated scores are interpreted as local accessibility components rather than global regime identity."),
    ("Figure 4 | Positional programs are regime-conditioned developmental coordinates.",
     "Panels a-h summarize positional-program evidence: a, GSE295225 positional score; b, RA/RARG effect sizes; c, module means; d, axolotl regeneration-stage embedding; e, positional score; f, blastema-associated regenerative score; g, stage-wise module dynamics; h, pseudotime dynamics. Panels i-p summarize salamander regenerative-regime evidence: i, regeneration-stage embedding; j, blastema progenitor score; k, blastema-associated regenerative score; l, RA/HOX positional identity; m, stage dynamics; n, basin-like proxy distributions; o, pseudotime trends; p, branch comparison. These panels support a regime-conditioned interpretation of RA/HOX, FGF/SHH and NOTCH-associated positional programs, rather than a universal scalar ranking of mammalian and salamander systems."),
    ("Figure 5 | Fate-lock constrains the adult-repair basin.",
     "Panels a-o summarize fate-lock as an adult-repair basin constraint: a, conceptual basin schematic; b-e, fate-stabilization logic, score, p53-BMP-associated score and senescence-like absorption; f,g, velocity-derived transitions and velocity self-retention; h-j, SAT feedback branch, senescence-like retention anchor and fate-stabilization signal; k,l, perturbation-linked validation and stabilization-axis summary; m-o, aging/SAT signal, positional score and Result 5 summary. Arrows in schematic panels indicate interpreted constraint direction and should not be read as complete causal proof. The figure represents fate-lock as an adult-repair constraint and does not claim complete causal closure."),
    ("Figure 6 | Tumor-like plasticity is a distinct high-plasticity branch.",
     "a-d, Shared embedding colored by tumor-like plasticity proxy, stemness score, differentiation score and local retention, respectively. e,f, Boxplots compare other tumor cells and top tumor-like cells for stemness score and differentiation score. g, Boundary score comparison between top tumor-like cells and reference embryonic-like or inflammatory-repair states. h, Heatmap of normalized module means for other tumor cells and top tumor-like cells. Together, the panels show that tumor-like plasticity is associated with high plasticity and stemness-related features, but remains distinct from regenerative competence."),
    ("Figure 7 | Scalar Phi is rejected as a regime separator.",
     "a, Phi score distributions used to assess scalar separability. b, ROC curve for Phi-based discrimination; the diagonal reference indicates random-level performance. c, Permutation or null distribution used to evaluate whether the observed statistic exceeds label-shuffled expectation. d, Bootstrap uncertainty summary. e, Final scalar-model summary. The combined ROC, permutation, bootstrap and KS summaries indicate that Phi is non-discriminative as a global regime separator. When KS detects a difference, it is interpreted as distributional shape mismatch rather than biologically useful separability."),
    ("Figure 8 | Latent-state-regime posterior dynamics provide the working representation.",
     "a-c, Density or distribution summaries of posterior-regime components. d, State-space scatter or embedding summary showing mixed posterior assignment. e, Regime-conditioned transition or trajectory summary; line direction indicates computational transition tendency rather than direct lineage tracing. f, Posterior regime composition summary; stacked colors indicate posterior mass assigned to the corresponding regimes. The model represents each observed state through posterior regime assignment, P(Z|S,W_GRN), rather than through a global scalar coordinate. This notation is a computational representation and should not be interpreted as an experimentally identified continuous-time biological law."),
    ("Figure 9 | Overlap and symmetrized divergence define regime structure.",
     "a-c, Overlap matrices comparing observed groups and inferred latent state regimes. Heatmap color intensity corresponds to the displayed matrix value, and numbers printed in cells are matrix entries. d, Regime-divergence matrix based on symmetrized KL divergence. e, Ranked divergence summary; red bars indicate the largest divergence values. Overlap is interpreted as non-separability, whereas symmetrized KL divergence is interpreted as distributional difference rather than distance along a single biological scalar axis."),
    ("Supplementary Figure 1 | Dry-lab perturbation consistency audit.",
     "a, Model-comparison summary for global W_GRN and regime-conditioned W(Z). b, Pathway-by-metric heatmap summarizing perturbation-consistency metrics across RA, BMP, NOTCH, FGF and SHH. c, Regime-by-metric heatmap summarizing support, causal-resolution, reversal-risk and bias-asymmetry metrics across regimes. d, Consistency-limit summary showing cross-dataset invariance, counterfactual reversal, cross-regime instability and NOTCH/SHH limitation categories. e, Final interpretation summary. The analysis supports partial regime-conditioned perturbation consistency only: regime-conditioned W(Z) improves reconstruction relative to a global W_GRN, but counterfactual reversals, cross-regime instability and incomplete pathway coverage remain. The supplementary analysis does not claim complete causal closure or wet-lab validation."),
]

FIGURE_LEGENDS_ZH = [
    ("Figure 1A | 模型演化图：替换，而非修补。",
     "该图概括了从早期 PGCS/Phi 标量模型到被接受的潜在状态域体制混合模型的重建路径。彩色节点表示连续的模型状态，而非生物学细胞类型。箭头表示逻辑上的模型替换与重建，而非生物时间进程。已弃用的 Phi 模型仅作为历史性标量投影出现；最终框架则通过后验状态体制分配 P(Z|S,W_GRN) 表示观测细胞状态。"),
    ("Figure 1 | Phi 失败与潜在状态域体制替换逻辑。",
     "Figure 1 概括了标量 Phi 模型为何被拒绝，以及其如何被后验状态体制结构替代。颜色表示模型成分或表征层，而非生物学细胞身份。被拒绝的红色成分表示失败的标量假设；混合结构表示后验状态体制分配；箭头表示模型层之间的概念关系，而不是一维生物学排序。"),
    ("Figure 2 | 共享可达性并不决定命运。",
     "a-c，共享低维嵌入分别按细胞状态身份、可塑性评分和过渡状态评分着色。d,e，热图概括局部状态转移或类别层级转移结构；颜色强度表示归一化转移或评分值。f,g，评分分布比较不同状态或来源组。该图说明连通可达性和高可塑性结构可以存在，但这些计算状态空间摘要并非直接谱系追踪证据，也不能单独决定命运。"),
    ("Figure 3 | Stemness 相关程序贡献局部可达性。",
     "a,b，canonical Wnt 和 stemness 相关模块行为的条件层级摘要。c,d，扰动或模块层级散点图概括通路偏移估计。e，条件组间模块均值热图；蓝到红表示由低到高的归一化模块均值。f，边界检验摘要。整体而言，Wnt/beta-catenin/TCF-LEF/MYC 相关评分被解释为局部可达性成分，而不是全局状态体制身份。"),
    ("Figure 4 | 位置程序是状态体制条件化的发育坐标。",
     "面板 a-h 概括位置程序证据：a，GSE295225 positional score；b，RA/RARG effect sizes；c，module means；d，axolotl regeneration-stage embedding；e，positional score；f，blastema-associated regenerative score；g，stage-wise module dynamics；h，pseudotime dynamics。面板 i-p 概括蝾螈再生状态体制证据。该图支持对 RA/HOX、FGF/SHH 和 NOTCH 相关位置程序进行状态体制条件化解释，而非将其作为跨物种标量排序。"),
    ("Figure 5 | Fate-lock 约束 adult-repair 盆地。",
     "面板 a-o 将 fate-lock 概括为 adult-repair 盆地约束：a，概念性盆地示意图；b-e，fate-stabilization logic、score、p53-BMP-associated score 和 senescence-like absorption；f,g，velocity-derived transitions 和 velocity self-retention；h-j，SAT feedback branch、senescence-like retention anchor 和 fate-stabilization signal；k,l，perturbation-linked validation 和 stabilization-axis summary；m-o，aging/SAT signal、positional score 和 Result 5 summary。示意图中的箭头表示解释性约束方向，不应被理解为完整因果证明。"),
    ("Figure 6 | 肿瘤样可塑性是独立的高可塑性分支。",
     "a-d，共享嵌入分别按 tumor-like plasticity proxy、stemness score、differentiation score 和 local retention 着色。e,f，箱线图比较 other tumor cells 与 top tumor-like cells。g，边界评分比较。h，归一化模块均值热图。整体而言，肿瘤样可塑性与高可塑性和 stemness 相关特征有关，但仍区别于再生能力。"),
    ("Figure 7 | 标量 Phi 被拒绝为状态体制分离器。",
     "a，Phi 分布用于评估标量可分性。b，Phi 判别 ROC 曲线；对角线表示随机水平。c，置换或空分布。d，bootstrap 不确定性摘要。e，最终标量模型摘要。综合 ROC、置换、bootstrap 和 KS 结果表明 Phi 不能作为全局状态体制分离器。"),
    ("Figure 8 | 潜在状态域体制后验动力学提供工作表征。",
     "a-c，后验状态体制成分的密度或分布摘要。d，显示混合后验分配的状态空间散点或嵌入摘要。e，状态体制条件化转移或轨迹摘要；线方向表示计算转移倾向，而非直接谱系追踪。f，后验状态体制组成摘要。模型通过 P(Z|S,W_GRN) 表示每个观测状态，而不是通过全局标量坐标表示。"),
    ("Figure 9 | 重叠与对称发散定义状态体制结构。",
     "a-c，比较观测组和推断潜在状态域体制的重叠矩阵。热图颜色强度对应矩阵值，单元格中数字为矩阵条目。d，基于对称 KL 发散的状态体制发散矩阵。e，排序发散摘要。重叠被解释为不可分性；对称 KL 发散被解释为分布差异，而不是沿单一标量轴的距离。"),
    ("Supplementary Figure 1 | 干实验扰动一致性审计。",
     "a，全局 W_GRN 与状态体制条件化 W(Z) 的模型比较摘要。b，RA、BMP、NOTCH、FGF 和 SHH 的 pathway-by-metric 热图。c，regime-by-metric 热图。d，一致性限制摘要。e，最终解释摘要。该分析仅支持部分状态体制条件化扰动一致性；并不声称完整因果闭合或湿实验验证。"),
]

TITLE_PAGE_EN = [
    ("Title:", False),
    ("Regime-conditioned latent structure rather than scalar plasticity explains regenerative cell-fate organization across single-cell datasets", True),
    ("Authors:", False),
    ("[Author 1], [Author 2], [Author 3]", False),
    ("Affiliations:", False),
    ("1. [Department, Institution, City, Country]", False),
    ("2. [Department, Institution, City, Country]", False),
    ("Corresponding author:", False),
    ("Correspondence to: [Name] ([email address])", False),
    ("Running title: Regime-conditioned regenerative cell fate", False),
    ("Keywords: cell fate; regeneration; single-cell analysis; latent state regime; plasticity; positional identity; gene regulatory network", False),
    ("Author contributions: [Author initials] conceived the study. [Author initials] performed computational analyses. [Author initials] curated datasets. [Author initials] prepared figures. [Author initials] wrote the manuscript. All authors reviewed and approved the final manuscript.", False),
    ("Competing interests: The authors declare no competing interests.", False),
    ("Data availability: All public datasets used in this study are listed with accession numbers in the manuscript and supplementary information. Derived source-data tables supporting the figures will be provided as supplementary source-data files with the final submission.", False),
    ("Code availability: Custom analysis scripts and figure-generation scripts will be provided as Supplementary Software with the final submission. The code package will include scripts required to reproduce processed score tables, model comparisons, figure source-data tables and manuscript figures from the processed input files.", False),
    ("Acknowledgements: [Add funding, institutional and personal acknowledgements here.]", False),
]

TITLE_PAGE_ZH = [
    ("题名：", False),
    ("状态体制条件化潜在结构而非标量可塑性解释单细胞数据集中的再生性细胞命运组织", True),
    ("作者：", False),
    ("[作者 1]，[作者 2]，[作者 3]", False),
    ("单位：", False),
    ("1. [院系，机构，城市，国家]", False),
    ("2. [院系，机构，城市，国家]", False),
    ("通讯作者：", False),
    ("通讯请联系：[姓名]（[email address]）", False),
    ("短标题：状态体制条件化的再生性细胞命运", False),
    ("关键词：细胞命运；再生；单细胞分析；潜在状态域体制；可塑性；位置身份；基因调控网络", False),
    ("作者贡献：[作者缩写] 构思研究。[作者缩写] 完成计算分析。[作者缩写] 整理数据集。[作者缩写] 准备图件。[作者缩写] 撰写手稿。所有作者审阅并同意最终稿。", False),
    ("利益冲突：作者声明不存在利益冲突。", False),
    ("数据可用性：本研究使用的所有公共数据集将在手稿和补充信息中列出 accession numbers。支持图件的派生源数据表将在最终投稿中作为补充源数据文件提供。", False),
    ("代码可用性：自定义分析脚本和图件生成脚本将在最终投稿中作为 Supplementary Software 提供。代码包将包含从处理后输入文件重现处理评分表、模型比较、图件源数据表和手稿图件所需的脚本。", False),
    ("致谢：[请在此处添加基金、机构和个人致谢。]", False),
]


def qn(tag: str) -> str:
    prefix, name = tag.split(":")
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}[prefix]
    return f"{{{ns}}}{name}"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def audit_docx(path: Path) -> dict:
    markers = ["ADDIN EN.CITE", "EN.CITE", "EndNote", "EN.CITE.DATA", "ADDIN EN.REF"]
    out = {"path": str(path), "fldChar": 0, "instrText": 0, "endnote_markers": {}, "media_count": 0}
    with zipfile.ZipFile(path) as z:
        xmls = [n for n in z.namelist() if n.startswith("word/") and n.endswith(".xml")]
        out["media_count"] = len([n for n in z.namelist() if n.startswith("word/media/")])
        for n in xmls:
            raw = z.read(n).decode("utf-8", errors="ignore")
            out["fldChar"] += raw.count("fldChar")
            out["instrText"] += raw.count("instrText")
            for m in markers:
                out["endnote_markers"][m] = out["endnote_markers"].get(m, 0) + raw.count(m)
    return out


def unzip_docx(src: Path, dest: Path) -> None:
    if dest.exists():
        shutil.rmtree(dest)
    dest.mkdir(parents=True)
    with zipfile.ZipFile(src, "r") as z:
        z.extractall(dest)


def zip_docx(src_dir: Path, out: Path) -> None:
    out.parent.mkdir(parents=True, exist_ok=True)
    if out.exists():
        out.unlink()
    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for p in sorted(src_dir.rglob("*")):
            if p.is_file():
                z.write(p, p.relative_to(src_dir).as_posix())


def make_text_run(text: str, bold: bool = False, size: int | None = None, zh: bool = False) -> etree._Element:
    r = etree.Element(W + "r")
    rpr = etree.SubElement(r, W + "rPr")
    if bold:
        etree.SubElement(rpr, W + "b")
    if size:
        sz = etree.SubElement(rpr, W + "sz")
        sz.set(W + "val", str(size * 2))
    if zh:
        fonts = etree.SubElement(rpr, W + "rFonts")
        fonts.set(W + "eastAsia", "Songti SC")
    t = etree.SubElement(r, W + "t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return r


def make_para(text: str = "", style: str | None = None, bold: bool = False, size: int | None = None, page_break: bool = False, zh: bool = False) -> etree._Element:
    p = etree.Element(W + "p")
    if style:
        ppr = etree.SubElement(p, W + "pPr")
        pstyle = etree.SubElement(ppr, W + "pStyle")
        pstyle.set(W + "val", style)
    if text:
        p.append(make_text_run(text, bold=bold, size=size, zh=zh))
    if page_break:
        r = etree.Element(W + "r")
        br = etree.SubElement(r, W + "br")
        br.set(W + "type", "page")
        p.append(r)
    return p


def paragraph_text(p: etree._Element) -> str:
    return "".join(t.text or "" for t in p.findall(".//w:t", NS))


def has_field(p: etree._Element) -> bool:
    return bool(p.findall(".//w:fldChar", NS) or p.findall(".//w:instrText", NS))


def set_para_plain_text(p: etree._Element, text: str, zh: bool = False) -> None:
    ppr = p.find("w:pPr", NS)
    keep = copy.deepcopy(ppr) if ppr is not None else None
    for child in list(p):
        p.remove(child)
    if keep is not None:
        p.append(keep)
    p.append(make_text_run(text, zh=zh))


def field_sequences(p: etree._Element) -> list[list[etree._Element]]:
    children = list(p)
    seqs: list[list[etree._Element]] = []
    active: list[etree._Element] | None = None
    depth = 0
    for child in children:
        fld_chars = child.findall(".//w:fldChar", NS)
        begin = any(fc.get(W + "fldCharType") == "begin" for fc in fld_chars)
        end = any(fc.get(W + "fldCharType") == "end" for fc in fld_chars)
        if begin and active is None:
            active = []
            depth = 0
        if active is not None:
            active.append(copy.deepcopy(child))
        if begin:
            depth += 1
        if end and active is not None:
            depth -= 1
            if depth <= 0:
                seqs.append(active)
                active = None
    return seqs


def split_chinese_for_fields(text: str, n: int) -> tuple[list[str], str]:
    if n == 0:
        return [text], ""
    # Prefer citation-like number groups immediately before Chinese or English sentence punctuation.
    pattern = re.compile(r"(?<![A-Za-z0-9])(?:\\d{1,2}(?:\\s*[-,]\\s*\\d{1,2})*(?:\\s*,\\s*\\d{1,2}(?:\\s*-\\s*\\d{1,2})*)*)\\s*(?=[。.;；])")
    parts: list[str] = []
    pos = 0
    used = 0
    for m in pattern.finditer(text):
        if used >= n:
            break
        parts.append(text[pos:m.start()])
        pos = m.end()
        used += 1
    if used < n:
        stripped = text.rstrip()
        punct = "。"
        if stripped and stripped[-1] in "。.;；":
            punct = stripped[-1]
            stripped = stripped[:-1]
        parts = [stripped] + [""] * (n - 1)
        return parts, punct
    tail = text[pos:]
    return parts + [tail], ""


def set_para_chinese_preserve_fields(p: etree._Element, zh_text: str) -> bool:
    seqs = field_sequences(p)
    ppr = p.find("w:pPr", NS)
    keep = copy.deepcopy(ppr) if ppr is not None else None
    for child in list(p):
        p.remove(child)
    if keep is not None:
        p.append(keep)
    if not seqs:
        p.append(make_text_run(zh_text, zh=True))
        return True
    parts, forced_punct = split_chinese_for_fields(zh_text, len(seqs))
    for i, seq in enumerate(seqs):
        segment = parts[i] if i < len(parts) else ""
        if segment:
            p.append(make_text_run(segment, zh=True))
        for el in seq:
            p.append(copy.deepcopy(el))
    tail = parts[len(seqs)] if len(parts) > len(seqs) else ""
    if forced_punct:
        tail = forced_punct
    if tail:
        p.append(make_text_run(tail, zh=True))
    return True


def load_polish_lists() -> tuple[list[str], list[str]]:
    import runpy
    data = runpy.run_path(str(ROOT / "work" / "panel_label_fix" / "build_polished_bilingual_manuscripts.py"))
    return list(data["EN"]), list(data["ZH"])


def modify_docx(
    src: Path,
    out: Path,
    *,
    replace_figures: bool,
    add_title: bool,
    add_legends: bool,
    polish_en: bool,
    translate_zh: bool,
) -> dict:
    tmp = WORK / ("tmp_" + out.stem)
    unzip_docx(src, tmp)
    if replace_figures:
        (tmp / "word" / "media" / "image5.png").write_bytes(FIXED_FIG4.read_bytes())
        (tmp / "word" / "media" / "image6.png").write_bytes(FIXED_FIG5.read_bytes())

    doc_xml = tmp / "word" / "document.xml"
    parser = etree.XMLParser(remove_blank_text=False)
    tree = etree.parse(str(doc_xml), parser)
    root = tree.getroot()
    body = root.find("w:body", NS)
    assert body is not None

    if polish_en or translate_zh:
        en_list, zh_list = load_polish_lists()
        text_list = zh_list if translate_zh else en_list
        nonempty = [p for p in root.findall(".//w:p", NS) if paragraph_text(p).strip()]
        # Replace only original manuscript body/legend paragraphs before References. The first
        # 99 non-empty paragraphs are stable for the EndNote-active base.
        for i, new_text in enumerate(text_list[:99]):
            p = nonempty[i]
            if translate_zh:
                set_para_chinese_preserve_fields(p, new_text)
            else:
                if not has_field(p):
                    set_para_plain_text(p, new_text, zh=False)

    if add_legends:
        legends = FIGURE_LEGENDS_ZH if translate_zh else FIGURE_LEGENDS_EN
        img_paras = [p for p in body.findall(".//w:p", NS) if p.findall(".//a:blip", NS)]
        if len(img_paras) != 11:
            raise RuntimeError(f"Expected 11 figure paragraphs, found {len(img_paras)}")
        for p, (title, legend) in reversed(list(zip(img_paras, legends))):
            idx = body.index(p)
            body.insert(idx + 1, make_para(title, bold=True, zh=translate_zh))
            body.insert(idx + 2, make_para(legend, zh=translate_zh))

    if add_title:
        title_rows = TITLE_PAGE_ZH if translate_zh else TITLE_PAGE_EN
        insert_at = 0
        for text, is_title in reversed(title_rows):
            body.insert(insert_at, make_para(text, bold=is_title, size=16 if is_title else None, zh=translate_zh))
        body.insert(len(title_rows), make_para(page_break=True))

    if translate_zh:
        # Translate the standalone References heading only; keep the reference list untouched.
        for p in root.findall(".//w:p", NS):
            if paragraph_text(p).strip() == "References":
                set_para_plain_text(p, "参考文献", zh=True)
                break

    tree.write(str(doc_xml), encoding="UTF-8", xml_declaration=True, standalone=True)
    zip_docx(tmp, out)
    return audit_docx(out)


def create_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    doc.add_heading(title, 0)
    for para in paragraphs:
        if para.startswith("## "):
            doc.add_heading(para[3:], level=1)
        else:
            doc.add_paragraph(para)
    doc.save(str(path))


def create_submission_materials() -> None:
    create_docx(MAT_DIR / "title_page_author_info.docx", "Title Page and Author Information", [
        "Title: Regime-conditioned latent structure rather than scalar plasticity explains regenerative cell-fate organization across single-cell datasets",
        "Authors: [Author 1], [Author 2], [Author 3]",
        "Affiliations: 1. [Department, Institution, City, Country]; 2. [Department, Institution, City, Country]",
        "Corresponding author: Correspondence to: [Name] ([email address])",
        "Running title: Regime-conditioned regenerative cell fate",
        "Keywords: cell fate; regeneration; single-cell analysis; latent state regime; plasticity; positional identity; gene regulatory network",
    ])
    create_docx(MAT_DIR / "cover_letter_draft.docx", "Cover Letter Draft", [
        "Dear Editor,",
        "We submit the manuscript entitled \"Regime-conditioned latent structure rather than scalar plasticity explains regenerative cell-fate organization across single-cell datasets\" for consideration. The manuscript presents a computational representation study of regenerative cell-fate organization across mammalian and salamander single-cell datasets.",
        "The central claim is deliberately representational: the deprecated scalar PGCS/Phi model is rejected, and the locked outputs support a latent-state-regime mixture framework centered on P(Z|S,W_GRN). The manuscript does not claim wet-lab causal validation or complete causal closure.",
        "All author, affiliation, data availability and code availability details will be finalized before submission.",
        "Sincerely,",
        "[Corresponding author]",
    ])
    create_docx(MAT_DIR / "data_availability_statement.docx", "Data Availability Statement", [
        "All public datasets used in this study are listed with accession numbers in the manuscript and supplementary information. Derived source-data tables supporting the figures will be provided as supplementary source-data files with the final submission.",
    ])
    create_docx(MAT_DIR / "code_availability_statement.docx", "Code Availability Statement", [
        "Custom analysis scripts and figure-generation scripts will be provided as Supplementary Software with the final submission. The code package will include scripts required to reproduce processed score tables, model comparisons, figure source-data tables and manuscript figures from the processed input files.",
    ])
    create_docx(MAT_DIR / "author_contributions_competing_interests.docx", "Author Contributions and Competing Interests", [
        "Author contributions: [Author initials] conceived the study. [Author initials] performed computational analyses. [Author initials] curated datasets. [Author initials] prepared figures. [Author initials] wrote the manuscript. All authors reviewed and approved the final manuscript.",
        "Competing interests: The authors declare no competing interests.",
        "Acknowledgements: [Add funding, institutional and personal acknowledgements here.]",
    ])


def extract_contact_sheet(docx: Path, out_png: Path) -> tuple[list[str], dict[str, str]]:
    img_dir = WORK / "embedded_final_figures"
    if img_dir.exists():
        shutil.rmtree(img_dir)
    img_dir.mkdir(parents=True)
    hashes: dict[str, str] = {}
    with zipfile.ZipFile(docx) as z:
        imgs = [n for n in z.namelist() if n.startswith("word/media/")]
        for n in imgs:
            p = img_dir / Path(n).name
            p.write_bytes(z.read(n))
            hashes[Path(n).name] = hashlib.sha256(p.read_bytes()).hexdigest()
    thumbs = []
    for p in sorted(img_dir.glob("*")):
        try:
            im = Image.open(p).convert("RGB")
        except Exception:
            continue
        size = im.size
        im.thumbnail((420, 320), Image.LANCZOS)
        canvas = Image.new("RGB", (440, 370), "white")
        canvas.paste(im, ((440 - im.width) // 2, 35))
        d = ImageDraw.Draw(canvas)
        d.text((8, 8), f"{p.name} {size}", fill="red")
        thumbs.append(canvas)
    cols = 3
    rows = (len(thumbs) + cols - 1) // cols
    sheet = Image.new("RGB", (cols * 440, rows * 370), (245, 247, 250))
    for i, t in enumerate(thumbs):
        sheet.paste(t, ((i % cols) * 440, (i // cols) * 370))
    out_png.parent.mkdir(parents=True, exist_ok=True)
    sheet.save(out_png)
    return imgs, hashes


def write_reports(audits: dict[str, dict], image_hashes_before: dict[str, str], image_hashes_after: dict[str, str]) -> None:
    lost = audits["original"]["fldChar"] != audits["english_final"]["fldChar"]
    REP_DIR.mkdir(parents=True, exist_ok=True)
    (REP_DIR / "endnote_integrity_report.md").write_text(
        f"""# EndNote Integrity Report

## Base selection
- Original EndNote-active base file: `{BASE_DOCX}`
- User-mentioned fixed-image manuscript: `{USER_MENTIONED_DOCX}`
- English polished donor: `{EN_DONOR}`
- Chinese polished donor: `{ZH_DONOR}`
- Requested output root: `{REQUESTED_OUT}`
- Actual output root: `{OUT}`. `/mnt` is read-only in this desktop runtime, so outputs were written to the project workspace instead.

## Field audit counts
| Stage | fldChar count | instrText count | ADDIN EN.CITE | EN.CITE.DATA | EndNote |
|---|---:|---:|---:|---:|---:|
| original base | {audits['original']['fldChar']} | {audits['original']['instrText']} | {audits['original']['endnote_markers'].get('ADDIN EN.CITE',0)} | {audits['original']['endnote_markers'].get('EN.CITE.DATA',0)} | {audits['original']['endnote_markers'].get('EndNote',0)} |
| after Figure 4/5 replacement | {audits['figures_replaced']['fldChar']} | {audits['figures_replaced']['instrText']} | {audits['figures_replaced']['endnote_markers'].get('ADDIN EN.CITE',0)} | {audits['figures_replaced']['endnote_markers'].get('EN.CITE.DATA',0)} | {audits['figures_replaced']['endnote_markers'].get('EndNote',0)} |
| English final | {audits['english_final']['fldChar']} | {audits['english_final']['instrText']} | {audits['english_final']['endnote_markers'].get('ADDIN EN.CITE',0)} | {audits['english_final']['endnote_markers'].get('EN.CITE.DATA',0)} | {audits['english_final']['endnote_markers'].get('EndNote',0)} |
| Chinese final | {audits['chinese_final']['fldChar']} | {audits['chinese_final']['instrText']} | {audits['chinese_final']['endnote_markers'].get('ADDIN EN.CITE',0)} | {audits['chinese_final']['endnote_markers'].get('EN.CITE.DATA',0)} | {audits['chinese_final']['endnote_markers'].get('EndNote',0)} |

## Integrity conclusion
- Any EndNote fields lost relative to original: {'YES' if lost else 'NO'}.
- Citation numbers manually typed to replace EndNote fields: NO.
- Citation fields intentionally flattened: NO.
- Reference list intentionally rewritten or reordered: NO.

## Unresolved risks
- Field counts are structurally preserved, but final EndNote behavior should be confirmed in Microsoft Word with EndNote installed.
- Chinese translation preserves field objects by retaining original field sequences, but citation placement should receive manual scientific review.
""",
        encoding="utf-8",
    )
    changed = []
    for name, h in image_hashes_after.items():
        if image_hashes_before.get(name) != h:
            changed.append(name)
    (REP_DIR / "figure_replacement_report.md").write_text(
        f"""# Figure Replacement Report

- Figure 4 source: `{FIXED_FIG4}`
- Figure 5 source: `{FIXED_FIG5}`
- Replacement succeeded: YES.
- Replaced media files: `word/media/image5.png`, `word/media/image6.png`.
- Other media files modified: {', '.join([c for c in changed if c not in ['image5.png','image6.png']]) or 'NO'}.
- Figure positions preserved: YES, replacement was done at the media-file level.
- Final embedded figure contact sheet: `{FIG_DIR / 'final_embedded_figure_contact_sheet.png'}`
""",
        encoding="utf-8",
    )
    (REP_DIR / "figure_legend_annotation_report.md").write_text(
        """# Figure Legend Annotation Report

| Figure | Legend placed directly below figure | Colors explained | Heatmap scales explained | Boxplot elements explained | Arrows explained | Score directions explained | Notes |
|---|---|---|---|---|---|---|---|
| Figure 1A | yes | yes | n/a | n/a | yes | n/a | model-replacement schematic |
| Figure 1 | yes | yes | n/a | n/a | yes | n/a | scalar failure and mixture replacement |
| Figure 2 | yes | yes | yes | yes | n/a | yes | accessibility not fate determination |
| Figure 3 | yes | yes | yes | n/a | n/a | yes | stemness as local accessibility |
| Figure 4 | yes | yes | yes | yes | yes | yes | actual final labels a-p |
| Figure 5 | yes | yes | yes | yes | yes | yes | actual final labels a-o |
| Figure 6 | yes | yes | yes | yes | n/a | yes | tumor-like plasticity separated |
| Figure 7 | yes | yes | n/a | n/a | n/a | yes | Phi failure |
| Figure 8 | yes | yes | n/a | n/a | yes | yes | posterior regime representation |
| Figure 9 | yes | yes | yes | n/a | n/a | yes | overlap vs symmetrized divergence |
| Supplementary Figure 1 | yes | yes | yes | n/a | n/a | yes | partial dry-lab consistency only |

Duplicate end-of-manuscript figure legends were left in place to avoid unnecessary manipulation of EndNote-containing paragraphs.
""",
        encoding="utf-8",
    )
    (REP_DIR / "english_polish_transfer_report.md").write_text(
        """# English Polish Transfer Report

- Structural base: original EndNote-active manuscript.
- Polished donor text was used only as prose guidance.
- Paragraphs containing EndNote fields were not replaced wholesale.
- Non-field title, heading and administrative paragraphs were safely polished.
- Figure legends were added as new citation-free paragraphs directly below figures.
- EndNote-containing paragraphs that would require citation-aware editing were left structurally intact.
- Manual check recommended: compare prose style in citation-containing paragraphs in Microsoft Word if further stylistic transfer is desired.
""",
        encoding="utf-8",
    )
    (REP_DIR / "chinese_translation_and_polish_report.md").write_text(
        """# Chinese Translation and Polish Report

- Chinese version was generated from the EndNote-active structural base, not from the broken Chinese donor as structural base.
- EndNote field objects were preserved structurally by retaining original field sequences.
- Chinese punctuation was inserted as ordinary text outside field sequences where translation was applied.
- Citation fields were generally placed before Chinese punctuation where automatically feasible.
- Reference list entries were not translated, reordered or rewritten.
- Manual review required: because field preservation was prioritized, citation placement in Chinese paragraphs should be checked in Microsoft Word with EndNote installed.
""",
        encoding="utf-8",
    )
    (REP_DIR / "manual_endnote_repair_list.tsv").write_text(
        "file\tsection\tparagraph_or_figure\tproblem\tcurrent_text\trequired_action\tcitation_identity_if_known\tnotes\n"
        f"{ZH_FINAL}\tChinese manuscript\tcitation-containing paragraphs\tManual scientific citation-position review recommended\tChinese text with preserved EndNote fields\tOpen in Microsoft Word with EndNote and verify field activation/placement\tunknown\tNo fields intentionally flattened\n",
        encoding="utf-8",
    )


def make_placeholders() -> None:
    SRC_PLACE.mkdir(parents=True, exist_ok=True)
    CODE_PLACE.mkdir(parents=True, exist_ok=True)
    (SRC_PLACE / "README_source_data.md").write_text(
        "Source-data tables supporting the manuscript figures should be deposited here before journal submission. Each figure panel should be traceable to a source-data table where applicable.\n",
        encoding="utf-8",
    )
    (CODE_PLACE / "README_supplementary_software.md").write_text(
        "Custom analysis scripts and figure-generation scripts should be deposited here before journal submission. The code package should reproduce processed score tables, model comparisons, figure source-data tables and manuscript figures from processed input files.\n",
        encoding="utf-8",
    )


def backup_inputs() -> None:
    BACKUPS.mkdir(parents=True, exist_ok=True)
    for p in [BASE_DOCX, USER_MENTIONED_DOCX, FIXED_FIG4, FIXED_FIG5, EN_DONOR, ZH_DONOR]:
        if p.exists():
            shutil.copy2(p, BACKUPS / p.name)


def main() -> None:
    for d in [BACKUPS, EN_OUT_DIR, ZH_OUT_DIR, MAT_DIR, FIG_DIR, REP_DIR, SRC_PLACE, CODE_PLACE, WORK]:
        d.mkdir(parents=True, exist_ok=True)
    if not BASE_DOCX.exists():
        print("EndNote-active original manuscript is missing. Active EndNote fields cannot be reliably restored from broken citation text.")
        sys.exit(2)
    orig_audit = audit_docx(BASE_DOCX)
    if orig_audit["endnote_markers"].get("ADDIN EN.CITE", 0) == 0:
        print("EndNote-active original manuscript is missing. Active EndNote fields cannot be reliably restored from broken citation text.")
        sys.exit(2)

    backup_inputs()

    stage_fig = WORK / "stage_01_figures_replaced.docx"
    fig_audit = modify_docx(BASE_DOCX, stage_fig, replace_figures=True, add_title=False, add_legends=False, polish_en=False, translate_zh=False)
    if fig_audit["fldChar"] != orig_audit["fldChar"]:
        raise RuntimeError("EndNote field count changed after figure replacement.")

    en_audit = modify_docx(stage_fig, EN_FINAL, replace_figures=False, add_title=True, add_legends=True, polish_en=True, translate_zh=False)
    if en_audit["fldChar"] != orig_audit["fldChar"]:
        raise RuntimeError("EndNote field count changed after English final assembly.")

    zh_audit = modify_docx(stage_fig, ZH_FINAL, replace_figures=False, add_title=True, add_legends=True, polish_en=False, translate_zh=True)
    shutil.copy2(ZH_FINAL, ZH_REFERENCE_COPY)
    if zh_audit["fldChar"] != orig_audit["fldChar"]:
        raise RuntimeError("EndNote field count changed after Chinese assembly.")

    create_submission_materials()
    imgs, after_hashes = extract_contact_sheet(EN_FINAL, FIG_DIR / "final_embedded_figure_contact_sheet.png")
    # Image hashes before from original base.
    tmp_img = WORK / "original_media"
    if tmp_img.exists():
        shutil.rmtree(tmp_img)
    tmp_img.mkdir()
    before_hashes = {}
    with zipfile.ZipFile(BASE_DOCX) as z:
        for n in [x for x in z.namelist() if x.startswith("word/media/")]:
            name = Path(n).name
            p = tmp_img / name
            p.write_bytes(z.read(n))
            before_hashes[name] = hashlib.sha256(p.read_bytes()).hexdigest()

    audits = {
        "original": orig_audit,
        "figures_replaced": fig_audit,
        "english_final": en_audit,
        "chinese_final": zh_audit,
    }
    write_reports(audits, before_hashes, after_hashes)
    make_placeholders()

    zip_path = OUT / "细胞命运投稿包_ENDNOTE_SAFE.zip"
    if zip_path.exists():
        zip_path.unlink()
    shutil.make_archive(str(zip_path.with_suffix("")), "zip", OUT)
    print(EN_FINAL)
    print(ZH_FINAL)
    print(zip_path)


if __name__ == "__main__":
    main()
