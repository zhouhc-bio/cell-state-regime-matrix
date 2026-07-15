from __future__ import annotations

import csv
import hashlib
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell
from docx.text.paragraph import Paragraph

try:
    from PIL import Image
except Exception:  # pragma: no cover
    Image = None


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
PKG = ROOT / "FINAL_SUBMISSION_PACKAGE"

BASE_DOCX = OUT / "submission_ready_manuscript_causal_consistency_updated.docx"
BASE_MD = OUT / "main_text_clean_version_causal_updated.md"
LEGEND_DOCX = OUT / "figure_legend_final_standardized.docx"
SUPP_CAPTION_MD = OUT / "supplementary_causal_closure_caption.md"

FINAL_SAFE_MD = OUT / "final_claim_safe_main_text.md"
FINAL_MANUSCRIPT = OUT / "FINAL_MANUSCRIPT_WITH_MAIN_FIGURES.docx"
FINAL_LEGENDS = OUT / "FINAL_FIGURE_LEGENDS.docx"
SUPP_DOCX = OUT / "SUPPLEMENTARY_MATERIAL.docx"
SOURCE_INDEX = OUT / "SUPPLEMENTARY_SOURCE_DATA_INDEX.tsv"
UNRESOLVED_REFS = OUT / "unresolved_reference_list.md"
ENDNOTE_TODO = OUT / "endnote_import_todo.md"
READINESS = OUT / "final_submission_readiness_report.md"

MAIN_FIG_DIR = OUT / "final_9figures_reconstructed"
SUPP_FIG_PNG = OUT / "FIGURE_CAUSAL_CLOSURE_FINAL.png"
SUPP_FIG_PDF = OUT / "FIGURE_CAUSAL_CLOSURE_FINAL.pdf"


MAIN_FIGURES = [
    (1, MAIN_FIG_DIR / "Final_Figure_1_Phi_failure_latent_regime_replacement.png"),
    (2, MAIN_FIG_DIR / "Final_Figure_2_shared_accessibility_reannotated.png"),
    (3, MAIN_FIG_DIR / "Final_Figure_3_stemness_local_accessibility.png"),
    (4, MAIN_FIG_DIR / "Final_Figure_4_regime_conditioned_positional_information.png"),
    (5, MAIN_FIG_DIR / "Final_Figure_5_fate_lock_adult_repair_basin.png"),
    (6, MAIN_FIG_DIR / "Final_Figure_6_tumor_like_plasticity_distinct_branch.png"),
    (7, MAIN_FIG_DIR / "Final_Figure_7_single_Phi_rejection.png"),
    (8, MAIN_FIG_DIR / "Final_Figure_8_latent_regime_posterior_dynamics.png"),
    (9, MAIN_FIG_DIR / "Final_Figure_9_overlap_symmetrized_divergence.png"),
]

SUPP_TABLES = [
    (1, OUT / "perturbation_deltaZ.tsv", "Pathway perturbation posterior shift summaries."),
    (2, OUT / "closure_model_comparison.tsv", "Global W versus regime-conditioned W(Z) closure-model comparison."),
    (3, OUT / "regime_conditioned_deltaZ_consistency.tsv", "Regime-conditioned deltaZ consistency by pathway and latent state regime."),
    (4, OUT / "counterfactual_bias_by_regime.tsv", "Counterfactual bias and direction-reversal summaries by regime."),
    (5, OUT / "edge_stability_across_regimes.tsv", "Regime-wise regulatory edge stability across W(Z) matrices."),
]

EXTRA_SUPP_SOURCE = [
    OUT / "causal_closure_score_summary.tsv",
    OUT / "cross_dataset_perturbation_similarity.tsv",
    OUT / "null_perturbation_control_results.tsv",
    OUT / "counterfactual_consistency_scores.tsv",
    OUT / "regime_conditioned_grn_summary.tsv",
    OUT / "W_matrix_by_regime.tsv",
    OUT / "pathway_response_conditioning_gates.tsv",
    OUT / "regime_axis_conditioning_gates.tsv",
    OUT / "perturbation_posterior_transfer_normalization.tsv",
]

FORBIDDEN_MAIN = [
    "causal proof",
    "causal closure established",
    "perturbations drive",
    "experimentally validated",
    "W(Z) demonstrates causality",
    "causal mechanism confirmed",
    "full causal closure",
]


def claim_safe_text(text: str) -> str:
    replacements = [
        (
            "Counterfactual direction reversals remained frequent (0.7), so these outputs support partial dry-lab perturbation consistency in regime-conditioned subspaces, not FULL_CLOSURE.",
            "Counterfactual direction reversals remained frequent (0.7), so these outputs support partial regime-conditioned perturbation consistency only.",
        ),
        (
            "the minimal sufficient representation supported by the locked outputs",
            "the minimal working representation supported by the locked outputs and evaluated baselines",
        ),
        (
            "the minimal sufficient representation supported by the current outputs",
            "the minimal working representation supported by the locked outputs and evaluated baselines",
        ),
        (
            "regime mixture is the minimal sufficient representation",
            "regime mixture is the minimal working representation supported by the locked outputs and evaluated baselines",
        ),
        (
            "provide the minimal sufficient representation",
            "provide the locked working representation",
        ),
        (
            "minimal sufficient representation",
            "minimal working representation supported by the locked outputs and evaluated baselines",
        ),
        (
            "not FULL_CLOSURE",
            "not full computational causal-consistency support",
        ),
        (
            "representation closure from causal closure",
            "representational consistency from computational causal-consistency support",
        ),
        (
            "A complete causal-closure claim",
            "A complete causal claim",
        ),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def set_base_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if name != "Heading 1" else 16)
        style.paragraph_format.space_after = Pt(6)


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.font.name = "Calibri"
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("1F4D78")
    p.paragraph_format.space_after = Pt(4)
    if subtitle:
        sp = doc.add_paragraph()
        sr = sp.add_run(subtitle)
        sr.font.size = Pt(10)
        sr.font.color.rgb = RGBColor.from_string("555555")
        sp.paragraph_format.space_after = Pt(12)


def insert_paragraph_after(paragraph: Paragraph, text: str | None = None, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def image_width(path: Path, max_width_in: float = 6.35, max_height_in: float = 6.45) -> float:
    if Image is None:
        return max_width_in
    with Image.open(path) as im:
        w, h = im.size
    width_for_height = max_height_in * (w / h)
    return min(max_width_in, width_for_height)


def add_picture_paragraph_after(anchor: Paragraph, image_path: Path, caption: str) -> Paragraph:
    fig_p = insert_paragraph_after(anchor)
    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_p.paragraph_format.page_break_before = True
    fig_p.paragraph_format.space_after = Pt(6)
    width = image_width(image_path)
    fig_p.add_run().add_picture(str(image_path), width=Inches(width))
    cap_p = insert_paragraph_after(fig_p)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap_p.paragraph_format.space_after = Pt(8)
    cap_run = cap_p.add_run(caption)
    cap_run.font.size = Pt(9)
    cap_run.font.italic = True
    return cap_p


def extract_main_legends() -> dict[int, str]:
    doc = Document(LEGEND_DOCX)
    legends: dict[int, str] = {}
    paras = [claim_safe_text(p.text.strip()) for p in doc.paragraphs if p.text.strip()]
    for idx, text in enumerate(paras):
        m = re.match(r"Figure\s+(\d+)\s+\|", text)
        if m:
            number = int(m.group(1))
            body = paras[idx + 1] if idx + 1 < len(paras) else ""
            text = text.replace("Scalar embedding failure", "Deprecated scalar proxy failure")
            text = text.replace("Scalar proxy", "Deprecated scalar proxy")
            text = text.replace("Phi", "deprecated scalar proxy")
            body = body.replace("Phi", "deprecated scalar proxy")
            legends[number] = f"{text} {body}".strip()
    return legends


def write_claim_safe_md() -> None:
    text = BASE_MD.read_text(encoding="utf-8")
    FINAL_SAFE_MD.write_text(claim_safe_text(text), encoding="utf-8")


def update_docx_claims(doc: Document) -> None:
    for p in doc.paragraphs:
        new_text = claim_safe_text(p.text)
        if new_text != p.text:
            p.text = new_text


def result_anchor(doc: Document, fig_number: int) -> Paragraph:
    starts: dict[int, int] = {}
    for i, p in enumerate(doc.paragraphs):
        m = re.match(rf"Result\s+(\d+)\.", p.text.strip())
        if m:
            starts[int(m.group(1))] = i
    if fig_number not in starts:
        raise RuntimeError(f"Could not locate Result {fig_number}.")
    start = starts[fig_number]
    next_indices = [idx for n, idx in starts.items() if idx > start]
    discussion = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Discussion"), len(doc.paragraphs))
    end = min(next_indices + [discussion]) if next_indices or discussion else len(doc.paragraphs)
    candidates = [p for p in doc.paragraphs[start + 1 : end] if p.text.strip()]
    if fig_number == 9:
        for p in candidates:
            if "Figure 9 closes" in p.text:
                return p
    if not candidates:
        return doc.paragraphs[start]
    return candidates[-1]


def build_manuscript_with_figures() -> None:
    doc = Document(BASE_DOCX)
    update_docx_claims(doc)
    legends = extract_main_legends()
    for number, path in reversed(MAIN_FIGURES):
        if not path.exists():
            continue
        anchor = result_anchor(doc, number)
        caption = legends.get(number, f"Figure {number}.")
        add_picture_paragraph_after(anchor, path, caption)
    doc.save(FINAL_MANUSCRIPT)


def build_final_legends_docx() -> None:
    legends = extract_main_legends()
    doc = Document()
    set_base_styles(doc)
    add_title(doc, "Final Figure Legends", "Locked 9-main-figure architecture; latent-state-regime mixture terminology standardized.")
    for i in range(1, 10):
        p = doc.add_paragraph(style="Heading 2")
        p.add_run(f"Figure {i}")
        doc.add_paragraph(legends.get(i, f"Figure {i} legend missing."))
    note = doc.add_paragraph(style="Heading 2")
    note.add_run("Supplementary figure note")
    doc.add_paragraph(
        "Supplementary Fig. 1 is reported in SUPPLEMENTARY_MATERIAL.docx and uses partial perturbation consistency / PARTIAL_CLOSURE wording only."
    )
    doc.save(FINAL_LEGENDS)


def read_tsv(path: Path, limit_rows: int | None = None) -> list[list[str]]:
    rows: list[list[str]] = []
    if not path.exists():
        return rows
    with path.open("r", encoding="utf-8", newline="") as fh:
        reader = csv.reader(fh, delimiter="\t")
        for i, row in enumerate(reader):
            rows.append(row)
            if limit_rows is not None and i >= limit_rows:
                break
    return rows


def tsv_shape(path: Path) -> tuple[int, int]:
    rows = read_tsv(path)
    if not rows:
        return (0, 0)
    return (max(len(rows) - 1, 0), len(rows[0]))


def add_small_table(doc: Document, rows: list[list[str]], max_cols: int = 6) -> None:
    if not rows:
        doc.add_paragraph("File missing or empty.")
        return
    selected = [row[:max_cols] for row in rows]
    n_cols = max(len(r) for r in selected)
    table = doc.add_table(rows=len(selected), cols=n_cols)
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(selected):
        for c_idx in range(n_cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = row[c_idx] if c_idx < len(row) else ""
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(7)
                para.paragraph_format.space_after = Pt(0)
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    doc.add_paragraph()


def build_supplementary_material() -> None:
    doc = Document()
    set_base_styles(doc)
    add_title(doc, "Supplementary Material", "Dry-lab perturbation consistency package; no new analysis or main-figure renumbering.")
    doc.add_heading("Supplementary Fig. 1", level=1)
    if SUPP_FIG_PNG.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(SUPP_FIG_PNG), width=Inches(image_width(SUPP_FIG_PNG, 6.35, 5.4)))
    caption = SUPP_CAPTION_MD.read_text(encoding="utf-8").strip() if SUPP_CAPTION_MD.exists() else ""
    caption = caption.replace("# ", "")
    caption = caption.replace(
        "Supplementary Figure 1 | Dry-lab perturbation consistency and partial causal closure of regime-conditioned dynamics.",
        "Supplementary Figure 1 | Dry-lab perturbation consistency and partial regime-conditioned perturbation consistency.",
    )
    caption = caption.replace("partial causal closure", "partial perturbation consistency")
    caption = caption.replace("causal closure score", "causal-consistency score")
    cap_p = doc.add_paragraph()
    cap_run = cap_p.add_run(caption)
    cap_run.font.size = Pt(9)
    cap_run.font.italic = True
    doc.add_page_break()
    doc.add_heading("Supplementary Tables 1-5", level=1)
    doc.add_paragraph(
        "The full supplementary tables are supplied as TSV source-data files in the submission package. "
        "Previews below show the header and first rows only where wide tables would otherwise be unreadable in Word."
    )
    for number, path, desc in SUPP_TABLES:
        doc.add_heading(f"Supplementary Table {number}. {path.name}", level=2)
        rows, cols = tsv_shape(path)
        doc.add_paragraph(f"{desc} Full source-data file: {path.name}. Rows: {rows}; columns: {cols}.")
        preview = read_tsv(path, limit_rows=6)
        add_small_table(doc, preview, max_cols=6)
    doc.add_heading("Supplementary source-data inventory", level=1)
    doc.add_paragraph("See SUPPLEMENTARY_SOURCE_DATA_INDEX.tsv for file-level placement, row counts and package paths.")
    doc.save(SUPP_DOCX)


def md5(path: Path) -> str:
    h = hashlib.md5()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def build_source_index() -> None:
    rows = []
    def add(path: Path, item: str, placement: str, role: str) -> None:
        exists = path.exists()
        n_rows, n_cols = tsv_shape(path) if exists and path.suffix.lower() == ".tsv" else (None, None)
        rows.append(
            {
                "item": item,
                "source_file": path.name,
                "placement": placement,
                "role": role,
                "rows": "" if n_rows is None else n_rows,
                "columns": "" if n_cols is None else n_cols,
                "status": "present" if exists else "missing",
                "md5": md5(path) if exists else "",
                "package_subdir": "supplementary_source_data",
            }
        )
    add(SUPP_FIG_PNG, "Supplementary Fig. 1", "Supplementary Figure", "Rendered figure PNG")
    add(SUPP_FIG_PDF, "Supplementary Fig. 1", "Supplementary Figure", "Vector/print PDF copy")
    for number, path, desc in SUPP_TABLES:
        add(path, f"Supplementary Table {number}", "Supplementary Table", desc)
    for path in EXTRA_SUPP_SOURCE:
        add(path, "Supplementary source data", "Source Data", "Additional dry-lab perturbation consistency support file")
    with SOURCE_INDEX.open("w", encoding="utf-8", newline="") as fh:
        fieldnames = ["item", "source_file", "placement", "role", "rows", "columns", "status", "md5", "package_subdir"]
        writer = csv.DictWriter(fh, fieldnames=fieldnames, delimiter="\t")
        writer.writeheader()
        writer.writerows(rows)


def extract_references() -> list[str]:
    doc = Document(BASE_DOCX)
    refs: list[str] = []
    in_refs = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "References":
            in_refs = True
            continue
        if in_refs and text == "Declarations":
            break
        if in_refs and text:
            refs.append(text)
    return refs


def build_reference_placeholders() -> None:
    refs = extract_references()
    lines = [
        "# Unresolved Reference List",
        "",
        "No existing .ris or .bib reference metadata file was found in the workspace. The manuscript contains formatted reference text, but importable reference metadata were not fabricated.",
        "",
    ]
    for ref in refs:
        lines.append(f"- {ref}")
    UNRESOLVED_REFS.write_text("\n".join(lines) + "\n", encoding="utf-8")
    ENDNOTE_TODO.write_text(
        "\n".join(
            [
                "# EndNote Import TODO",
                "",
                "Reference metadata files were not available. Before journal submission:",
                "",
                "1. Resolve each formatted manuscript reference in EndNote, Zotero or another reference manager.",
                "2. Export verified records as RIS and BibTeX.",
                "3. Re-run reference formatting according to the target journal style.",
                "4. Do not infer DOI, PMID or accession metadata from incomplete text.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def forbidden_hits(path: Path) -> list[tuple[str, int, str]]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    hits: list[tuple[str, int, str]] = []
    for phrase in FORBIDDEN_MAIN:
        for m in re.finditer(re.escape(phrase), text, flags=re.IGNORECASE):
            line_no = text[: m.start()].count("\n") + 1
            line = text.splitlines()[line_no - 1]
            hits.append((phrase, line_no, line))
    return hits


def copy_package_files() -> None:
    PKG.mkdir(exist_ok=True)
    main_fig_pkg = PKG / "main_figures"
    supp_pkg = PKG / "supplementary_source_data"
    ref_pkg = PKG / "references"
    for d in [main_fig_pkg, supp_pkg, ref_pkg]:
        d.mkdir(exist_ok=True)
    root_files = [
        FINAL_MANUSCRIPT,
        FINAL_LEGENDS,
        SUPP_DOCX,
        SOURCE_INDEX,
        FINAL_SAFE_MD,
        READINESS,
    ]
    for f in root_files:
        if f.exists():
            shutil.copy2(f, PKG / f.name)
    for f in MAIN_FIG_DIR.glob("Final_Figure_*.*"):
        if f.suffix.lower() in {".png", ".pdf", ".tif", ".tiff", ".jpg", ".jpeg", ".svg"}:
            shutil.copy2(f, main_fig_pkg / f.name)
    for f in [SUPP_FIG_PNG, SUPP_FIG_PDF] + [p for _, p, _ in SUPP_TABLES] + EXTRA_SUPP_SOURCE:
        if f.exists():
            shutil.copy2(f, supp_pkg / f.name)
    if SOURCE_INDEX.exists():
        shutil.copy2(SOURCE_INDEX, supp_pkg / SOURCE_INDEX.name)
    for f in [UNRESOLVED_REFS, ENDNOTE_TODO]:
        if f.exists():
            shutil.copy2(f, ref_pkg / f.name)
            shutil.copy2(f, PKG / f.name)


def package_missing_files() -> list[str]:
    required = [
        FINAL_MANUSCRIPT,
        FINAL_LEGENDS,
        SUPP_DOCX,
        SOURCE_INDEX,
        FINAL_SAFE_MD,
        SUPP_FIG_PNG,
        SUPP_FIG_PDF,
        *[p for _, p, _ in SUPP_TABLES],
        *[p for _, p in MAIN_FIGURES],
    ]
    return [str(p.relative_to(ROOT)) for p in required if not p.exists()]


def build_readiness_report() -> None:
    main_fig_present = [p for _, p in MAIN_FIGURES if p.exists()]
    supp_fig_count = int(SUPP_FIG_PNG.exists() or SUPP_FIG_PDF.exists())
    missing = package_missing_files()
    hits = forbidden_hits(FINAL_SAFE_MD)
    reference_metadata = list(ROOT.glob("**/*.ris")) + list(ROOT.glob("**/*.bib"))
    reference_ready = bool(reference_metadata)
    main_figures_inserted = FINAL_MANUSCRIPT.exists() and len(main_fig_present) == 9
    supp_ready = SUPP_DOCX.exists() and SOURCE_INDEX.exists() and all(p.exists() for _, p, _ in SUPP_TABLES)
    causal_safe = not hits
    ready = main_figures_inserted and supp_ready and causal_safe and reference_ready and not missing
    lines = [
        "# Final Submission Readiness Report",
        "",
        "## Package summary",
        f"- Main figure count: {len(main_fig_present)}",
        f"- Supplementary figure count: {supp_fig_count}",
        "- Locked main-figure architecture: 9 main figures; no Figure 10 added.",
        "- Supplementary causal analysis placement: Supplementary Fig. 1 and Supplementary Tables 1-5.",
        "",
        "## Causal claim safety",
        f"- Status: {'pass' if causal_safe else 'review required'}",
        "- Main text was micro-edited to avoid full causal closure or wet-lab validation claims.",
    ]
    if hits:
        lines.append("- Forbidden phrase hits:")
        for phrase, line_no, line in hits:
            lines.append(f"  - {phrase} at line {line_no}: {line}")
    else:
        lines.append("- Forbidden overclaim scan in final_claim_safe_main_text.md: no hits.")
    lines += [
        "",
        "## Reference status",
        "- Importable RIS/BibTeX metadata files found: " + ("yes" if reference_ready else "no"),
    ]
    if not reference_ready:
        lines.append("- Generated unresolved_reference_list.md and endnote_import_todo.md. DOI/PMID/citation metadata were not fabricated.")
    lines += [
        "",
        "## Missing files",
    ]
    if missing:
        lines.extend([f"- {m}" for m in missing])
    else:
        lines.append("- None among required assembly inputs.")
    lines += [
        "",
        "## Unresolved risks",
        "- Reference metadata remain unresolved unless an external reference manager export is supplied.",
        "- Journal-specific figure-resolution, file-format and source-data naming rules may require a final portal-specific check.",
        "- Supplementary Table 5 is provided as TSV source data because the full table is too large for readable Word embedding.",
        "",
        "## Verdict",
        f"MAIN_TEXT_READY: {'yes' if causal_safe else 'no'}",
        f"MAIN_FIGURES_INSERTED: {'yes' if main_figures_inserted else 'no'}",
        f"SUPPLEMENTARY_PACKAGE_READY: {'yes' if supp_ready else 'no'}",
        f"REFERENCE_PACKAGE_READY: {'yes' if reference_ready else 'no'}",
        f"CAUSAL_CLAIM_SAFE: {'yes' if causal_safe else 'no'}",
        f"READY_FOR_SUBMISSION: {'yes' if ready else 'no'}",
    ]
    READINESS.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    write_claim_safe_md()
    build_manuscript_with_figures()
    build_final_legends_docx()
    build_supplementary_material()
    build_source_index()
    build_reference_placeholders()
    build_readiness_report()
    copy_package_files()
    # Copy the readiness report once more after it is created.
    shutil.copy2(READINESS, PKG / READINESS.name)


if __name__ == "__main__":
    main()
