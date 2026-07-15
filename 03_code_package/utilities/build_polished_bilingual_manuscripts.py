from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn


ROOT = Path("/Users/hanchengdezhuanqiangongju/Documents/Codex/2026-06-18/task-reconstruct-and-continue-analysis-of")
BASE = ROOT / "Manuscript_Figure4_Figure5_panel_labels_fixed.docx"
EN_DOCX = ROOT / "Regenerative_cell_fate_latent_regime_ENGLISH_POLISHED.docx"
ZH_DOCX = ROOT / "Regenerative_cell_fate_latent_regime_CHINESE_POLISHED.docx"
EN_MD = ROOT / "Regenerative_cell_fate_latent_regime_ENGLISH_POLISHED.md"
ZH_MD = ROOT / "Regenerative_cell_fate_latent_regime_CHINESE_POLISHED.md"
LOG = ROOT / "bilingual_polish_change_log.md"


EN = [
    "Regenerative cell fate is encoded by latent-state-regime mixtures rather than a scalar state",
    "Abstract",
    "Repair and regeneration are frequently initiated by injury-induced plasticity, but mammalian wound repair and salamander blastema formation resolve into different biological outcomes. We asked whether this divergence could be reduced to a single scalar order parameter linking stemness, fate-lock, positional information and gene-regulatory structure. The scalar PGCS/Phi representation failed as a regime-separating model: ROC performance was near random, permutation testing did not support a stable mean shift, bootstrap intervals crossed zero and the Kolmogorov-Smirnov statistic captured distributional shape differences rather than separability. We therefore organized the locked outputs as a latent-state-regime mixture, P(Z|S,W_GRN), with Z restricted to adult_repair, embryonic_reactivation, salamander_blastema and salamander_intact. In this representation, stemness, fate-lock, positional programs and tumor-like plasticity contribute to posterior regime probability structure without reducing regeneration to a single coordinate. A supplementary dry-lab perturbation audit supports partial regime-conditioned perturbation consistency, but not global causal closure or wet-lab validation. These results support a model in which regenerative cell fate is represented by overlapping latent state regimes rather than by a universal scalar biological axis.",
    "Introduction",
    "Cell fate is commonly formulated as motion through a constrained landscape in which differentiation, plasticity and stability occupy distinct but partially accessible regions of state space 1,2 . Single-cell genomics now makes it possible to approximate such landscapes from molecular profiles, trajectories, RNA velocity, regulatory programs and perturbation-linked responses 3-6 . A central representational challenge remains: comparable degrees of plasticity can occur during repair, regeneration, senescence-associated remodeling and cancer-like state transitions without implying equivalent fate outcomes 7-9 .",
    "Regeneration provides a stringent setting in which to test this problem. Mammalian tissues often resolve injury through inflammatory repair, fibrosis or senescence-enriched stabilization, whereas salamander limb injury can generate a blastema capable of pattern-restoring regeneration 10-13 . Developmental positional programs, including RA-associated proximodistal identity, FGF/SHH-linked limb patterning and Notch-associated boundary or transitional-state regulation, provide candidate coordinates for regenerative organization, but they should not be interpreted as a universal scalar ranking of species or outcomes 14-17 . Likewise, stemness-associated Wnt/beta-catenin/TCF-LEF/MYC activity can increase local accessibility without determining whether a state becomes regenerative, senescent or tumor-like 9,18,19 .",
    "The earlier PGCS/Phi framework attempted to compress these components into a single scalar representation. That representation is not used here as the final biological model. In the locked validation layer, Phi retained distributional structure but failed to separate regimes: the AUC was approximately 0.480, permutation and bootstrap statistics were unstable, and the significant KS result was non-discriminative. These failures are biologically informative because they reject the view that salamander blastema is merely a higher-score version of mammalian repair, or that tumor-like plasticity is equivalent to regenerative competence 9,11 .",
    "The accepted representation for the locked outputs is a latent-state-regime mixture in which observed cell states carry posterior mass over adult_repair, embryonic_reactivation, salamander_blastema and salamander_intact regimes 20,21 . Rather than imposing a global scalar coordinate, the model summarizes local state tendencies as posterior-weighted contributions from regime-specific transition components 22 . For regime k, p_k(S)=P(Z=k|S,W_GRN), and the expected local transition tendency is summarized as E[Delta S|S]=sum_k p_k(S)F_k(S). This expression is a compact computational summary of regime-conditioned structure, not evidence that a continuous-time biological law has been experimentally identified.",
    "Results",
    "Model lineage: scalar failure requires model replacement",
    "We first make the model lineage explicit (Fig. 1A). The earlier PGCS/Phi model assumed that regenerative state could be compressed into a single order parameter, an idea consistent with classical landscape intuition but not guaranteed by cross-species regeneration biology 1,10 . The failure layer combined random-level ROC performance, permutation instability, bootstrap uncertainty and KS shape mismatch without separability. This sequence supports replacement rather than refinement: the final model is a latent-state-regime mixture, not a recalibrated scalar.",
    "The fate-control framework is organized around posterior regime mixtures",
    "Fig. 1 anchors the manuscript around posterior regime assignment rather than scalar ordering. The biological scaffold includes stemness maintenance, fate-lock/destemness, developmental positional information, species-divergent regeneration and tumor-like plasticity 14,18,23 . The representational layer changes the interpretation of these modules: they are read through P(Z|S,W_GRN), rather than as additive components of a global Phi axis.",
    "The four locked regimes are adult_repair, embryonic_reactivation, salamander_blastema and salamander_intact. This vocabulary prevents three common over-collapses: mammalian repair is not a lower-score salamander program, salamander blastema is not simply high plasticity, and tumor-like plasticity is not regeneration 10,11 .",
    "Shared accessibility is compatible with distinct posterior regimes",
    "The mammalian wound and related plasticity analyses support shared accessibility without requiring shared fate determination (Fig. 2). In the locked analysis, GSE153596 provides the mammalian wound-trajectory context and was analyzed with trajectory, RNA-velocity and CellRank-based state-transition methods 24 . High-plasticity regions can remain connected in inferred state space while separating into distinct posterior regimes 3,25,26 .",
    "This result is central to the replacement model. Accessibility is necessary for fate change, but it is not equivalent to fate determination. Under P(Z|S,W_GRN), similar accessibility can coexist with adult-repair retention, salamander-blastema posterior mass or tumor-like plasticity 27,28 .",
    "Stemness-associated programs are local accessibility components",
    "Stemness-associated programs provide a local accessibility layer rather than a global fate driver (Fig. 3). Wnt/beta-catenin/TCF-LEF/MYC activity is a well-established contributor to stemness maintenance, transcriptional competence and proliferative programs 18,19,29 . In the locked figure system, perturbation-linked summaries are consistent with local increases in Wnt output and stemness-associated state scores, but these changes do not by themselves assign cells to a regenerative regime.",
    "This distinction avoids a frequent interpretive error. A stemness increase may expand reachable state space while still resolving into adult repair, tumor-like plasticity or a non-regenerative state, depending on the surrounding posterior regime mixture. Stemness therefore enters S, but it is not a standalone determinant of Z 27 .",
    "Positional programs define regime-conditioned developmental coordinates",
    "Fig. 4 integrates positional information under a regime-conditioned interpretation. RA-associated proximodistal identity, FGF/SHH-linked limb patterning and NOTCH-associated boundary or transitional-state regulation provide candidate developmental coordinates for regenerative organization 14,17,30 . These programs are not treated as a universal species-ranking axis; instead, they define positional and patterning features whose meaning depends on posterior regime context.",
    "This distinction is particularly important for salamander regeneration. Salamander blastema formation is represented as a distinct latent regenerative regime, not as a higher value on a mammalian repair axis 10,11 . Positional-program activity is therefore consistent with blastema organization, but it does not by itself establish global regenerative competence in every tissue or species 10 .",
    "Fate-lock defines an adult-repair basin constraint",
    "Fig. 5 frames fate-lock as an adult-repair basin constraint. BMP-SMAD signaling, p53/p21, p16/Rb and senescence-associated programs are mechanistically linked to differentiation bias, cell-cycle arrest, tumor suppression and damage-associated tissue remodeling 7,31-33 . In the model, these processes bias mammalian repair toward adult-repair-dominant posterior mixtures.",
    "The basin interpretation remains deliberately restrained. Fate-lock is not claimed to be a universal aging law, and the current analysis does not close the causal path from individual pathway perturbations to irreversible cell fate 34 . Instead, it provides a computationally organized constraint: adult repair can retain plasticity while remaining biased toward stabilized, non-blastema outcomes.",
    "Tumor-like plasticity is a separate high-plasticity branch",
    "Fig. 6 preserves tumor-like plasticity as a boundary branch rather than a regenerative regime. Cancer and tumor-like cell states can show stemness, self-renewal, lineage ambiguity and local retention, but these properties do not by themselves make them regenerative 11,27,35 . This separation is essential because a model that equates plasticity with regeneration would incorrectly merge tumor-like and blastema-like states.",
    "In the final framework, tumor-like plasticity is treated as a negative boundary condition outside the locked four-regime posterior vocabulary, even when it shares components of accessibility or stemness. High plasticity without compatible positional and regime context should therefore not be interpreted as regenerative competence 36 .",
    "The deprecated scalar embedding fails as a regime separator",
    "The scalar validation layer is shown in Fig. 7. In the locked outputs, Phi did not behave as a reliable classifier or biological order parameter: ROC performance was near random (AUC approximately 0.480), the permutation mean-shift test was not significant, bootstrap intervals overlapped zero and the KS result reflected distributional shape differences rather than separability. These results make Phi a deprecated scalar proxy rather than a partial predictor 47 .",
    "This failure is the turning point of the manuscript. Because salamander blastema, mammalian repair and intact salamander tissue were not separated by a stable scalar representation, the downstream framework shifts to a representation that preserves overlap, posterior uncertainty and mixture structure.",
    "Latent-state-regime posterior dynamics provide the locked representation",
    "Fig. 8 presents the accepted dynamical representation. Regime identity is encoded by P(Z|S,W_GRN), where each observed state carries posterior mass over regime-specific components rather than being assigned to a single scalar coordinate 6,22 . For regime k, p_k(S)=P(Z=k|S,W_GRN), and the expected local transition tendency is summarized as E[Delta S|S]=sum_k p_k(S)F_k(S). This notation describes regime-conditioned computational structure, not an experimentally identified continuous-time biological law 5 .",
    "The model is presented as an evaluated representation rather than a complete causal regulatory circuit. It does not require separable clusters, scalar thresholds or manually specified switching variables, and it does not establish wet-lab causal closure 34 .",
    "Overlap and divergence jointly define the final regime structure",
    "Fig. 9 closes the representation problem by separating overlap from identity. Overlap matrices show that observed species/regime groups and inferred latent state regimes are not cleanly separable in measured state space, consistent with the failure of scalar compression in this analysis. Symmetrized KL divergence, by contrast, captures distributional differences among regimes; here, divergence is interpreted as distributional structure rather than distance along a biological axis 37 .",
    "Together, overlap and divergence support a non-compressible representation. The locked outputs are structured enough to support posterior regime reconstruction under the evaluated model, but not separable enough to support a single biological coordinate.",
    "Supplementary dry-lab perturbation consistency is partial and regime-conditioned",
    "The supplementary perturbation audit asked whether pathway-linked posterior-shift summaries were directionally compatible with the learned regulatory structure using existing dry-lab outputs (Supplementary Fig. 1; Supplementary Tables 1-5). Perturb-seq and related single-cell perturbation approaches provide a useful basis for computational consistency checks, but such analyses do not by themselves establish wet-lab validation in the present manuscript 34,38,39 .",
    "The global W_GRN model was insufficient, whereas regime-conditioned W(Z) reduced mean squared error from 0.00835 to 0.00312 in the locked comparison table. The improvement was partial: RA and BMP showed the strongest available perturbation-consistency support, whereas NOTCH, FGF and SHH were limited by weaker coverage or inconsistent directionality in the available dry-lab evidence. Counterfactual direction reversals remained frequent. The supported conclusion is partial regime-conditioned perturbation consistency, not global causal closure.",
    "Discussion",
    "The central conclusion is representational. Regenerative cell fate was not well described by the deprecated PGCS/Phi scalar because scalar compression lost the overlap, uncertainty and cross-species divergence observed in the locked outputs. This result remains compatible with landscape views of cell-fate organization, but it shifts the operative representation from a single coordinate to posterior regime probability structure 1,2,22 . Plasticity is permissive, not instructive.",
    "The latent-state-regime mixture clarifies how major biological modules should be interpreted. Stemness-associated programs expand local accessibility; fate-lock constrains adult repair; positional programs provide regime-conditioned developmental coordinates; salamander blastema is treated as a distinct latent regenerative regime; and tumor-like plasticity remains a separate high-plasticity boundary branch 11,18,23,27,36 . The value of the model is that these components can overlap without being forced into a scalar hierarchy. Regenerative competence therefore requires accessibility to be interpreted together with positional information, fate-lock state and posterior regime context.",
    "The most important negative result is the failure of scalar separability. The significant KS statistic does not rescue Phi because it reflects distributional shape difference rather than a biologically usable threshold. Near-random ROC performance and unstable resampling results likewise make Phi unsuitable for regime assignment. These failures support a replacement model, not a softened version of the old scalar framework.",
    "The supplementary perturbation layer should be interpreted cautiously. Regime-conditioned W(Z) improved posterior-shift reconstruction relative to a global W_GRN, but the analysis remains dry-lab, partial and limited by counterfactual reversals and cross-regime instability. Thus, the manuscript supports partial regime-conditioned perturbation consistency, not complete causal explanation 34 .",
    "The framework points to testable experimental directions rather than completed validation. The central experimental challenge is not to maximize plasticity, but to organize access to regime-specific state space. Rather than attempting to maximize a scalar regenerative score, perturbation studies should test whether pathway modulation changes posterior regime access in a regime-conditioned manner 38 . Direct lineage tracing, spatial morphogen measurements, matched perturbation time courses and multiome measurements will be required to determine whether the posterior regimes correspond to causal biological transitions in vivo 36 .",
    "Limitations",
    "This study is a computational representation study, not a complete causal theory of regeneration. The final model is constrained by processed state scores, proxy modules, cross-species alignment, orthology mapping and batch/species entanglement 40,41 . Dataset labels, treatment annotations and perturbation metadata are therefore used only within the limits supported by the source records and are not treated as independent wet-lab validation 40,42 .",
    "The regulatory-network layer is treated as a learned and interpretable computational object. The manuscript does not claim that every inferred edge has direct experimental support or that the inferred W_GRN fully reconstructs the underlying biological regulatory circuit 5,6 . Chromatin memory, lineage history and spatial morphogen gradients remain relevant missing layers rather than completed measurements 43-46 .",
    "The dry-lab perturbation analysis supports partial regime-conditioned consistency only. It should not be read as evidence that pathway changes are sufficient to cause the observed regime transitions in vivo. Stronger causal validation would require matched wet-lab perturbation time series, lineage tracing and multiome or spatial readouts in the same biological contexts 38,44,45 .",
    "Methods",
    "Project-state reconstruction",
    "The manuscript was reconstructed from locked project-state files, claim-safe main text, figure-to-story mapping, inconsistency reports and submission-readiness reports. Reconstruction preserved the accepted latent-state-regime mixture model and treated scalar-positive Phi artifacts as historical or failure-context outputs.",
    "State representation",
    "Cell-state structure was represented through processed state variables and module scores summarizing stemness/accessibility, transitional behavior, fate-lock and developmental positional information. These state components were interpreted as inputs to posterior regime assignment rather than as a scalar developmental axis 11,22 .",
    "Scalar model invalidation",
    "The deprecated Phi model was evaluated using ROC performance, permutation testing, bootstrap confidence intervals and distributional comparison. The locked outputs included random-level AUC, unstable mean-shift statistics and KS-detected shape mismatch without separability 47 .",
    "Latent-state-regime mixture model",
    "The accepted representation used P(Z|S,W_GRN), with Z restricted to adult_repair, embryonic_reactivation, salamander_blastema and salamander_intact. For regime k, p_k(S)=P(Z=k|S,W_GRN), and the expected local transition tendency is summarized as E[Delta S|S]=sum_k p_k(S)F_k(S), where F_k(S) denotes a regime-conditioned transition component. This notation is a compact modeling summary of regime-conditioned structure, not evidence that a continuous-time biological law has been experimentally identified 6,22 .",
    "Trajectory and posterior analyses",
    "Trajectory, velocity and terminal-state summaries were interpreted as computational evidence for accessibility, retention and posterior regime probability structure, not as direct lineage-history evidence 3,26,48 . Dataset records, figure source files and locked project outputs were used only within their documented scope and were not treated as newly validated external evidence 49,50 .",
    "Regulatory-network and perturbation-consistency analyses",
    "The learned regulatory-network object W_GRN was used as part of the posterior-regime representation and perturbation-consistency audit 51 . Regime-conditioned W(Z) was compared with a global W_GRN model using reconstruction error, directionality, counterfactual agreement and null controls. This analysis was used only as dry-lab perturbation-consistency support, not as complete causal validation 5,6,34,38 .",
    "Figure reconstruction and claim safety",
    "The final figure system was held to nine main figures plus Supplementary Fig. 1. Fig. 1A was used as the model-evolution map and did not create a new main-result figure number. Claims were audited to remove scalar-positive Phi interpretation, wet-lab validation language and complete causal-closure language.",
    "Data availability",
    "All public datasets used in this study will be listed with accession numbers in the final submission. Derived source-data tables supporting the figures will be provided as supplementary source-data files with the manuscript.",
    "Code availability",
    "Custom analysis scripts and figure-generation scripts will be provided as Supplementary Software with the final submission. The code package will include scripts required to reproduce the processed score tables, model comparisons, figure source-data tables and manuscript figures from the processed input files.",
    "Acknowledgements",
    "Acknowledgements will be finalized before submission.",
    "Author contributions",
    "Author contributions will be finalized before submission.",
    "Competing interests",
    "The authors declare no competing interests.",
    "Figure legends",
    "Figure 1A | Model evolution map: replacement, not refinement.",
    "The earlier PGCS/Phi scalar model treated cell fate as if it could be approximated by a single order parameter, an assumption related to classical landscape views of cell-fate organization but not guaranteed by cross-species single-cell comparisons 1,52 . Colored boxes mark successive model states: old scalar assumption, failure layer, transition, accepted mixture model and final claim. Arrows indicate logical replacement rather than refinement or temporal progression. The locked validation layer rejected Phi as a regime separator: discrimination was near random, permutation and bootstrap summaries were unstable, and the KS result reflected distributional shape difference rather than separability 47 . The accepted framework is a latent-state-regime mixture, P(Z|S,W_GRN), in which posterior regime probability structure replaces a global scalar axis 6,22 .",
    "Figure 1 | Phi failure and latent-state-regime replacement logic.",
    "Figure 1 introduces the transition from scalar ordering to posterior regime probability structure. Phi is shown only as a failed scalar projection. Colors separate the rejected scalar hypothesis, the accepted posterior-mixture representation, the dynamical replacement and non-compressible outcome structure. Arrows show conceptual relationships among mixture components and do not imply a one-dimensional ordering. The accepted interpretation is that observed states carry posterior mass over adult_repair, embryonic_reactivation, salamander_blastema and salamander_intact regimes.",
    "Figure 2 | Shared accessibility does not determine fate.",
    "The figure summarizes connected accessibility and high-plasticity structure while separating accessibility from fate determination. Categorical colors denote inferred state groups, whereas continuous colorbars report the indicated score or probability. Heatmaps encode normalized transition or absorption values, with darker/larger color intensity indicating higher values on the displayed scale. Boxplots show distributional summaries with points representing cells or state observations. Trajectory, pseudotime and velocity-derived summaries are interpreted as computational state-space evidence, not as direct lineage validation 3,26,48 .",
    "Figure 3 | Stemness is a local accessibility component.",
    "Stemness-associated Wnt/beta-catenin/TCF-LEF/MYC programs contribute to local accessibility and proliferative competence, but they do not define global regime identity 18,19,29 . Point and line colors denote perturbation or pathway conditions as labeled in each panel. Higher module-score or color intensity indicates stronger activity of the displayed score. Diverging heatmap colors indicate relative module-score direction around the scale midpoint.",
    "Figure 4 | Positional programs are regime-conditioned developmental coordinates.",
    "Panels a-h summarize positional-program evidence: a GSE295225 positional score, b RA/RARG effect sizes, c module means, d axolotl regeneration-stage embedding, e positional score, f blastema-associated regenerative score, g stage-wise module dynamics and h pseudotime dynamics. Panels i-p summarize salamander regenerative-regime evidence: i regeneration-stage embedding, j blastema progenitor score, k blastema-associated regenerative score, l RA/HOX positional identity, m stage dynamics, n basin-like proxy distributions, o pseudotime trends and p branch comparison. RA/HOX-, FGF/SHH- and NOTCH-associated positional or patterning programs are interpreted within posterior regime context, not as a universal scalar ranking of mammalian and salamander systems 14,17,30 .",
    "Figure 5 | Fate-lock constrains the adult-repair basin.",
    "Panels a-o summarize fate-lock as an adult-repair basin constraint: a basin constraint schematic; b-e fate-stabilization logic, score, p53-BMP-associated score and senescence-like absorption; f-g velocity-derived transitions and velocity self-retention; h-j SAT feedback branch, senescence-like retention anchor and fate-stabilization signal; k-l perturbation-linked validation and stabilization-axis summary; and m-o aging/SAT signal, positional score and Result 5 summary. The basin schematic is conceptual, whereas scatter plots, heatmaps, bar charts and boxplots summarize locked source-data outputs. The figure does not claim complete causal closure.",
    "Figure 6 | Tumor-like plasticity is a separate branch.",
    "Tumor-like plasticity is shown as a high-plasticity boundary branch distinct from regenerative competence 27,36 . Continuous colorbars encode the indicated tumor-like, stemness, differentiation or local-retention score, with higher intensity corresponding to higher score values. Boxplots compare score distributions between groups; heatmap colors show module means. The tumor-like branch is not interpreted as salamander_blastema or regenerative competence.",
    "Figure 7 | Scalar Phi is rejected as a regime separator.",
    "ROC, permutation, bootstrap and KS summaries show that Phi is non-discriminative as a global order parameter 47 . Line colors distinguish empirical curves, null or random-reference expectations and observed statistics as labeled. The ROC diagonal indicates random-level performance; permutation curves show the label-shuffle null; bootstrap intervals show uncertainty around estimated shifts. The significant KS statistic is interpreted as distributional shape difference, not separability.",
    "Figure 8 | Latent-state-regime posterior dynamics provide the working representation.",
    "The model represents each observed state through posterior regime assignment, P(Z|S,W_GRN), rather than through a global scalar coordinate. Colored density curves and stacked bars indicate posterior mass assigned to adult_repair, embryonic_reactivation, salamander_blastema and salamander_intact. Higher density or bar height reflects greater posterior mass within the displayed summary. For regime k, p_k(S)=P(Z=k|S,W_GRN), and the expected local transition tendency is summarized as E[Delta S|S]=sum_k p_k(S)F_k(S). This notation describes a regime-conditioned computational representation, not an experimentally identified continuous-time biological law 6,22 .",
    "Figure 9 | Overlap and symmetrized divergence define regime structure.",
    "Overlap is interpreted as non-separability, whereas symmetrized KL divergence is interpreted as distributional difference rather than scalar distance 22,37 . Heatmap color intensity corresponds to the displayed matrix value, with numbers printed in cells. Blue overlap matrices quantify shared posterior or state-space occupancy; red bar lengths rank the strongest symmetrized KL divergence values. Higher divergence indicates greater distributional difference, not movement along a biological axis.",
    "Supplementary Figure 1 | Dry-lab perturbation consistency audit.",
    "Existing perturbation-derived outputs were used to ask whether pathway-linked posterior-shift summaries were directionally compatible with regime-conditioned structure. Heatmap colors encode metric values, with higher intensity indicating stronger directional consistency, null separation, support or bias as labeled; gray cells denote a shared model-level value rather than a pathway-specific estimate. The model-comparison panel reports global W_GRN versus regime-conditioned W(Z), and the consistency-limit panel summarizes unavailable cross-dataset comparison, frequent counterfactual reversals, cross-regime instability and limited NOTCH/SHH evidence. Regime-conditioned W(Z) reduced reconstruction error relative to a global W_GRN, but counterfactual reversals and cross-regime instability remained; the analysis supports partial regime-conditioned perturbation consistency only, not complete causal closure or wet-lab validation.",
]


ZH = [
    "再生性细胞命运由潜在调控状态混合表征，而非单一标量状态",
    "摘要",
    "修复与再生通常起始于损伤诱导的可塑性，但哺乳动物创伤修复与蝾螈芽基形成并不会收敛到同一生物学结局。本研究检验了这一差异是否能够被压缩为一个单一标量序参量，并同时连接 stemness、fate-lock、位置信息和基因调控网络结构。标量 PGCS/Phi 表征未能作为区分调控状态的模型：ROC 表现接近随机水平，置换检验不支持稳定的均值偏移，bootstrap 置信区间跨越零值，Kolmogorov-Smirnov 统计量反映的是分布形状差异而非可分性。因此，我们将锁定输出组织为潜在调控状态混合模型 P(Z|S,W_GRN)，其中 Z 限定为 adult_repair、embryonic_reactivation、salamander_blastema 和 salamander_intact。在这一表征中，stemness、fate-lock、位置程序和肿瘤样可塑性共同塑造后验调控状态结构，但再生并不被还原为单一坐标。补充性干实验扰动审计支持部分 regime-conditioned perturbation consistency，但不支持全局因果闭合或湿实验验证。整体而言，这些结果支持一种模型：再生性细胞命运由相互重叠的潜在调控状态表示，而不是由普适的标量生物学轴表示。",
    "引言",
    "细胞命运常被表述为在受约束景观中的运动，其中分化、可塑性和稳定性占据状态空间中不同但部分可达的区域 1,2 。单细胞组学使我们能够从分子图谱、轨迹、RNA velocity、调控程序和扰动相关响应中近似重建此类景观 3-6 。然而，一个核心表征问题仍然存在：相近程度的可塑性可以出现在组织修复、再生、衰老相关重塑和癌样状态转换中，但这并不意味着它们具有等价的命运结局 7-9 。",
    "再生为检验这一问题提供了严格情境。哺乳动物组织在损伤后常通过炎症性修复、纤维化或富集衰老程序的稳定化过程结束，而蝾螈肢体损伤则可形成能够支持模式恢复的芽基 10-13 。发育位置信息程序，包括 RA 相关近远轴身份、FGF/SHH 相关肢体模式化以及 Notch 相关边界或过渡状态调控，为再生组织化提供候选坐标，但不应被解释为跨物种或跨结局的普适标量排序 14-17 。同样，Wnt/beta-catenin/TCF-LEF/MYC 相关 stemness 活性能够增加局部可达性，却不能单独决定一个状态将进入再生、衰老还是肿瘤样结局 9,18,19 。",
    "早期 PGCS/Phi 框架试图将这些成分压缩为单一标量表征。本文不将该表征作为最终生物学模型。在锁定验证层中，Phi 保留了一定的分布结构，却无法区分调控状态：AUC 约为 0.480，置换和 bootstrap 统计均不稳定，显著的 KS 结果也不具备判别性。这一失败具有生物学意义，因为它否定了将蝾螈芽基视为哺乳动物修复高分版本的解释，也否定了将肿瘤样可塑性等同于再生能力的解释 9,11 。",
    "锁定输出所接受的表征是潜在调控状态混合模型，其中观测到的细胞状态在 adult_repair、embryonic_reactivation、salamander_blastema 和 salamander_intact 四种状态上携带后验质量 20,21 。模型并不施加全局标量坐标，而是将局部状态倾向概括为来自特定调控状态转移成分的后验加权贡献 22 。对于状态 k，p_k(S)=P(Z=k|S,W_GRN)，期望局部转移倾向写作 E[Delta S|S]=sum_k p_k(S)F_k(S)。该表达式是对 regime-conditioned structure 的紧凑计算概括，并不意味着连续时间生物学定律已经被实验证实。",
    "结果",
    "模型谱系：标量失败要求模型替换",
    "我们首先明确模型谱系（Fig. 1A）。早期 PGCS/Phi 模型假定再生状态可以被压缩为单一序参量；这一思路与经典景观直觉相符，但并不能由跨物种再生生物学自动保证 1,10 。失败层包括接近随机的 ROC 表现、置换检验不稳定、bootstrap 不确定性以及缺乏可分性的 KS 分布形状差异。该序列支持模型替换，而不是对旧标量进行重新校准：最终模型是潜在调控状态混合模型，而非再标定的标量模型。",
    "命运控制框架围绕后验调控状态混合组织",
    "Fig. 1 将全文锚定在后验调控状态分配，而不是标量排序。生物学框架包括 stemness maintenance、fate-lock/destemness、发育位置信息、物种差异性再生以及肿瘤样可塑性 14,18,23 。真正改变的是表征层：这些模块通过 P(Z|S,W_GRN) 被解释，而不是作为全局 Phi 轴的加性成分。",
    "四种锁定状态为 adult_repair、embryonic_reactivation、salamander_blastema 和 salamander_intact。该词汇体系避免了三类常见混淆：哺乳动物修复不是低分版蝾螈程序，蝾螈芽基不只是高可塑性，肿瘤样可塑性也不是再生 10,11 。",
    "共享可达性可以与不同后验调控状态共存",
    "哺乳动物创伤及相关可塑性分析支持共享可达性，但并不要求共享命运决定（Fig. 2）。在锁定分析中，GSE153596 提供哺乳动物创伤轨迹背景，并通过轨迹、RNA velocity 和 CellRank 状态转移方法进行分析 24 。高可塑性区域可以在推断的状态空间中保持连通，同时分离到不同后验调控状态 3,25,26 。",
    "这一结果是模型替换的核心。可达性是命运改变的必要条件，但并不等同于命运决定。在 P(Z|S,W_GRN) 下，相似的可达性可以与 adult-repair 保留、salamander-blastema 后验质量或肿瘤样可塑性共存 27,28 。",
    "Stemness 相关程序是局部可达性成分",
    "Stemness 相关程序提供局部可达性层，而不是全局命运驱动因子（Fig. 3）。Wnt/beta-catenin/TCF-LEF/MYC 活性是 stemness 维持、转录能力和增殖程序的重要组成部分 18,19,29 。在锁定图件系统中，扰动相关摘要与 Wnt 输出和 stemness 相关状态评分的局部升高一致，但这些变化本身并不能把细胞分配到再生状态。",
    "这一点避免了常见解释错误。Stemness 升高可能扩展可达状态空间，但其最终仍可根据周围后验调控状态混合而进入 adult repair、tumor-like plasticity 或非再生状态。因此，stemness 进入 S，但不是 Z 的独立决定因子 27 。",
    "位置程序定义 regime-conditioned 的发育坐标",
    "Fig. 4 在 regime-conditioned 解释下整合位置信息。RA 相关近远轴身份、FGF/SHH 相关肢体模式化以及 NOTCH 相关边界或过渡状态调控，为再生组织化提供候选发育坐标 14,17,30 。这些程序不被处理为普适的物种排序轴；相反，它们定义了位置信息和模式化特征，其含义依赖后验调控状态背景。",
    "这一点对于蝾螈再生尤其重要。蝾螈芽基形成被表示为一种独立的潜在再生状态，而不是哺乳动物修复轴上的更高取值 10,11 。因此，位置程序活性与芽基组织化相一致，但它本身并不能在所有组织或物种中建立全局再生能力 10 。",
    "Fate-lock 定义 adult-repair 盆地约束",
    "Fig. 5 将 fate-lock 表述为 adult-repair 盆地约束。BMP-SMAD 信号、p53/p21、p16/Rb 和衰老相关程序在机制上与分化偏向、细胞周期停滞、肿瘤抑制以及损伤相关组织重塑相联系 7,31-33 。在模型中，这些过程将哺乳动物修复偏置到 adult-repair 主导的后验混合。",
    "盆地解释保持有意的保守性。本文并不声称 fate-lock 是普适衰老定律，也不声称当前分析已闭合从单个通路扰动到不可逆细胞命运的因果路径 34 。相反，它提供了一个计算组织化的约束：adult repair 可以保留可塑性，同时仍偏向稳定化、非芽基样结局。",
    "肿瘤样可塑性是独立的高可塑性分支",
    "Fig. 6 将肿瘤样可塑性保留为边界分支，而不是再生状态。癌症和肿瘤样细胞状态可以显示 stemness、自我更新、谱系模糊性和局部保留，但这些性质本身并不能使其具有再生性 11,27,35 。这一分离非常关键，因为将可塑性等同于再生的模型会错误地合并肿瘤样状态和芽基样状态。",
    "在最终框架中，肿瘤样可塑性被视为锁定四状态后验词汇之外的负向边界条件，即使它共享可达性或 stemness 的某些成分也是如此。缺乏相容位置背景和调控状态背景的高可塑性，不应被解释为再生能力 36 。",
    "已弃用的标量嵌入不能区分调控状态",
    "标量验证层见 Fig. 7。在锁定输出中，Phi 并不表现为可靠分类器或生物学序参量：ROC 表现接近随机（AUC 约 0.480），置换均值偏移检验不显著，bootstrap 区间跨越零值，KS 结果反映的是分布形状差异而非可分性。因此，Phi 是已弃用的标量代理，而不是部分预测因子 47 。",
    "这一失败构成全文的转折点。由于蝾螈芽基、哺乳动物修复和完整蝾螈组织不能被稳定标量表征分离，后续框架转向能够保留重叠、后验不确定性和混合结构的表征。",
    "潜在调控状态后验动力学提供锁定表征",
    "Fig. 8 展示被接受的动力学表征。调控状态身份由 P(Z|S,W_GRN) 编码，其中每个观测状态在特定调控状态成分上携带后验质量，而不是被分配到单一标量坐标 6,22 。对于状态 k，p_k(S)=P(Z=k|S,W_GRN)，期望局部转移倾向写作 E[Delta S|S]=sum_k p_k(S)F_k(S)。该记号描述 regime-conditioned 的计算结构，而不是已经通过实验识别出的连续时间生物学定律 5 。",
    "该模型被呈现为经过评估的表征，而不是完整的因果调控回路。它不需要可分离聚类、标量阈值或人工指定的切换变量，也不建立湿实验层面的因果闭合 34 。",
    "重叠与发散共同定义最终调控状态结构",
    "Fig. 9 通过区分重叠与身份来闭合表征问题。重叠矩阵显示，观测物种/状态群体和推断的潜在调控状态在测量状态空间中并不能被干净分离，这与本分析中标量压缩的失败一致。相反，对称 KL 发散捕捉调控状态之间的分布差异；此处，发散被解释为分布结构，而不是生物学轴上的距离 37 。",
    "总体而言，重叠与发散支持一种不可压缩表征。锁定输出具有足够结构以支持被评估模型下的后验调控状态重建，但并不具备足够可分性以支持单一生物学坐标。",
    "补充性干实验扰动一致性是部分且调控状态条件化的",
    "补充扰动审计检验了通路相关后验偏移摘要是否与既有干实验输出中的学习调控结构方向一致（Supplementary Fig. 1；Supplementary Tables 1-5）。Perturb-seq 及相关单细胞扰动方法为计算一致性检查提供了有用基础，但在本文中，这类分析本身并不构成湿实验验证 34,38,39 。",
    "全局 W_GRN 模型不足，而 regime-conditioned W(Z) 在锁定比较表中将均方误差从 0.00835 降至 0.00312。该改进是部分的：RA 和 BMP 显示最强的可用扰动一致性支持，而 NOTCH、FGF 和 SHH 受限于覆盖不足或可用干实验证据中的方向不一致。反事实方向反转仍然频繁。可支持的结论是部分 regime-conditioned perturbation consistency，而不是全局因果闭合。",
    "讨论",
    "本文的核心结论是表征性的。再生性细胞命运不能被已弃用的 PGCS/Phi 标量充分描述，因为标量压缩丢失了锁定输出中观察到的重叠、不确定性和跨物种发散。该结果仍与细胞命运组织的景观观点相容，但操作性表征从单一坐标转向后验调控状态结构 1,2,22 。可塑性是许可性条件，而不是指令性条件。",
    "潜在调控状态混合模型澄清了主要生物学模块的解释方式。Stemness 相关程序扩展局部可达性；fate-lock 约束 adult repair；位置程序提供 regime-conditioned 的发育坐标；蝾螈芽基被视为独立潜在再生状态；肿瘤样可塑性则保持为独立的高可塑性边界分支 11,18,23,27,36 。该模型的价值在于允许这些成分重叠，而不强迫它们进入标量层级。因此，再生能力需要将可达性与位置信息、fate-lock 状态和后验调控状态背景共同解释。",
    "最重要的负结果是标量可分性的失败。显著的 KS 统计量无法挽救 Phi，因为它反映的是分布形状差异，而不是生物学上可用的阈值。接近随机的 ROC 表现和不稳定的重采样结果同样使 Phi 不适合作为状态分配工具。这些失败支持模型替换，而不是旧标量框架的弱化版本。",
    "补充扰动层需要谨慎解释。Regime-conditioned W(Z) 相对于全局 W_GRN 改善了后验偏移重构，但该分析仍是干实验、部分的，并受反事实反转和跨状态不稳定性的限制。因此，本文支持部分 regime-conditioned perturbation consistency，而不是完整因果解释 34 。",
    "该框架指出了未来可检验的实验方向，而不是声称已经完成验证。核心实验挑战不是最大化可塑性，而是组织对特定调控状态空间的访问。与其试图最大化一个标量再生评分，扰动研究应测试通路调节是否以 regime-conditioned 方式改变后验调控状态访问 38 。直接谱系追踪、空间形态发生因子测量、匹配的扰动时间序列以及 multiome 测量，将是判断这些后验调控状态是否对应体内因果生物学转变所必需的 36 。",
    "局限性",
    "本研究是计算表征研究，而不是完整的再生因果理论。最终模型受处理后的状态评分、代理模块、跨物种对齐、直系同源映射以及批次/物种纠缠所限制 40,41 。因此，数据集标签、处理注释和扰动元数据仅在源记录支持的范围内使用，并不被视为独立湿实验验证 40,42 。",
    "调控网络层被视为已学习且可解释的计算对象。本文并不声称每一条推断边都有直接实验证据，也不声称推断的 W_GRN 完整重建了底层生物学调控回路 5,6 。染色质记忆、谱系历史和空间形态发生因子梯度仍是相关但尚未完成测量的层级 43-46 。",
    "干实验扰动分析仅支持部分 regime-conditioned consistency。它不应被解读为通路改变足以在体内导致观测到的调控状态转换。更强的因果验证需要在相同生物学背景中进行匹配的湿实验扰动时间序列、谱系追踪以及 multiome 或空间读出 38,44,45 。",
    "方法",
    "项目状态重建",
    "手稿由锁定项目状态文件、claim-safe 主文、图-叙事映射、不一致性报告和投稿准备度报告重建而来。重建过程保留了被接受的潜在调控状态混合模型，并将所有 scalar-positive Phi 元素仅作为历史或失败背景输出处理。",
    "状态表征",
    "细胞状态结构通过处理后的状态变量和模块评分表示，这些变量概括 stemness/accessibility、过渡行为、fate-lock 和发育位置信息。这些状态成分被解释为后验调控状态分配的输入，而不是标量发育轴 11,22 。",
    "标量模型失效判定",
    "已弃用的 Phi 模型通过 ROC 表现、置换检验、bootstrap 置信区间和分布比较进行评估。锁定输出包括随机水平的 AUC、不稳定的均值偏移统计，以及 KS 检测到的无可分性分布形状不匹配 47 。",
    "潜在调控状态混合模型",
    "被接受的表征使用 P(Z|S,W_GRN)，其中 Z 限定为 adult_repair、embryonic_reactivation、salamander_blastema 和 salamander_intact。对于状态 k，p_k(S)=P(Z=k|S,W_GRN)，期望局部转移倾向写作 E[Delta S|S]=sum_k p_k(S)F_k(S)，其中 F_k(S) 表示 regime-conditioned transition component。该记号是对 regime-conditioned structure 的紧凑建模概括，而不是连续时间生物学定律已经被实验证实的证据 6,22 。",
    "轨迹与后验分析",
    "轨迹、velocity 和终末状态摘要被解释为可达性、保留和后验调控状态结构的计算证据，而不是直接谱系历史证据 3,26,48 。数据集记录、图件源文件和锁定项目输出仅在其记录范围内使用，并不被视为新的外部验证证据 49,50 。",
    "调控网络和扰动一致性分析",
    "学习得到的调控网络对象 W_GRN 被用作后验调控状态表征和扰动一致性审计的一部分 51 。Regime-conditioned W(Z) 与全局 W_GRN 模型在重构误差、方向性、反事实一致性和空模型控制方面进行比较。该分析仅作为干实验扰动一致性支持，而不作为完整因果验证 5,6,34,38 。",
    "图件重建与声明安全",
    "最终图件系统保持为九个主图加 Supplementary Fig. 1。Fig. 1A 作为模型演化图使用，并不创建新的主结果图号。所有声明均经过审计，以移除正向标量 Phi 解释、湿实验验证语言和完整因果闭合语言。",
    "数据可用性",
    "本研究使用的所有公共数据集将在最终投稿中列出 accession numbers。支持图件的派生源数据表将作为补充源数据文件随稿件提供。",
    "代码可用性",
    "自定义分析脚本和图件生成脚本将在最终投稿中作为 Supplementary Software 提供。代码包将包括从处理后输入文件重现处理评分表、模型比较、图件源数据表和稿件图件所需的脚本。",
    "致谢",
    "致谢内容将在投稿前最终确定。",
    "作者贡献",
    "作者贡献将在投稿前最终确定。",
    "利益冲突",
    "作者声明不存在利益冲突。",
    "图注",
    "Figure 1A | 模型演化图：替换，而非修补。",
    "早期 PGCS/Phi 标量模型将细胞命运近似为单一序参量；这一假设与经典细胞命运景观观点有关，但并不能由跨物种单细胞比较自动保证 1,52 。彩色框标记连续模型状态：旧标量假设、失败层、过渡、被接受的混合模型和最终声明。箭头表示逻辑替换，而不是渐进式修补或时间先后。锁定验证层拒绝了 Phi 作为状态分离器的解释：判别接近随机，置换和 bootstrap 摘要不稳定，KS 结果反映分布形状差异而非可分性 47 。被接受的框架是潜在调控状态混合 P(Z|S,W_GRN)，其中后验调控状态结构取代全局标量轴 6,22 。",
    "Figure 1 | Phi 失败与潜在调控状态替换逻辑。",
    "Figure 1 展示从标量排序到后验调控状态结构的转变。Phi 仅作为失败的标量投影出现。颜色区分被拒绝的标量假设、被接受的后验混合表征、动力学替换和不可压缩结局结构。箭头显示混合成分之间的概念关系，并不暗示一维排序。被接受的解释是：观测状态在 adult_repair、embryonic_reactivation、salamander_blastema 和 salamander_intact 状态上携带后验质量。",
    "Figure 2 | 共享可达性不决定命运。",
    "该图概括了连通可达性和高可塑性结构，同时将可达性与命运决定分开。分类颜色表示推断状态群体，连续色条报告相应评分或概率。热图编码归一化转移或吸收值，颜色越深或强度越大表示显示尺度上的数值越高。箱线图显示分布摘要，点表示细胞或状态观测。轨迹、pseudotime 和 velocity 摘要被解释为计算状态空间证据，而不是直接谱系验证 3,26,48 。",
    "Figure 3 | Stemness 是局部可达性成分。",
    "Stemness 相关 Wnt/beta-catenin/TCF-LEF/MYC 程序促进局部可达性和增殖能力，但并不定义全局调控状态身份 18,19,29 。点和线颜色表示各面板标注的扰动或通路条件。模块评分或颜色强度越高，表示所显示评分活性越强。发散热图颜色表示围绕尺度中点的相对模块评分方向。",
    "Figure 4 | 位置程序是 regime-conditioned 的发育坐标。",
    "面板 a-h 概括位置程序证据：a GSE295225 positional score，b RA/RARG effect sizes，c module means，d axolotl regeneration-stage embedding，e positional score，f blastema-associated regenerative score，g stage-wise module dynamics，h pseudotime dynamics。面板 i-p 概括蝾螈再生状态证据：i regeneration-stage embedding，j blastema progenitor score，k blastema-associated regenerative score，l RA/HOX positional identity，m stage dynamics，n basin-like proxy distributions，o pseudotime trends，p branch comparison。RA/HOX、FGF/SHH 和 NOTCH 相关位置或模式化程序在后验调控状态背景中解释，而不是作为哺乳动物和蝾螈系统的普适标量排序 14,17,30 。",
    "Figure 5 | Fate-lock 约束 adult-repair 盆地。",
    "面板 a-o 将 fate-lock 概括为 adult-repair 盆地约束：a 盆地约束示意图；b-e fate-stabilization logic、score、p53-BMP-associated score 和 senescence-like absorption；f-g velocity-derived transitions 和 velocity self-retention；h-j SAT feedback branch、senescence-like retention anchor 和 fate-stabilization signal；k-l perturbation-linked validation 和 stabilization-axis summary；m-o aging/SAT signal、positional score 和 Result 5 summary。盆地示意图为概念性表示，而散点图、热图、条形图和箱线图概括锁定源数据输出。该图不声称完整因果闭合。",
    "Figure 6 | 肿瘤样可塑性是独立分支。",
    "肿瘤样可塑性被显示为区别于再生能力的高可塑性边界分支 27,36 。连续色条编码所示 tumor-like、stemness、differentiation 或 local-retention score，颜色强度越高表示评分越高。箱线图比较组间评分分布；热图颜色表示模块均值。肿瘤样分支不被解释为 salamander_blastema 或再生能力。",
    "Figure 7 | 标量 Phi 被拒绝为状态分离器。",
    "ROC、置换、bootstrap 和 KS 摘要显示，Phi 作为全局序参量不具备判别性 47 。线条颜色区分经验曲线、空模型或随机参考期望以及标注的观测统计量。ROC 对角线表示随机水平表现；置换曲线显示标签打乱空分布；bootstrap 区间显示估计偏移的不确定性。显著 KS 统计量被解释为分布形状差异，而不是可分性。",
    "Figure 8 | 潜在调控状态后验动力学提供工作表征。",
    "模型通过后验调控状态分配 P(Z|S,W_GRN) 表示每个观测状态，而不是通过全局标量坐标表示。彩色密度曲线和堆叠条形图显示分配给 adult_repair、embryonic_reactivation、salamander_blastema 和 salamander_intact 的后验质量。更高密度或条形高度表示显示摘要中的后验质量更高。对于状态 k，p_k(S)=P(Z=k|S,W_GRN)，期望局部转移倾向写作 E[Delta S|S]=sum_k p_k(S)F_k(S)。该记号描述 regime-conditioned 的计算表征，而不是已经实验识别出的连续时间生物学定律 6,22 。",
    "Figure 9 | 重叠与对称发散定义调控状态结构。",
    "重叠被解释为不可分性，而对称 KL 发散被解释为分布差异，而不是标量距离 22,37 。热图颜色强度对应显示矩阵值，单元格中印有数值。蓝色重叠矩阵量化共享后验或状态空间占据；红色条形长度排序最强的对称 KL 发散值。更高发散表示更大的分布差异，而不是沿生物学轴的移动。",
    "Supplementary Figure 1 | 干实验扰动一致性审计。",
    "既有扰动衍生输出用于检验通路相关后验偏移摘要是否与 regime-conditioned structure 方向一致。热图颜色编码指标值，颜色强度越高表示方向一致性、空模型分离、支持或偏差越强；灰色单元格表示共享的模型层级数值，而不是通路特异估计。模型比较面板报告全局 W_GRN 与 regime-conditioned W(Z)，一致性限制面板概括不可用的跨数据集比较、频繁反事实反转、跨状态不稳定和有限的 NOTCH/SHH 证据。Regime-conditioned W(Z) 相对于全局 W_GRN 降低了重构误差，但反事实反转和跨状态不稳定仍然存在；该分析仅支持部分 regime-conditioned perturbation consistency，不支持完整因果闭合或湿实验验证。",
]


def set_text(paragraph, text: str, chinese: bool = False) -> None:
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    if chinese:
        run.font.name = "Songti SC"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")


def build_docx(texts: list[str], output: Path, chinese: bool = False) -> None:
    doc = Document(str(BASE))
    nonempty = [p for p in doc.paragraphs if p.text.strip()]
    expected = "References"
    if len(nonempty) < len(texts) or nonempty[99].text.strip() != expected:
        raise RuntimeError("Unexpected manuscript paragraph structure; aborting language replacement.")
    for i, text in enumerate(texts):
        set_text(nonempty[i], text, chinese=chinese)
    if chinese:
        set_text(nonempty[99], "参考文献", chinese=True)
    doc.save(str(output))


def export_md(texts: list[str], refs: list[str], output: Path, chinese: bool = False) -> None:
    lines: list[str] = []
    for i, t in enumerate(texts):
        if i == 0:
            lines.append(f"# {t}")
        elif t in {"Abstract", "Introduction", "Results", "Discussion", "Limitations", "Methods", "Data availability", "Code availability", "Acknowledgements", "Author contributions", "Competing interests", "Figure legends", "References", "摘要", "引言", "结果", "讨论", "局限性", "方法", "数据可用性", "代码可用性", "致谢", "作者贡献", "利益冲突", "图注", "参考文献"}:
            lines.append(f"\n## {t}")
        elif i in {9, 11, 14, 17, 20, 23, 26, 29, 32, 35, 38, 47, 52, 54, 56, 58, 60, 62, 64}:
            lines.append(f"\n### {t}")
        else:
            lines.append(f"\n{t}")
    lines.extend([""] + refs)
    output.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def main() -> None:
    if len(EN) != 99 or len(ZH) != 99:
        raise RuntimeError(f"Unexpected text list lengths: EN={len(EN)}, ZH={len(ZH)}")

    build_docx(EN, EN_DOCX, chinese=False)
    build_docx(ZH, ZH_DOCX, chinese=True)

    src_doc = Document(str(BASE))
    nonempty = [p.text.strip() for p in src_doc.paragraphs if p.text.strip()]
    refs = nonempty[100:]
    export_md(EN + ["References"], refs, EN_MD)
    export_md(ZH + ["参考文献"], refs, ZH_MD, chinese=True)

    LOG.write_text(
        """# Bilingual polish change log

- Generated an English polished manuscript from the figure-fixed DOCX.
- Generated a Chinese academic manuscript from the same figure-fixed DOCX.
- Preserved all embedded figures, including the repaired Figure 4 and Figure 5.
- Preserved the reference entries; no DOI, PMID or citation metadata were fabricated.
- Preserved the locked latent-state-regime mixture framing and the conservative partial-closure interpretation.
- Chinese manuscript translates body text, administrative sections and figure legends; figure image content remains unchanged.
""",
        encoding="utf-8",
    )

    print(EN_DOCX)
    print(ZH_DOCX)
    print(EN_MD)
    print(ZH_MD)
    print(LOG)


if __name__ == "__main__":
    main()
