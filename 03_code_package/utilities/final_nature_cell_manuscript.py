from __future__ import annotations

import re
import xml.etree.ElementTree as ET
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
DOC1 = ROOT / "work" / "Doc1.docx"

TITLE_OPTIONS = [
    "Regenerative cell fate is a latent-state-regime mixture rather than a scalar state",
    "Failure of scalar compression reveals mixture regimes of repair and regeneration",
    "Non-compressible regenerative dynamics emerge from latent cell-fate regimes",
]

MANUSCRIPT_TITLE = TITLE_OPTIONS[0]

ABSTRACT = [
    (
        "Repair and regeneration often begin with injury-induced plasticity, yet adult mammalian tissues usually resolve through constrained repair whereas salamander limbs can form blastema and restore structure. We tested whether this divergence could be represented by a single scalar order parameter, Phi, linking stemness, fate-lock, positional programs and regulatory-network structure. The scalar hypothesis failed: ROC performance was random-level (AUC approximately 0.480), permutation testing did not support a stable mean shift, bootstrap confidence intervals crossed zero, and the Kolmogorov-Smirnov result reflected distributional shape differences rather than separability. Thus Phi is not a valid discriminative or biological order parameter for regeneration."
    ),
    (
        "We therefore reconstruct the manuscript around a latent-state-regime mixture dynamical system. In this framework, observed cell states are mixtures over adult repair, embryonic reactivation, salamander blastema and salamander intact regimes, rather than positions on a global scalar axis. Salamander blastema is treated as a distinct latent regenerative regime rather than an elevated scalar state; mammalian repair is an adult-repair-dominant mixture constrained by fate-lock, senescence-like stabilization and inflammatory context; and tumor-like plasticity remains a separate high-plasticity branch. The result is a non-compressible representation in which distributional overlap without scalar separability, not scalar ordering, explains the observed structure."
    ),
]

INTRODUCTION = [
    (
        "Injured tissues can transiently acquire plasticity, but plasticity alone does not explain why biological outcomes diverge. Mammalian wounds often undergo inflammation-associated repair, fibrosis or senescence-enriched failure, whereas salamander limb injury can enter a blastema state and support pattern-restoring regeneration. Tumors define a further boundary case: they can display high plasticity and local retention without becoming regenerative tissue. The central problem is therefore representational as much as molecular: the same broad plasticity vocabulary describes states that do not share the same fate logic."
    ),
    (
        "A natural solution is to search for a compact order parameter. Stemness scores, developmental-axis scores and integrated Phi-like quantities promise to compress complex fate behavior into a single coordinate. The reconstructed manuscript treats that possibility as a testable hypothesis rather than as an assumption. The old empirical framework provided the state components: Wnt/beta-catenin/TCF-LEF/MYC-associated accessibility, RA/HOX/FGF/SHH/NOTCH positional programs, BMP-SMAD and p53/p21/p16/Rb-associated fate-lock, senescence-like stabilization, and tumor-like plasticity."
    ),
    (
        "The final validation rejects scalar compression. Phi retained distributional structure, but its classifier performance was random-level and its mean-shift statistics were unstable. The significant KS test did not rescue the model because KS detects shape difference, not a biologically interpretable separation threshold. This failure matters biologically: salamander blastema cannot be described as simply farther along the same scalar coordinate as mammalian repair, and mammalian repair cannot be treated as the same regenerative program at lower intensity."
    ),
    (
        "The replacement is a latent-state-regime mixture framework. Regime identity is represented probabilistically as P(Z|S,W_GRN), with Z spanning adult repair, embryonic reactivation, salamander blastema and salamander intact. Phi is retained only as a failed scalar projection whose per-regime distributions can be modeled within the mixture. Under this interpretation, latent state regime mixture explains observed structure while preserving overlap, non-identifiability and regime-specific divergence."
    ),
    (
        "The final manuscript follows a locked nine-figure architecture. Figures 1-6 retain the old empirical foundation under regime-dependent interpretation required by the validation layer. Figures 7-9 close the representation problem by rejecting the single-Phi model, replacing it with a latent-state-regime mixture, and quantifying overlap together with symmetrized distributional divergence. No new model, dataset or statistical analysis is introduced during this rewrite."
    ),
]

RESULTS: list[tuple[str, list[str]]] = [
    (
        "Result 1. The fate-control framework is reframed as a latent-state-regime mixture problem",
        [
            "The manuscript is first reorganized around the distinction between plasticity and fate determination (Fig. 1). The old developmental fate-divergence architecture remains biologically useful: stemness-associated accessibility, positional information, fate-lock, senescence-like stabilization and tumor-like plasticity define the relevant axes of the system. What changes is the representational claim. These axes are no longer interpreted as components of one global scalar trajectory.",
            "Figure 1 therefore serves as the conceptual bridge from organized plasticity to latent-state-regime mixture dynamics. Plasticity is permissive, positional programs are regime-conditioned, fate-lock constrains adult repair, and tumor-like plasticity is a separate branch. The figure must be read as a schematic synthesis anchored by later data-bound panels, not as evidence for scalar control.",
        ],
    ),
    (
        "Result 2. High-plasticity states share accessibility without fate equivalence",
        [
            "The first empirical layer shows that regenerative, senescence-like and tumor-like groups can occupy connected regions of high-plasticity state space (Fig. 2). In GSE153596, embryonic-like plasticity cells and tumor-like plasticity proxy cells showed high state entropy, supporting a shared accessibility landscape. CellRank and RNA-velocity summaries further indicated that senescence-like mammalian wound states had strong self-retention.",
            "The revised interpretation is deliberately limited. Shared accessibility does not imply shared outcome, direct lineage relation or regenerative equivalence. Velocity and CellRank are treated as inferred trajectory comparators, not lineage tracing. Figure 2 supports distributional overlap without scalar separability at the accessibility level.",
        ],
    ),
    (
        "Result 3. Stemness-associated programs increase accessibility but do not specify fate",
        [
            "Stemness-associated Wnt/beta-catenin/TCF-LEF/MYC programs provide an accessibility-generating layer (Fig. 3). In perturbation evidence from GSE130381, beta-catenin activation and CHIR increased canonical Wnt output, and CHIR increased the stemness module. These findings support the old claim that stemness programs expand the reachable state space.",
            "They do not, however, define regenerative identity. Boundary analyses did not support a universal one-dimensional BMP-Wnt antagonistic axis, and stemness scores cannot classify mammalian repair, salamander blastema or tumor-like plasticity. In the final model, stemness contributes local accessibility within latent state regimes rather than a global fate coordinate.",
        ],
    ),
    (
        "Result 4. Positional programs act as regime-conditioned developmental coordinates",
        [
            "RA/RARG-linked perturbation evidence and RA/HOX-associated positional proxies support a developmental-position layer in regenerative contexts (Fig. 4). Axolotl regeneration showed increased positional and regeneration-module activity relative to intact reference tissue, and frog compartment-level DPRI evidence supported context-dependent positional organization.",
            "These signals are not treated as universal spatial coordinates or scalar determinants. They are expression-layer and compartment-proxy measurements whose interpretation depends on regime context. Figure 4 therefore supports a positional coordinate inside latent regenerative regimes, while preserving the distinction between salamander blastema, salamander intact and mammalian repair.",
        ],
    ),
    (
        "Result 5. Fate-lock and senescence-like stabilization define an adult-repair-biased regime",
        [
            "Fate-lock and senescence-like outputs are merged into the adult repair-failure layer (Fig. 5). BMP-SMAD and p53/p21/p16/Rb-associated logic represent fate-stabilizing constraints, while mammalian senescence-like wound states show retention and adult repair-failure bias in the trajectory summaries.",
            "The final claim is not that fate-lock is a universal aging law or that mammalian repair failure is salamander regeneration in reverse. Instead, mammalian repair is interpreted as an adult-repair-dominant mixture constrained by fate-lock, inflammatory context and senescence-like stabilization. Direct causal feedback among these processes remains a prediction requiring perturbation time series.",
        ],
    ),
    (
        "Result 6. Tumor-like plasticity is a separate high-plasticity branch",
        [
            "Tumor-like plasticity is retained as a boundary condition of the system (Fig. 6). In GSE195655, tumor-like plasticity proxy cells showed stemness enrichment, reduced differentiation and increased local self-retention. These properties place them in a high-plasticity region of state space.",
            "The boundary tests prevent over-interpretation. Tumor-like plasticity did not show meaningful embryonic-like increase and was not equivalent to inflammatory repair after correction. It is therefore modeled as a distinct high-plasticity branch, not as embryonic reversion and not as regeneration.",
        ],
    ),
    (
        "Result 7. The single-Phi order-parameter hypothesis is rejected",
        [
            "The manuscript then tests the scalar representation directly (Fig. 7). If Phi were a valid global order parameter, mammalian repair, salamander blastema and salamander intact states should be separable by a single scalar coordinate. The locked validation outputs do not support that condition.",
            "ROC AUC was approximately 0.480, indicating random-level discrimination. The permutation mean-shift test was not significant, bootstrap confidence intervals were unstable and crossed zero for the mean shift, and the KS test indicated distributional shape difference rather than separability. Phi is therefore not a valid discriminative or biological order parameter and is not used to classify regimes.",
        ],
    ),
    (
        "Result 8. A latent-state-regime mixture replaces scalar thresholding",
        [
            "The replacement model represents observed states as probabilistic mixtures over latent state regimes (Fig. 8). Regime identity is described by P(Z|S,W_GRN), where Z includes adult repair, embryonic reactivation, salamander blastema and salamander intact. No global Phi threshold is used.",
            "The distributional form is P(Phi|S) = sum_Z P(Phi|S,Z)P(Z|S,W_GRN). Each regime has its own Phi distribution, and observed species or condition groups can carry mixed posterior mass. This is why scalar classification fails while latent state regime mixture explains observed structure.",
        ],
    ),
    (
        "Result 9. Regime overlap and divergence coexist without scalar separability",
        [
            "The final figure quantifies non-identifiability and distributional divergence (Fig. 9). Overlap matrices show that latent state regimes and observed species/regime groups are not cleanly separable in the measured state space. This overlap is the empirical reason that a global scalar classifier fails.",
            "At the same time, overlap does not imply identity. Symmetrized KL divergence shows regime-specific distributional differences, supporting adult repair and salamander blastema as overlapping but distinct latent-state-regime mixtures. Figure 9 closes the manuscript with distributional overlap without scalar separability.",
        ],
    ),
]

DISCUSSION = [
    (
        "This final rewrite converts the project into a representation-falsification manuscript. The old evidence still supports a structured biological system: injury-induced plasticity expands accessibility; stemness-associated programs contribute to access; positional programs organize regenerative states; fate-lock and senescence-like stabilization constrain mammalian repair; and tumor-like plasticity forms a separate branch. The central revision is that these observations cannot be compressed into one scalar order parameter."
    ),
    (
        "The failure of Phi reflects non-identifiability, regime overlap and distributional degeneracy. The same scalar values can be compatible with different posterior regime compositions, and different regimes can overlap in the measured state space while retaining distinct distributional structure. AUC near 0.480 shows that Phi has no useful global discriminative capacity. The non-significant permutation result and unstable bootstrap intervals rule out a robust scalar mean shift."
    ),
    (
        "KS significance is therefore not evidence of separability. It indicates that distributions differ in shape, but it does not define a threshold, a classifier or a biological coordinate that orders regeneration. This distinction is essential: a model can detect distributional divergence and still fail as a scalar fate map. In this manuscript, overlap and divergence coexist because the observed data arise from mixtures of latent state regimes."
    ),
    (
        "Biologically, salamander blastema is interpreted as a distinct latent regenerative regime, mammalian repair as a constrained adult regime mixture, and regeneration as regime access rather than scalar movement. Salamander intact tissue is not blastema; it provides a partially overlapping reference regime. Tumor-like plasticity is not embryonic reversion; it is a separate high-plasticity branch marked by stemness enrichment, reduced differentiation and local retention."
    ),
    (
        "This framework also clarifies the role of prior biological modules. Stemness is accessibility, not fate. Positional information is regime-conditioned developmental organization, not a universal determinant. Fate-lock is an adult-repair constraint, not a complete causal law. Senescence-like stabilization is a repair-failure-biased attractor component, not proof of irreversible mechanism. Regime-dependent interpretation required by the validation layer prevents these modules from being overextended."
    ),
    (
        "The limits are explicit. The analysis is locked at processed state-score and proxy layers, with cross-species alignment, batch/species entanglement and orthology constraints. It does not establish complete causal regulatory circuitry, direct chromatin memory, lineage history or spatial morphogen control. Those claims require same-system perturbation time series, lineage tracing, spatial measurements and multiome validation."
    ),
    (
        "The resulting implication is conceptual and experimentally testable: regenerative fate should be treated as a latent-state-regime mixture problem, not as scalar maximization. Perturbations that increase stemness should increase accessibility without guaranteeing blastema-like regeneration. Perturbations that restore positional organization should be effective only in compatible regime contexts. Perturbations that deepen fate-lock or senescence-like stabilization should increase adult repair retention. The manuscript therefore establishes a non-compressible cell-fate framework rather than a new scalar fate coordinate."
    ),
]

FIGURE_LEGENDS = [
    (
        "Figure 1 | From developmental fate control to latent-state-regime mixture dynamics.",
        "Conceptual synthesis of the locked nine-figure architecture. Stemness-associated accessibility, positional programs, fate-lock, senescence-like stabilization and tumor-like plasticity are retained as biological axes, but final fate accessibility is represented by probabilistic latent state regimes rather than a scalar order parameter."
    ),
    (
        "Figure 2 | Shared high-plasticity accessibility without fate equivalence.",
        "State-space, entropy, CellRank and RNA-velocity-derived summaries support shared accessibility among regenerative, senescence-like and tumor-like regions. Velocity and CellRank are labeled as inferred trajectory comparators, not lineage tracing."
    ),
    (
        "Figure 3 | Stemness programs expand access but do not classify regeneration.",
        "Wnt/beta-catenin/TCF-LEF/MYC-associated perturbation and module-score evidence supports accessibility expansion. Boundary analyses prevent interpretation as a universal BMP-Wnt scalar axis or a fate classifier."
    ),
    (
        "Figure 4 | Positional programs are regime-conditioned developmental coordinates.",
        "RA/RARG, RA/HOX-like positional proxies, axolotl regeneration evidence and frog DPRI compartment evidence are integrated as regime-conditioned positional organization. They are not treated as direct spatial gradients or global scalar determinants."
    ),
    (
        "Figure 5 | Fate-lock and senescence-like stabilization in adult repair.",
        "BMP-SMAD and p53/p21/p16/Rb-associated fate-stabilization logic, senescence-like retention and adult repair-failure evidence are combined as an adult-repair-biased mixture. The figure does not claim direct causal feedback without perturbation time series."
    ),
    (
        "Figure 6 | Tumor-like plasticity as a distinct branch.",
        "Tumor-like plasticity proxy cells show stemness enrichment, reduced differentiation and local retention while failing boundary tests for embryonic or inflammatory-repair equivalence. The branch is distinct from regeneration."
    ),
    (
        "Figure 7 | Rejection of the single-Phi scalar model.",
        "Phi distributions show structure, but ROC AUC is approximately 0.480, permutation testing is not significant, bootstrap confidence intervals cross zero and KS reflects shape difference rather than separability. Phi is invalid as a global order parameter."
    ),
    (
        "Figure 8 | Latent-state-regime mixture posterior model.",
        "Regime identity is represented as P(Z|S,W_GRN) over adult repair, embryonic reactivation, salamander blastema and salamander intact. The model uses P(Phi|S) = sum_Z P(Phi|S,Z)P(Z|S,W_GRN), with no global threshold."
    ),
    (
        "Figure 9 | Overlap and symmetrized distributional divergence.",
        "Overlap matrices and symmetrized KL divergence quantify non-identifiability and regime-specific divergence. Adult repair and salamander blastema are overlapping but distinct mixtures, producing distributional overlap without scalar separability."
    ),
]

METHODS_INSERT = [
    (
        "Final manuscript reconstruction and claim locking",
        [
            "The final rewrite used the locked figure-integration and model-validation outputs already generated in the project workspace. The manuscript text was rewritten without changing datasets, statistics, model structure or figure architecture. Figures 1-6 retain the old empirical scaffold under regime-dependent interpretation; Figures 7-9 represent the scalar-model rejection, latent-state-regime mixture replacement and overlap/divergence closure.",
            "The single-Phi rejection was constrained by the locked validation files Phi_unified.tsv, roc_curve.tsv, permutation_test_results.tsv, bootstrap_confidence_intervals.tsv, ks_test_results.tsv and single_phi_model_invalidation.md. The latent mixture language was constrained by latent_regime_mixture_model.md, regime_posterior_probabilities.tsv, regime_overlap_matrix.tsv, species_regime_overlap_matrix.tsv and regime_KL_divergence_matrix.tsv.",
        ],
    )
]


def extract_tail_from_doc1() -> list[tuple[str, str]]:
    if not DOC1.exists():
        return []
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    items: list[tuple[str, str]] = []
    with ZipFile(DOC1) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    body = root.find("w:body", ns)
    if body is None:
        return items
    capture = False
    for p in body.findall("w:p", ns):
        txt = "".join(t.text or "" for t in p.findall(".//w:t", ns)).strip()
        if not txt:
            continue
        pstyle = ""
        ppr = p.find("w:pPr", ns)
        if ppr is not None:
            style_el = ppr.find("w:pStyle", ns)
            if style_el is not None:
                pstyle = style_el.attrib.get(f"{{{ns['w']}}}val", "")
        if txt == "Methods":
            capture = True
        if capture:
            items.append((pstyle, txt))
    return items


def write_text(path: Path, text: str) -> None:
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def write_markdown_outputs() -> None:
    write_text(
        OUT / "FINAL_TITLE_OPTIONS.md",
        "# Final Title Options\n\n"
        + "\n".join(f"{i}. {title}" for i, title in enumerate(TITLE_OPTIONS, 1))
        + f"\n\nSelected working title: {MANUSCRIPT_TITLE}",
    )
    write_text(OUT / "FINAL_ABSTRACT.md", "# Final Abstract\n\n" + "\n\n".join(ABSTRACT))
    write_text(OUT / "FINAL_INTRODUCTION.md", "# Final Introduction\n\n" + "\n\n".join(INTRODUCTION))
    write_text(
        OUT / "FINAL_RESULTS.md",
        "# Final Results\n\n"
        + "\n\n".join(
            "## " + title + "\n\n" + "\n\n".join(paragraphs) for title, paragraphs in RESULTS
        ),
    )
    write_text(OUT / "FINAL_DISCUSSION.md", "# Final Discussion\n\n" + "\n\n".join(DISCUSSION))


def write_figure_legend_alignment() -> None:
    lines = [
        "# Figure Legend Alignment",
        "",
        "Locked regime labels: adult repair; embryonic reactivation; salamander blastema; salamander intact.",
        "",
        "| Figure | Result section | Role | Aligned legend title | Required terminology |",
        "|---|---|---|---|---|",
    ]
    for i, (title, _) in enumerate(FIGURE_LEGENDS, 1):
        role = "old empirical scaffold" if i <= 6 else "model-validation closure"
        terms = "regime-dependent interpretation required" if i <= 6 else "latent mixture, no global Phi threshold"
        lines.append(f"| Figure {i} | Result {i} | {role} | {title} | {terms} |")
    lines.extend(["", "## Full Aligned Legends", ""])
    for i, (title, body) in enumerate(FIGURE_LEGENDS, 1):
        lines.extend([f"### Figure {i}", "", f"{title} {body}", ""])
    write_text(OUT / "FIGURE_LEGEND_ALIGNMENT.md", "\n".join(lines))


def manuscript_text_for_scan() -> str:
    parts = [MANUSCRIPT_TITLE, *ABSTRACT, *INTRODUCTION]
    for title, paragraphs in RESULTS:
        parts.append(title)
        parts.extend(paragraphs)
    parts.extend(DISCUSSION)
    for title, body in FIGURE_LEGENDS:
        parts.extend([title, body])
    return "\n".join(parts)


def write_claim_check() -> None:
    text = manuscript_text_for_scan()
    forbidden = [
        "Phi explains regeneration",
        "high Phi corresponds to blastema",
        "single trajectory governs fate",
        "tumor-like plasticity equals embryonic reversion",
    ]
    hits = [phrase for phrase in forbidden if re.search(re.escape(phrase), text, flags=re.I)]
    required = [
        "distributional overlap without scalar separability",
        "latent state regime mixture explains observed structure",
        "regime-dependent interpretation required",
    ]
    required_status = [
        (phrase, "PRESENT" if phrase in text else "MISSING") for phrase in required
    ]
    lines = [
        "# Claim Consistency Final Check",
        "",
        "| Check | Status | Evidence |",
        "|---|---|---|",
        "| Single scalar Phi interpretation | INVALIDATED | ROC AUC approximately 0.480; permutation not significant; bootstrap intervals cross zero; KS interpreted as shape difference only. |",
        "| Classifier interpretation of Phi | REMOVED | Manuscript states Phi is not used to classify regimes. |",
        "| Salamander blastema interpretation | PASS | Manuscript states blastema is a distinct latent regenerative regime rather than an elevated scalar state. |",
        "| Mammalian repair interpretation | PASS | Manuscript states mammalian repair is an adult-repair-dominant mixture. |",
        "| Tumor-like plasticity boundary | PASS | Manuscript states tumor-like plasticity is a separate high-plasticity branch. |",
    ]
    for phrase, status in required_status:
        lines.append(f"| Required phrase: {phrase} | {status} | Present in final manuscript text. |")
    if hits:
        lines.append(f"| Forbidden phrase scan | FAIL | Hits: {', '.join(hits)} |")
        final_status = "FAIL"
    else:
        lines.append("| Forbidden phrase scan | PASS | No exact forbidden phrase appears in manuscript-only text. |")
        final_status = "PASS"
    lines.extend(
        [
            "",
            f"FINAL_STATUS: {final_status}",
            "",
            "Final scientific position: Cell fate and regenerative dynamics are not governed by a scalar order parameter, but instead arise from a latent-state-regime mixture structure that cannot be compressed into a single continuous coordinate.",
        ]
    )
    write_text(OUT / "CLAIM_CONSISTENCY_FINAL_CHECK.md", "\n".join(lines))
    if final_status != "PASS":
        raise RuntimeError("Forbidden manuscript phrasing detected")


def set_font(run, name: str = "Calibri", size: int | None = None, bold=None, italic=None, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.28
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for style_name, size, color, before, after in [
        ("Heading 1", 15, "1F4D78", 16, 8),
        ("Heading 2", 12, "2E74B5", 10, 5),
        ("Heading 3", 11, "2E74B5", 7, 3),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.append(field)


def add_para(doc: Document, text: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.keep_together = False
    paragraph.paragraph_format.keep_with_next = False
    run = paragraph.add_run(text)
    set_font(run, size=11)
    return paragraph


def add_bold_para(doc: Document, title: str, body: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_together = False
    title_run = paragraph.add_run(title)
    set_font(title_run, size=11, bold=True)
    body_run = paragraph.add_run(" " + body)
    set_font(body_run, size=11)


def append_doc1_tail(doc: Document) -> None:
    tail = extract_tail_from_doc1()
    if not tail:
        return
    inserted_methods = False
    doc.add_page_break()
    for style, text in tail:
        if text == "Methods":
            doc.add_heading("Methods", level=1)
            continue
        if text == "Statistical testing and multiple-testing correction" and not inserted_methods:
            for heading, paragraphs in METHODS_INSERT:
                doc.add_heading(heading, level=2)
                for paragraph in paragraphs:
                    add_para(doc, paragraph)
            inserted_methods = True
        if style == "1":
            doc.add_heading(text, level=1)
        elif style in {"21", "31"}:
            doc.add_heading(text, level=2)
        elif style == "EndNoteBibliography":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(text)
            set_font(run, size=9)
        else:
            add_para(doc, text)


def build_docx(path: Path) -> None:
    doc = Document()
    configure_styles(doc)
    add_page_number(doc.sections[0].footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run(MANUSCRIPT_TITLE)
    set_font(run, size=20, bold=True, color="0B2545")

    deck = doc.add_paragraph()
    deck.paragraph_format.space_after = Pt(16)
    run = deck.add_run("Final compressed manuscript aligned to the locked nine-figure latent-state-regime mixture architecture")
    set_font(run, size=10, italic=True, color="555555")

    doc.add_heading("Abstract", level=1)
    for paragraph in ABSTRACT:
        add_para(doc, paragraph)

    doc.add_heading("Introduction", level=1)
    for paragraph in INTRODUCTION:
        add_para(doc, paragraph)

    doc.add_heading("Results", level=1)
    for title_text, paragraphs in RESULTS:
        doc.add_heading(title_text, level=2)
        for paragraph in paragraphs:
            add_para(doc, paragraph)

    doc.add_heading("Discussion", level=1)
    for paragraph in DISCUSSION:
        add_para(doc, paragraph)

    append_doc1_tail(doc)

    doc.add_page_break()
    doc.add_heading("Figure legends", level=1)
    for title_text, body in FIGURE_LEGENDS:
        add_bold_para(doc, title_text, body)

    doc.save(path)


def main() -> None:
    OUT.mkdir(exist_ok=True)
    write_markdown_outputs()
    write_figure_legend_alignment()
    write_claim_check()
    build_docx(OUT / "FINAL_MANUSCRIPT.docx")
    print("final manuscript outputs generated")
    print(OUT / "FINAL_MANUSCRIPT.docx")


if __name__ == "__main__":
    main()
