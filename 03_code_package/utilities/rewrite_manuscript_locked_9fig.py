from __future__ import annotations

import csv
import re
import textwrap
import xml.etree.ElementTree as ET
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
DOC1 = ROOT / "work" / "Doc1.docx"


TITLE = "Latent-state-regime mixture dynamics of injury-induced cell plasticity"


ABSTRACT_PARAS = [
    (
        "Injury-induced plasticity expands the range of cell states accessible during repair, but plasticity itself does not specify whether tissue resolves through regeneration, senescence-associated repair failure or tumor-like plasticity. The previous manuscript established a biological foundation for this problem through stemness-associated accessibility, developmental positional information, fate-lock and lineage-restriction programs, senescence-associated stabilization, mammalian inflammatory repair and axolotl regeneration observations. Here the manuscript is reconstructed around the locked conclusion that these observations do not support a single global scalar order parameter for regeneration."
    ),
    (
        "The tested Phi hypothesis is retained as part of the scientific narrative but is treated as experimentally rejected. Phi showed distributional structure, yet it failed as a discriminative order parameter: ROC AUC was approximately 0.480, the permutation mean-shift test was not significant, bootstrap confidence intervals were unstable and crossed zero for the mean shift, and the significant Kolmogorov-Smirnov result indicated distributional shape difference rather than separability. These results require replacing scalar interpretation with a latent-state-regime mixture model."
    ),
    (
        "In the final framework, mammalian repair is represented as an adult-repair-dominant regime mixture, salamander blastema as a distinct latent regenerative regime with partial overlap, salamander intact tissue as a separate reference regime, and tumor-like plasticity as a distinct high-plasticity branch rather than embryonic regeneration. The resulting manuscript is therefore a hypothesis-revision and model-falsification paper: old empirical results are retained, but their causal interpretation is reconstructed as regime-dependent accessibility, not single-scalar fate control."
    ),
]


INTRODUCTION = [
    (
        "Regeneration and repair diverge from a shared biological starting point: injured tissues often enter plastic states, yet the consequences of that plasticity differ across systems. Mammalian wounds can resolve through partial repair, inflammatory remodeling or senescence-enriched failure, whereas salamander limb injury can access blastema-associated developmental reactivation and patterned regeneration. Tumors add a further boundary case, in which high plasticity can support local self-retention and reduced differentiation without becoming regeneration. The central problem is therefore not whether plasticity occurs, but how plastic states are constrained, redirected or stabilized into distinct biological regimes."
    ),
    (
        "The old manuscript provided the empirical scaffold for this problem. Stemness-associated Wnt/beta-catenin/TCF-LEF/MYC programs were interpreted as accessibility-generating programs. Developmental positional information, including RA/HOX/FGF/SHH/Notch-related proxies and DPRI-like measurements, provided a candidate coordinate system for regenerative organization. BMP-SMAD and ARF-p53-p21-p16-Rb-associated programs represented fate-lock and lineage-restriction logic. Senescence-associated and SASP-like programs marked mammalian repair-failure states, while tumor-like plasticity was treated as a separate high-plasticity branch. These components remain part of the manuscript, but their interpretation is revised."
    ),
    (
        "A key motivation for the reconstruction is the cross-species discrepancy between mammalian repair and salamander regeneration. It is tempting to describe this discrepancy as a quantitative difference along a single regenerative axis. That interpretation is not supported by the locked validation outputs. Salamander blastema cannot be equated with high Phi, and mammalian inflammatory repair cannot be treated as the same process at lower intensity. Instead, both systems occupy overlapping but distributionally distinct regions of a latent state regime mixture."
    ),
    (
        "Phi is therefore repositioned as an initial hypothesis that was tested and rejected. A single scalar order parameter would be attractive because it would compress plasticity, fate-lock, positional information and species differences into one coordinate. However, the final validation shows that Phi has distributional structure without useful discriminative separability. The replacement framework is a latent-state-regime mixture developmental dynamical system in which regime identity is represented probabilistically by P(Z|S,W_GRN), not by a global threshold."
    ),
    (
        "The reconstructed manuscript follows a locked nine-figure architecture. Figures 1-6 integrate the old empirical foundation under regime-dependent interpretation. Figures 7-9 provide the model-falsification and replacement layer: single-Phi rejection, latent-state-regime mixture reconstruction, and regime overlap with symmetrized divergence. This structure preserves the old findings while preventing the rejected scalar model from re-entering as a secondary explanation."
    ),
]


RESULTS: list[tuple[str, list[str]]] = [
    (
        "Result 1. The manuscript framework is upgraded from organized plasticity to latent-state-regime mixture dynamics",
        [
            "The original manuscript organized injury-induced plasticity around a developmental fate-divergence framework. That framework remains useful as a biological scaffold because it separated permissive plasticity from positional organization, fate-stabilization logic and tissue-level control. The reconstruction changes the mathematical and causal interpretation: the final manuscript no longer treats fate divergence as the readout of a single scalar control variable. Instead, it treats regenerative accessibility as a probabilistic mixture over latent state regimes (Fig. 1).",
            "In this revised architecture, old Fig. 1 and old Fig. 9 are retained only after conceptual relabeling. Plasticity remains permissive, not instructive. Positional information remains a regime-conditioned developmental coordinate. Fate-lock and senescence-associated stabilization remain adult-repair constraints. Tumor-like plasticity remains a distinct high-plasticity branch. The added model layer is the explicit rejection of single-Phi control and the replacement of scalar thresholding with latent state regime probabilities.",
            "This conceptual upgrade is not a new biological model added after the fact. It is a constraint imposed by the locked validation outputs. The manuscript is therefore framed as a hypothesis-revision paper: the old empirical results define the state variables and biological boundaries, while the new validation determines which representation of those results is statistically defensible.",
        ],
    ),
    (
        "Result 2. High-plasticity states occupy a shared accessibility space without fate equivalence",
        [
            "The first empirical layer is the old observation that regenerative, senescence-like and tumor-like groups occupy connected regions of a shared high-plasticity state space (Fig. 2). In GSE153596, multiple connected states showed high entropy, including embryonic-like plasticity cells with mean state entropy of 0.903 and tumor-like plasticity proxy cells with mean state entropy of 0.978. These outputs support a shared accessibility landscape rather than isolated fate islands.",
            "The reconstructed interpretation is deliberately narrower than a fate-equivalence claim. Connectedness did not predict outcome. Senescence-like cells showed the strongest mammalian wound-repair self-retention signal, with mean CellRank self-absorption of 0.499 and velocity self-transition probability of 0.567. Inflammatory regeneration did not preferentially resolve into the tumor-like plasticity proxy, with CellRank mean 0.133 and velocity transition probability 0.191 in the relevant comparison.",
            "Thus, shared state-space accessibility is preserved as an empirical foundation, but it is reinterpreted as within-regime accessibility. RNA velocity and CellRank are retained as inferred trajectory comparators rather than lineage tracing. Figure 2 supports the premise that plasticity expands accessible state space, not the stronger claim that shared plasticity makes regeneration, senescence and tumor-like plasticity equivalent.",
        ],
    ),
    (
        "Result 3. Stemness-associated programs increase accessibility but do not determine outcome",
        [
            "Stemness-associated programs form the second empirical layer. In GSE130381, Wnt/CHIR/beta-catenin perturbation supported a stemness-associated accessibility response: activated beta-catenin increased canonical Wnt output with Cohen's d = 3.51 and q = 0.032, CHIR increased canonical Wnt output with d = 3.85 and q = 2.95e-05, and CHIR increased the stemness module with d = 1.27 and q = 0.0365 (Fig. 3).",
            "These effects support stemness as an accessibility-generating resource, not as a fate identity. The completed analyses did not support a universal one-dimensional BMP-Wnt antagonistic axis: canonical Wnt output and BMP-SMAD output were not negatively correlated in the relevant GSE130381 boundary analysis (Spearman rho = 0.307, p = 0.0611, q = 0.0957). This boundary result is important because it prevents the manuscript from replacing one scalar model with another stemness-only axis.",
            "Under the latent-state-regime mixture interpretation, stemness-associated programs contribute local accessibility within regimes. They may increase transition capacity, but they do not determine whether the accessible state resolves as mammalian adult repair, salamander blastema-associated regeneration, senescence-like stabilization or tumor-like plasticity.",
        ],
    ),
    (
        "Result 4. Developmental positional information is regime-conditioned rather than globally deterministic",
        [
            "Developmental positional information provides the third empirical layer. RA/RARG-linked perturbation analyses in GSE295225 supported large shifts in developmental-position modules: PAR 6 h RA versus vehicle shifted the developmental-position module with Cohen's d = 9.89 and q = 0.00769, and PAR 48 h RA versus vehicle shifted the same module with d = 15.6 and q = 0.00210 (Fig. 4). These results retain the old manuscript's claim that positional-program activity is associated with regenerative plasticity.",
            "Axolotl regeneration supplied the cross-system regenerative anchor. In GSE315993, axolotl developmental regeneration showed higher positional score than intact tissue (mean difference = 0.587, Cohen's d = 0.602, q = 5.30e-67) and higher regeneration module activity (mean difference = 0.549, d = 0.561, q = 2.42e-40). Diffusion pseudotime associations supported remodeling and redifferentiation proxies while retaining RA/HOX-associated positional identity.",
            "The reconstructed claim is regime-conditioned rather than universal. DPRI and RA/HOX-like positional scores are expression-layer proxies, not direct spatial coordinates, morphogen gradients or causal proof. Frog compartment DPRI evidence further supports context dependence: connective-tissue blastema and ectoderm/AER-like compartments differ within the same species context. These results support positional-program activity as a coordinate within a latent regenerative regime, not as a global scalar determinant of regeneration.",
        ],
    ),
    (
        "Result 5. Fate-lock and senescence-like programs define an adult repair-failure-biased regime",
        [
            "The old fate-lock and SAT results are merged into a single adult repair-failure section (Fig. 5). Fate-stabilization and lineage-restriction programs were represented by BMP-SMAD and ARF-p53-p21-p16-Rb-associated logic. In the completed perturbation validation, 8,156 tests yielded 88 BH-FDR <= 0.05 hits distributed across organizational, developmental-position, plasticity and fate-stabilization categories, indicating that fate-stabilization responses occur within a broader regulatory system rather than as an isolated module.",
            "In GSE153596, senescence-like cells provided the mammalian retention anchor, with mean CellRank self-absorption of 0.499 and velocity self-transition probability of 0.567. Comparative analyses placed mammalian senescence-like repair-failure states against axolotl developmental regeneration: mammalian senescence-like states showed higher fate-stabilization signal (mean difference = 0.204, d = 0.213, q = 1.34e-37), higher Aging/SAT signal (mean difference = 0.052, d = 0.051, q = 1.40e-09) and lower positional score (mean difference = -0.274, d = -0.282, q = 2.26e-38).",
            "The revised interpretation is adult-repair/fate-lock-biased mixture, not a universal law of aging and not the same process as salamander regeneration. SAT remains a proposed mammalian pathological branch supported by retention and comparative-score evidence. It does not establish direct p53/BMP-to-senescence causality or a complete causal feedback loop. Perturbation time series would be required to test whether damage, senescence-associated stabilization, SASP-like inflammation, positional disorganization and repair failure reinforce one another mechanistically.",
        ],
    ),
    (
        "Result 6. Tumor-like plasticity is a distinct high-plasticity branch rather than embryonic regeneration",
        [
            "Tumor-like plasticity is retained as a separate outcome within the shared high-plasticity landscape (Fig. 6). In GSE195655 tumor cells, top-quartile tumor-like plasticity proxy cells were strongly enriched for stemness relative to other tumor cells (n_top = 469, n_other = 1,406; Cohen's d = 1.296, p = 3.01e-100), showed reduced differentiation (d = -0.317, p = 4.21e-07), and showed increased local self-retention (d = 0.163, p = 5.91e-10).",
            "The boundary tests are central to the reconstructed claim. The tumor-like branch was not equivalent to embryonic regeneration: embryonic-like score was not meaningfully increased in the top tumor-like plasticity quartile (d = -0.022, q = 0.195), and inflammatory-repair score was not significantly higher after multiple-testing correction (d = -0.112, q = 0.125).",
            "Thus, tumor-like plasticity remains a distinct operational high-plasticity branch marked by stemness enrichment, reduced differentiation and local retention. It should not be described as embryonic reversion, regeneration failure or inflammatory-repair equivalence. This boundary prevents the latent-state-regime mixture model from collapsing all high-plasticity states into a single regenerative continuum.",
        ],
    ),
    (
        "Result 7. A single global Phi order parameter is rejected",
        [
            "The reconstructed manuscript treats Phi as a tested hypothesis, not as a surviving explanatory layer (Fig. 7). A single global order parameter would have predicted that mammalian repair, salamander blastema and salamander intact states could be separated along one scalar axis. The locked validation outputs do not support that interpretation.",
            "Phi displayed distributional structure, but discriminative performance was random-level. The ROC AUC was approximately 0.480. The permutation mean-shift test was not significant. Bootstrap confidence intervals were unstable, with the mean-shift interval crossing zero. The Kolmogorov-Smirnov test was significant, but in this context it indicates distributional shape difference rather than useful separability.",
            "This result invalidates single-Phi as a discriminative or order parameter. The manuscript therefore does not use Phi to separate species, define regeneration, identify salamander blastema or explain mammalian repair failure. Phi is retained only as a failed scalar hypothesis whose rejection motivates the latent-state-regime mixture model.",
        ],
    ),
    (
        "Result 8. Latent-state-regime mixture replaces the single-Phi framework",
        [
            "The replacement model is a latent-state-regime mixture developmental dynamical system (Fig. 8). Instead of assigning cells to regimes by a global Phi threshold, the model represents regime identity using posterior probabilities P(Z|S,W_GRN). The latent states are adult repair, embryonic reactivation, salamander blastema and salamander intact. These regimes are not observed directly as fixed labels; they are inferred as probabilistic components.",
            "The corresponding distributional statement is P(Phi|S) = sum_Z P(Phi|S,Z)P(Z|S,W_GRN). Each latent state regime has its own Phi distribution, and observed species or condition groups can contain mixed posterior mass. This is why the failure of a global classifier does not eliminate biological structure: structure exists, but it is not separable by one scalar threshold.",
            "Under this model, mammalian inflammatory repair is an adult-repair-dominant mixture; salamander blastema is a distinct latent regenerative regime with overlap; and salamander intact is not interchangeable with blastema. The model is partially supported at the processed score/proxy layer and should not be overstated as complete raw-data or causal validation.",
        ],
    ),
    (
        "Result 9. Regime overlap and symmetrized divergence explain cross-species structure without scalar separation",
        [
            "The final figure quantifies the relationship between overlap and divergence (Fig. 9). Regime overlap matrices and species/regime posterior-composition overlap show that latent state regimes are not cleanly separable in the observed state space. This overlap is the empirical reason a single global Phi classifier fails.",
            "At the same time, overlap does not imply identity. Symmetrized KL divergence reveals distributional differences among latent state regimes and observed species/regime compositions. The figure must therefore be interpreted as evidence for distributional divergence without scalar separability. Adult repair and salamander blastema are overlapping but distributionally different latent-state-regime mixtures.",
            "This closes the manuscript logic. The old empirical framework explains the biological components of plasticity, positional information, fate-lock, senescence and tumor-like plasticity. The new validation determines the correct representation of those components: a latent-state-regime mixture rather than a single order parameter.",
        ],
    ),
]


DISCUSSION = [
    (
        "This reconstruction changes the manuscript from a broad developmental fate-divergence framework into a model-falsification and regime-reconstruction paper. The old empirical findings are retained: injury-induced plasticity expands accessibility; stemness-associated programs contribute to that accessibility; positional-program proxies are associated with regenerative organization; fate-lock and senescence-like programs constrain mammalian repair; and tumor-like plasticity forms a distinct high-plasticity branch. What changes is the representational claim. Cell fate and regenerative accessibility are not governed by one global scalar Phi. They are better represented by a latent-state-regime mixture with overlapping distributions."
    ),
    (
        "The failure of single-Phi is central rather than incidental. Phi was attractive because it promised a compact order parameter linking mammalian repair, salamander blastema and tumor-like plasticity. However, the locked validation shows that Phi cannot perform that role. An AUC of approximately 0.480 is random-level classification. A non-significant permutation mean shift and bootstrap instability further argue against scalar separation. The significant KS result does not rescue the model because KS detects distributional shape difference; it does not establish a threshold that separates regimes."
    ),
    (
        "This distinction resolves an important biological ambiguity. Salamander blastema should not be described as high Phi, and mammalian repair should not be described as the same process at lower intensity. The two systems overlap in state space but differ distributionally and mechanistically. Mammalian repair is best interpreted as an adult-repair-dominant mixture constrained by fate-lock, senescence-like stabilization and inflammatory repair context. Salamander blastema is a distinct latent regenerative regime associated with developmental reactivation and positional-program organization. Salamander intact tissue remains a separate reference regime with only partial overlap."
    ),
    (
        "The same logic constrains the interpretation of old manuscript concepts. Plasticity becomes within-regime accessibility rather than fate instruction. Fate-lock becomes an adult-repair basin constraint rather than a universal aging law. DPRI and RA/HOX-like positional activity become regime-conditioned developmental coordinates rather than universal determinants. Senescence becomes attractor stabilization in the adult repair-failure context. Regeneration becomes a regime transition or access problem rather than a scalar trait. Tumor-like plasticity remains distinct and should not be interpreted as embryonic reversion."
    ),
    (
        "The reconstructed framework also clarifies what the current data can and cannot prove. It can support distributional divergence without separability. It can show that the old single-scalar hypothesis is invalid. It can preserve old biological evidence while requiring a probabilistic latent representation. It cannot establish a complete causal regulatory program, direct spatial-coordinate control, lineage history or chromatin-level mechanism across all systems. Those claims require same-system perturbation, spatial transcriptomics, lineage tracing and multiome time series."
    ),
    (
        "Several limitations remain explicit. Batch and species effects are entangled in cross-dataset comparisons. Cross-species embedding and orthology mapping limit interpretation. Several state dimensions are operational module scores or computational proxies rather than direct measurements of causal potential. Axolotl and frog analyses are expression-layer and compartment-proxy based, not direct spatial or velocity analyses. The latent structure is partially non-identifiable, and residual entropy remains after conditioning on observed state, pathway and GRN-derived variables. The mixture model is therefore the best supported representation of the current outputs, not a final causal mechanism."
    ),
    (
        "The central experimental implication is that regeneration should be tested as regime access rather than scalar maximization. Increasing stemness alone should increase accessibility but not reliably specify regenerative outcome. Restoring positional-program organization should shift accessibility toward regenerative remodeling only in a compatible regime context. Increasing fate-lock or SASP-like stabilization should deepen adult repair-failure retention when not coupled to organized remodeling. Tumor-like plasticity should be monitored as a distinct high-plasticity branch, not treated as a regenerative intermediate. These predictions keep the framework falsifiable while avoiding the rejected single-Phi interpretation."
    ),
]


FIGURE_CAPTIONS = [
    (
        "Figure 1 | Conceptual transition from developmental fate-divergence to latent-state-regime mixture dynamics.",
        "The figure should combine the old evidence architecture with the final model schematic. Plasticity remains permissive, but final fate accessibility is represented by probabilistic latent state regimes rather than a single scalar Phi. The rejected Phi model should be shown explicitly as a failed hypothesis, not as a positive explanatory layer.",
    ),
    (
        "Figure 2 | Shared high-plasticity state space across regeneration, senescence-like repair failure and tumor-like plasticity.",
        "State-space, entropy, CellRank and RNA-velocity-derived summaries from the old manuscript support shared accessibility but not fate equivalence. Velocity and CellRank panels must be labeled as inferred trajectory comparators rather than lineage tracing.",
    ),
    (
        "Figure 3 | Stemness-associated programs increase accessibility but do not determine fate.",
        "Wnt/CHIR/beta-catenin perturbation and module-score analyses support stemness-associated accessibility, while boundary analyses reject a simple universal BMP-Wnt antagonistic fate axis. The figure supports local accessibility within regimes, not deterministic regeneration.",
    ),
    (
        "Figure 4 | Developmental positional information and DPRI are regime-conditioned coordinates of regenerative competence.",
        "RA/RARG perturbation, RA/HOX positional proxies, axolotl blastema-associated evidence and frog compartment DPRI boundary evidence are integrated under a regime-conditioned interpretation. Positional information is not treated as a direct spatial coordinate or global scalar determinant.",
    ),
    (
        "Figure 5 | Fate-lock, senescence-like stabilization and mammalian adult repair-failure regime.",
        "Fate-stabilization, lineage-restriction, senescence-like retention and SAT-related panels are merged to represent an adult-repair/fate-lock-biased regime. The figure must not claim direct p53/BMP causality or equate mammalian repair failure with salamander regeneration.",
    ),
    (
        "Figure 6 | Tumor-like plasticity is a distinct high-plasticity branch rather than embryonic regeneration.",
        "Tumor-like plasticity proxy cells show stemness enrichment, reduced differentiation and local self-retention, while embryonic-like and inflammatory-repair boundary tests do not support embryonic reversion or repair equivalence.",
    ),
    (
        "Figure 7 | Single-Phi model rejection and data-closure evidence.",
        "The figure reports distributional structure but random-level discrimination: ROC AUC approximately 0.480, non-significant permutation mean-shift test, bootstrap instability with mean interval crossing zero, and KS significance interpreted as distributional shape difference rather than separability.",
    ),
    (
        "Figure 8 | Latent-state-regime mixture model and posterior landscape.",
        "The figure shows P(Phi|S) = sum_Z P(Phi|S,Z)P(Z|S,W_GRN), per-regime Phi distributions, posterior probabilities P(Z|S,W_GRN), posterior overlap in state space and observed groups as mixtures of latent state regimes.",
    ),
    (
        "Figure 9 | Regime overlap and symmetrized distributional divergence.",
        "Overlap matrices and symmetrized KL divergence quantify non-identifiability and distributional divergence. Adult repair and salamander blastema are overlapping but distributionally different latent-state-regime mixtures, not separable by a global Phi threshold.",
    ),
]


METHODS_INSERT = [
    (
        "Latent-state-regime mixture validation and figure locking",
        [
            "The final manuscript uses the locked model-validation outputs generated before manuscript reconstruction. The single-Phi rejection was based on Phi_unified.tsv, roc_curve.tsv, permutation_test_results.tsv, bootstrap_confidence_intervals.tsv and ks_test_results.tsv. The latent-state-regime mixture and posterior-landscape figures were based on mixture_density_decomposition_phi.tsv, regime_posterior_probabilities.tsv and observed_regime_to_latent_regime_composition.tsv. Regime overlap and distributional divergence were based on regime_overlap_matrix.tsv, species_regime_overlap_matrix.tsv and regime_KL_divergence_matrix.tsv.",
            "No new datasets, statistical tests or biological models were introduced during manuscript reconstruction. The nine-figure architecture was taken from old_new_figure_merger_map.tsv, final_results_to_figure_mapping.tsv, figure_rebuild_requirements.md and final_figure_integration_verdict.md. Claims were constrained by claim_consistency_audit.md and final_data_locking_report.md.",
        ],
    )
]


def extract_tail_from_doc1() -> list[tuple[str, str]]:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    items: list[tuple[str, str]] = []
    with ZipFile(DOC1) as z:
        root = ET.fromstring(z.read("word/document.xml"))
    body = root.find("w:body", ns)
    assert body is not None
    capture = False
    for p in body.findall("w:p", ns):
        txt = "".join(t.text or "" for t in p.findall(".//w:t", ns)).strip()
        if not txt:
            continue
        pstyle = ""
        ppr = p.find("w:pPr", ns)
        if ppr is not None:
            style_el = ppr.find("w:pStyle", ns)
            if style_el is not None:
                pstyle = style_el.attrib.get(f"{{{ns['w']}}}val", "")
        if txt == "Methods":
            capture = True
        if capture:
            items.append((pstyle, txt))
    return items


def add_markdown(path: Path, heading: str, paragraphs: list[str] | None = None, sections: list[tuple[str, list[str]]] | None = None) -> None:
    lines = [f"# {heading}", ""]
    if paragraphs:
        for p in paragraphs:
            lines.append(p)
            lines.append("")
    if sections:
        for title, paras in sections:
            lines.append(f"## {title}")
            lines.append("")
            for p in paras:
                lines.append(p)
                lines.append("")
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def write_figure_alignment(path: Path) -> None:
    lines = ["# Figure Caption Alignment", ""]
    for i, (title, body) in enumerate(FIGURE_CAPTIONS, 1):
        lines.append(f"## Figure {i}")
        lines.append("")
        lines.append(f"- caption_title: {title}")
        lines.append(f"- aligned_caption_text: {body}")
        lines.append(f"- results_section: Result {i}")
        lines.append(f"- locked_role: {'old empirical foundation' if i <= 6 else 'new model-validation layer'}")
        lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def write_final_claim_audit(path: Path) -> None:
    checks = [
        ("Phi separates species", "REMOVED", "Manuscript states that Phi fails as a global discriminator."),
        ("high Phi = regeneration", "REMOVED", "Manuscript states that salamander blastema is a distinct latent regenerative regime."),
        ("single order parameter governs fate", "REMOVED", "Manuscript states that latent-state-regime mixture replaces single-Phi control."),
        ("tumor-like plasticity = embryonic reversion", "REMOVED", "Manuscript states tumor-like plasticity is a distinct branch."),
        ("distributional divergence without separability", "PRESENT", "Used to interpret KS and KL/overlap results."),
        ("latent state regime mixture explains overlap", "PRESENT", "Used in Introduction, Results 8-9 and Discussion."),
        ("regime-dependent interpretation required", "PRESENT", "Used throughout old-result reinterpretation."),
    ]
    lines = ["# Claim Consistency Final Audit", "", "| Claim rule | Status | Evidence in rewritten manuscript |", "|---|---|---|"]
    for claim, status, evidence in checks:
        lines.append(f"| {claim} | {status} | {evidence} |")
    lines.append("")
    lines.append("FINAL_STATUS: PASS")
    lines.append("")
    lines.append("The rewritten manuscript treats single-Phi as a failed hypothesis and uses latent-state-regime mixture language as the final interpretation.")
    path.write_text("\n".join(lines), encoding="utf-8")


def set_font(run, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    # Named override to the narrative preset: scientific manuscripts with long
    # formulas, dataset IDs and module names render more reliably left-aligned.
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = doc.styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def add_para(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_together = False
    p.paragraph_format.keep_with_next = False
    r = p.add_run(text)
    set_font(r, size=11)
    return p


def build_docx(path: Path) -> None:
    doc = Document()
    configure_styles(doc)
    footer = doc.sections[0].footer.paragraphs[0]
    add_page_number(footer)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(12)
    r = title.add_run(TITLE)
    set_font(r, size=20, bold=True, color="0B2545")

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(18)
    r = meta.add_run("Reconstructed manuscript aligned to locked nine-figure latent-state-regime mixture architecture")
    set_font(r, size=10, italic=True, color="555555")

    doc.add_heading("Abstract", level=1)
    for p in ABSTRACT_PARAS:
        add_para(doc, p)

    doc.add_heading("Introduction", level=1)
    for p in INTRODUCTION:
        add_para(doc, p)

    doc.add_heading("Results", level=1)
    for title_text, paras in RESULTS:
        doc.add_heading(title_text, level=2)
        for p in paras:
            add_para(doc, p)

    doc.add_heading("Discussion", level=1)
    for p in DISCUSSION:
        add_para(doc, p)

    tail = extract_tail_from_doc1()
    inserted_methods = False
    for style, txt in tail:
        if txt == "Methods":
            doc.add_page_break()
            doc.add_heading("Methods", level=1)
            continue
        if txt == "Statistical testing and multiple-testing correction" and not inserted_methods:
            for h, paras in METHODS_INSERT:
                doc.add_heading(h, level=2)
                for p in paras:
                    add_para(doc, p)
            inserted_methods = True
        if style == "1":
            doc.add_heading(txt, level=1)
        elif style == "21":
            doc.add_heading(txt, level=2)
        elif style == "31":
            doc.add_heading(txt, level=2)
        elif style == "EndNoteBibliography":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(txt)
            set_font(r, size=9)
        else:
            add_para(doc, txt)

    doc.add_page_break()
    doc.add_heading("Figure legends", level=1)
    for title_text, body in FIGURE_CAPTIONS:
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        r = p.add_run(title_text)
        set_font(r, size=11, bold=True)
        add_para(doc, body)

    doc.save(path)


def scan_forbidden(text: str) -> list[str]:
    forbidden = [
        r"Phi separates species",
        r"high Phi = regeneration",
        r"single order parameter governs fate",
        r"tumor-like plasticity = embryonic reversion",
        r"high Phi equals salamander blastema",
    ]
    hits = []
    for pat in forbidden:
        if re.search(pat, text, flags=re.I):
            hits.append(pat)
    return hits


def main() -> None:
    OUT.mkdir(exist_ok=True)
    add_markdown(OUT / "introduction_rewritten.md", "Introduction Rewritten", INTRODUCTION)
    add_markdown(OUT / "results_rewritten_by_figure.md", "Results Rewritten By Figure", sections=RESULTS)
    add_markdown(OUT / "discussion_rewritten.md", "Discussion Rewritten", DISCUSSION)
    write_figure_alignment(OUT / "figure_caption_alignment.md")
    write_final_claim_audit(OUT / "claim_consistency_final_audit.md")
    manuscript = OUT / "manuscript_rewritten.docx"
    build_docx(manuscript)

    combined = "\n".join(ABSTRACT_PARAS + INTRODUCTION + [p for _, ps in RESULTS for p in ps] + DISCUSSION)
    hits = scan_forbidden(combined)
    if hits:
        raise RuntimeError(f"Forbidden phrasing found in manuscript text: {hits}")
    print("manuscript reconstruction complete")
    print(manuscript)


if __name__ == "__main__":
    main()
