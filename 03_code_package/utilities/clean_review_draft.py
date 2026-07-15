#!/usr/bin/env python3
"""Create a clean review draft from the inline-figure manuscript."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path.cwd()
INPUT = ROOT / "Regenerative_cell_fate_latent_regime_manuscript_WITH_INLINE_FIGURES.docx"
OUT_DOCX = ROOT / "Regenerative_cell_fate_latent_regime_CLEAN_REVIEW_DRAFT.docx"
OUT_MD = ROOT / "Regenerative_cell_fate_latent_regime_CLEAN_REVIEW_DRAFT.md"
CHANGE_LOG = ROOT / "clean_draft_change_log.md"


FORBIDDEN_SOFTENINGS = {
    "proves": "supports",
    "proven": "supported",
    "demonstrates causality": "is consistent with computational causal-consistency",
    "fully explains": "provides a computational framework for",
    "establishes universally": "supports in the evaluated setting",
    "confirms": "is consistent with",
    "solves regeneration": "provides a computational framework for regeneration studies",
    "guarantees rejuvenation": "suggests a testable direction for rejuvenation studies",
    "causal proof": "computational causal-consistency support",
    "experimentally validated causality": "dry-lab perturbation-consistency support",
    "causal closure established": "partial computational causal-consistency support",
}


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing | .//w:pict"))


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def replace_text_preserving_runs(paragraph, replacements: list[tuple[str, str]]) -> int:
    """Replace paragraph text by rebuilding runs only when needed.

    This is used only for text-only paragraphs; drawing paragraphs are left intact.
    """

    original = paragraph.text
    updated = original
    count = 0
    for old, new in replacements:
        new_updated = updated.replace(old, new)
        if new_updated != updated:
            count += updated.count(old)
            updated = new_updated
    if updated == original:
        return 0

    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = updated
    else:
        paragraph.add_run(updated)
    return count


def callout_replacements(text: str) -> list[tuple[str, str]]:
    pairs: list[tuple[str, str]] = []
    if "Supplementary Figure 1" in text:
        pairs.append(("Supplementary Figure 1", "Supplementary Fig. 1"))
    if "Figure 1A" in text:
        pairs.append(("Figure 1A", "Fig. 1A"))
    for n in range(1, 10):
        if f"Figure {n}" in text:
            pairs.append((f"Figure {n}", f"Fig. {n}"))
    return pairs


def standardize_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles_by_name = {style.name: style for style in doc.styles}

    def get_style(style_name: str):
        return styles_by_name.get(style_name)

    def set_font(style_name: str, size: float, bold: bool | None = None) -> None:
        style = get_style(style_name)
        if style is None:
            return
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        if bold is not None:
            style.font.bold = bold

    for style_name in ["Normal", "Body Text", "First Paragraph", "Compact"]:
        style = get_style(style_name)
        if style is not None:
            style.font.name = "Arial"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            style.font.size = Pt(10.5)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing = 1.08

    set_font("Heading 1", 16, True)
    set_font("Heading 2", 13, True)
    set_font("Heading 3", 11, True)

    if get_style("Heading 1") is not None:
        get_style("Heading 1").paragraph_format.space_before = Pt(0)
        get_style("Heading 1").paragraph_format.space_after = Pt(10)
    if get_style("Heading 2") is not None:
        get_style("Heading 2").paragraph_format.space_before = Pt(12)
        get_style("Heading 2").paragraph_format.space_after = Pt(6)
    if get_style("Heading 3") is not None:
        get_style("Heading 3").paragraph_format.space_before = Pt(8)
        get_style("Heading 3").paragraph_format.space_after = Pt(4)


def clean_docx() -> tuple[int, int, int]:
    doc = Document(INPUT)
    standardize_styles(doc)

    in_figure_legends = False
    callout_edits = 0
    softened = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text == "Figure legends":
            in_figure_legends = True
        elif text == "References":
            in_figure_legends = False

        if paragraph.style.name.startswith("Heading"):
            if paragraph.style.name == "Heading 1":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        if paragraph_has_drawing(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(4)
            continue

        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.08

        if not text:
            continue

        # Main-text figure callouts only; figure legend headings remain "Figure X |".
        if not in_figure_legends:
            callout_edits += replace_text_preserving_runs(
                paragraph, callout_replacements(paragraph.text)
            )

        # Light claim-safety softening only when explicitly overstrong phrases occur.
        lowered = paragraph.text.lower()
        replacements = []
        for old, new in FORBIDDEN_SOFTENINGS.items():
            if old in lowered:
                pattern = re.compile(re.escape(old), re.IGNORECASE)
                new_text = pattern.sub(new, paragraph.text)
                if new_text != paragraph.text:
                    replacements.append((paragraph.text, new_text))
                    break
        if replacements:
            softened += replace_text_preserving_runs(paragraph, replacements)

    # Remove pure duplicated blank paragraphs without touching image paragraphs.
    removed_blank = 0
    previous_blank = False
    for paragraph in list(doc.paragraphs):
        is_blank = not paragraph.text.strip() and not paragraph_has_drawing(paragraph)
        if is_blank and previous_blank:
            delete_paragraph(paragraph)
            removed_blank += 1
        previous_blank = is_blank

    doc.save(OUT_DOCX)
    return callout_edits, softened, removed_blank


def iter_markdown_blocks(docx_path: Path) -> list[str]:
    doc = Document(docx_path)
    blocks: list[str] = []
    image_counter = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style = paragraph.style.name
        has_image = paragraph_has_drawing(paragraph)
        if has_image:
            image_counter += 1
            blocks.append(f"<!-- Inline figure retained in DOCX: image {image_counter} -->")
            continue
        if not text:
            continue
        if style == "Heading 1":
            blocks.append(f"# {text}")
        elif style == "Heading 2":
            blocks.append(f"## {text}")
        elif style == "Heading 3":
            blocks.append(f"### {text}")
        elif style == "Compact":
            blocks.append(f"- {text}")
        else:
            blocks.append(text)
    return blocks


def write_markdown() -> None:
    blocks = iter_markdown_blocks(OUT_DOCX)
    OUT_MD.write_text("\n\n".join(blocks) + "\n", encoding="utf-8")


def write_change_log(callout_edits: int, softened: int, removed_blank: int) -> None:
    CHANGE_LOG.write_text(
        "\n".join(
            [
                "# Clean Draft Change Log",
                "",
                "- Standardized heading hierarchy, paragraph spacing and figure-placeholder paragraph alignment.",
                f"- Standardized main-text figure callouts to `Fig. X` / `Supplementary Fig. 1` where applicable (`{callout_edits}` replacements).",
                f"- Removed duplicated blank paragraphs where present (`{removed_blank}` removed).",
                f"- Applied claim-safety softening only for explicitly overstrong wording (`{softened}` replacements).",
                "- Preserved all `[REF::...]` citation slots; no real references, DOI, PMID or fabricated bibliography entries were added.",
                "- Kept all inline figures in the DOCX and kept the Figure legends section separated from the main text.",
                "",
                "Remaining before true submission final:",
                "",
                "- Manually resolve all `[REF::...]` citation slots against verified reference metadata.",
                "- Perform journal-specific figure upload/export checks and final source-data packaging.",
                "- Fill acknowledgements, author contributions and competing-interest placeholders.",
                "",
            ]
        ),
        encoding="utf-8",
    )


def main() -> None:
    callout_edits, softened, removed_blank = clean_docx()
    write_markdown()
    write_change_log(callout_edits, softened, removed_blank)


if __name__ == "__main__":
    main()
